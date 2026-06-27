# -*- coding: utf-8 -*-
"""Generate 2026-06-28 World Cup predictions as a formatted .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json, os

OUT = r"D:\ai\世界杯\2026-worldcup-predictions\2026年6月28日_6场预测.docx"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# colour helpers
DARK  = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
ACCENT= RGBColor(0xC0, 0x39, 0x2B)   # red
GOLD  = RGBColor(0xD4, 0xA0, 0x17)   # gold
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0x27, 0xAE, 0x60)
RED_BG     = RGBColor(0xE7, 0x4C, 0x3C)

# ─── helper functions ───
def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    """Set cell text and formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)[2:]}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # vertical centering
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def set_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" '
            f'w:color="{val.get("color","CCCCCC")}"/>')
        borders.append(element)
    tcPr.append(borders)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('6月28日 六场预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('小组赛收官日  ·  L / K / J 组  ·  第3轮', bold=False, size=Pt(11), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(18))

# meta line
add_para('生成时间: 2026年6月28日 北京时间 01:00  |  数据源: FIFA API + ESPN API  |  框架: CLAUDE.md v17',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# ═══════════════════════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════════════════════
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '时间', '小组', '比赛', '形势', '首选比分', '半场', '备选比分', '冷门\n风险'],
    ['1', '05:00', 'L', '🇭🇷 克罗地亚 vs 🇬🇭 加纳', '克3分/加4分 生死战', '1-1', '0-0', '克1-0 / 加1-0', '中'],
    ['2', '05:00', 'L', '🇵🇦 巴拿马 vs 🏴󠁧󠁢󠁥󠁮󠁧󠁿 英格兰', '巴淘汰/英4分', '英格兰 3-0', '1-0', '2-0 / 4-0', '极低'],
    ['3', '07:30', 'K', '🇨🇴 哥伦比亚 vs 🇵🇹 葡萄牙', '哥6分/葡4分 争头名', '1-1', '0-0', '葡2-1 / 哥1-0', '低'],
    ['4', '07:30', 'K', '🇨🇩 刚果金 vs 🇺🇿 乌兹别克', '刚1分/乌淘汰', '刚果金 2-0', '1-0', '1-0 / 3-0', '低'],
    ['5', '10:00', 'J', '🇩🇿 阿尔及利亚 vs 🇦🇹 奥地利', '双方3分 争第2', '0-0', '0-0', '奥1-0 / 阿1-0', '中高'],
    ['6', '10:00', 'J', '🇯🇴 约旦 vs 🇦🇷 阿根廷', '约淘汰/阿6分头名', '阿根廷 3-0', '2-0', '2-0 / 4-0', '极低'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# column widths in inches (landscape: ~10.5" usable)
col_widths = [0.28, 0.45, 0.28, 2.10, 1.50, 0.75, 0.55, 1.20, 0.50]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

# header row
for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

# data rows
for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        col = summary_data[0][j]
        align = 'left' if j == 3 else 'center'
        bold = (j == 5)  # bold the primary prediction
        if j == 3:
            set_cell_left(table.rows[i].cells[j], text, bold=False, size=Pt(9))
        else:
            set_cell(table.rows[i].cells[j], text, bold=bold, size=Pt(9), bg=bg)

# cold column highlight
cold_cells = [(5, 8, RED_BG, WHITE), (6, 8, RED_BG, WHITE), (3, 8, RED_BG, WHITE)]
for ri, ci, cbg, ccol in cold_cells:
    pass  # will highlight below

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. METRICS
# ═══════════════════════════════════════════════════
add_heading_styled('战绩指标（6月27日回顾）', level=2)
metrics = [['指标', '数值'],
           ['方向正确率', '5/6 (83%)'],
           ['首选比分命中', '1/6 (17%)'],
           ['备选比分命中', '2/2 (100%)'],
           ['差距≤1球', '4/6 (67%)']]
mt = doc.add_table(rows=len(metrics), cols=2)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(metrics[0]):
    set_cell(mt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(metrics[1:], 1):
    set_cell(mt.rows[i].cells[0], row[0], bold=True, size=Pt(9))
    set_cell(mt.rows[i].cells[1], row[1], size=Pt(9))
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. GROUP STANDINGS
# ═══════════════════════════════════════════════════
add_heading_styled('小组积分（赛前）', level=2)

groups_data = {
    'L组': [
        ['#', '球队', '场', '胜', '平', '负', '进/失', '净胜', '积分', '状态'],
        ['1', '英格兰', '2', '1', '1', '0', '4/2', '+2', '4', '✓'],
        ['2', '加纳', '2', '1', '1', '0', '1/0', '+1', '4', '✓'],
        ['3', '克罗地亚', '2', '1', '0', '1', '3/4', '-1', '3', '—'],
        ['4', '巴拿马', '2', '0', '0', '2', '0/2', '-2', '0', 'E'],
    ],
    'K组': [
        ['#', '球队', '场', '胜', '平', '负', '进/失', '净胜', '积分', '状态'],
        ['1', '哥伦比亚', '2', '2', '0', '0', '4/1', '+3', '6', 'Q'],
        ['2', '葡萄牙', '2', '1', '1', '0', '6/1', '+5', '4', 'Q'],
        ['3', '刚果金', '2', '0', '1', '1', '1/2', '-1', '1', '—'],
        ['4', '乌兹别克', '2', '0', '0', '2', '1/8', '-7', '0', 'E'],
    ],
    'J组': [
        ['#', '球队', '场', '胜', '平', '负', '进/失', '净胜', '积分', '状态'],
        ['1', '阿根廷', '2', '2', '0', '0', '5/0', '+5', '6', 'Q'],
        ['2', '奥地利', '2', '1', '0', '1', '3/3', '0', '3', '—'],
        ['3', '阿尔及利亚', '2', '1', '0', '1', '2/4', '-2', '3', '—'],
        ['4', '约旦', '2', '0', '0', '2', '2/5', '-3', '0', 'E'],
    ],
}

for gname, rows in groups_data.items():
    add_para(gname, bold=True, size=Pt(11), color=DARK)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(rows[0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(rows[1:], 1):
        for j, text in enumerate(row):
            c = 'center'
            if j == 1: c = 'left'
            bg = LIGHT_GRAY if i % 2 == 0 else None
            set_cell(t.rows[i].cells[j], text, size=Pt(8), bg=bg)
            if j == 1:
                set_cell_left(t.rows[i].cells[j], text, size=Pt(8))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 4. BRACKET PATHS
# ═══════════════════════════════════════════════════
add_heading_styled('淘汰赛路径', level=2)

bracket_data = [
    ['组', '第1名', '第2名', '第1 → 对手', '第2 → 对手', '第1难度', '第2难度'],
    ['L', '英格兰/加纳', '加纳/克罗地亚', '3rd-place 🟢', '葡萄牙(K2) 🔴', '可打', '死亡级'],
    ['K', '哥伦比亚/葡萄牙', '葡萄牙/哥伦比亚', '3rd-place 🟢', 'L2(英/加/克) ⚠️', '可打', '中等'],
    ['J', '阿根廷(已锁)', '奥/阿争', '3rd-place 🟢', '🇪🇸 西班牙(H1) 💀', '可打', '死亡级'],
]
bt = doc.add_table(rows=len(bracket_data), cols=len(bracket_data[0]))
bt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(bracket_data[0]):
    set_cell(bt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(bracket_data[1:], 1):
    for j, text in enumerate(row):
        c = 'left' if j in [1,2,3,4] else 'center'
        color = RED_BG if '💀' in text else None
        fg = WHITE if color else None
        if c == 'left':
            set_cell_left(bt.rows[i].cells[j], text, size=Pt(8), color=fg)
        else:
            set_cell(bt.rows[i].cells[j], text, size=Pt(8), color=fg, bg=color)
        if color:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(color)[2:]}"/>')
            bt.rows[i].cells[j]._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 5. MATCH DETAILS
# ═══════════════════════════════════════════════════

matches = [
    {
        'num': 1, 'time': '05:00', 'group': 'L',
        'home': '克罗地亚', 'away': '加纳',
        'home_flag': '🇭🇷', 'away_flag': '🇬🇭',
        'situation': '克罗地亚3分(须赢才稳) / 加纳4分(打平出线)  |  L2=葡萄牙💀→双方都想避免',
        'venue': '—',
        'factors': [
            ('加纳已验证大巴(0-0英格兰)', '加纳 🔥🔥', '加纳大巴是本届顶级 — 英格兰78%控球=0进球'),
            ('淘汰赛路径: L2=葡萄牙🔴', '双方都不想第2', '打平=加纳L1→安全路径'),
            ('莫德里奇回归首发', '克罗地亚 🔥', '前两场轮换→本场全主力'),
            ('克罗地亚防线老化(4失球)', '加纳', '速度型反击是克罗地亚最怕的'),
            ('加纳自主进攻极弱(2场1球)', '克罗地亚', '加纳0-0英格兰=极品大巴但00进攻'),
        ],
        'strength': '克罗地亚: 体系型(莫德里奇单人超巨年龄折扣)  |  加纳: 大巴型(已验证0-0英格兰)',
        'africa_resilience': [('低位防守', '✅✅ 0-0英格兰'), ('速度反击', '✅✅ 耶伦基'), ('前30分钟', '✅✅ 2场均不丢球'), ('被压制不崩盘', '✅✅ 对英格兰全场'), ('定位球', '✅')],
        'upset_risk': '中 — 加纳大巴已在本届验证。克罗地亚必须赢但可能攻不破',
        'predictions': [('首选', '1-1', '0-0', '克罗地亚先进(莫德里奇定位球)→加纳反击扳平→双方默契保持'),
                        ('备选', '克罗地亚 1-0', '0-0', '莫德里奇定位球制胜'),
                        ('备选', '加纳 1-0', '0-0', '耶伦基反击偷一个→大巴死守')],
    },
    {
        'num': 2, 'time': '05:00', 'group': 'L',
        'home': '巴拿马', 'away': '英格兰',
        'home_flag': '🇵🇦', 'away_flag': '🏴󠁧󠁢󠁥󠁮󠁧󠁿',
        'situation': '巴拿马0分已淘汰 / 英格兰4分(打平即头名)  |  身价比 英格兰€1.31B vs 巴拿马~€15M ≈ 87:1',
        'venue': '—',
        'factors': [
            ('身价比87:1', '英格兰 🔥🔥🔥', '极端差距—强队方不做小比分预测'),
            ('巴拿马2场0进球已淘汰', '英格兰', '攻防能力近零'),
            ('英格兰可能轮换保护主力', '巴拿马 微弱', '但替补仍是碾压级'),
            ('英格兰G3反弹(已过0-0加纳综合征)', '英格兰', 'G2低迷→G3反弹是规律'),
        ],
        'strength': '英格兰: 体系型(凯恩1人超巨→缺少第二破局点) 但身价碾压  |  巴拿马: 不适用',
        'africa_resilience': None,
        'upset_risk': '极低 — 身价87:1+对手已淘汰+0进球能力',
        'predictions': [('首选', '英格兰 3-0', '1-0', '身价碾压→轮换替补仍碾压'),
                        ('备选', '英格兰 2-0', '1-0', '保守轮换不影响结果'),
                        ('备选', '英格兰 4-0', '2-0', '凯恩刷进球→金靴冲刺')],
    },
    {
        'num': 3, 'time': '07:30', 'group': 'K',
        'home': '哥伦比亚', 'away': '葡萄牙',
        'home_flag': '🇨🇴', 'away_flag': '🇵🇹',
        'situation': '哥伦比亚6分(Q) / 葡萄牙4分(Q) — 双方均已出线，争头名  |  身价 葡萄牙€957M vs 哥伦比亚~€250M',
        'venue': '—',
        'factors': [
            ('哥伦比亚打平即头名', '哥伦比亚 🔥', '平局=K1=3rd路径🟢→不冒险'),
            ('葡萄牙必须赢才能头名', '哥伦比亚 🔥', '葡压上→哥反击→萨拉赫式剧本'),
            ('身价差3.8:1(葡萄牙€957M)', '葡萄牙 🔥', 'C罗+B费+莱昂=超级巨星型'),
            ('哥伦比亚2场仅丢1球', '哥伦比亚', '防线稳定'),
            ('K2=L2(英/加/克)→非死亡级', '双方都能接受', '平局双赢→K2对手可控'),
        ],
        'strength': '葡萄牙: 超级巨星型(C罗+B费+莱昂≥3超巨)  |  哥伦比亚: 体系型(仅迪亚斯一个超巨)',
        'africa_resilience': None,
        'upset_risk': '低 — 双方已出线→平局双赢→默契球概率高',
        'predictions': [('首选', '1-1', '0-0', '哥打平=头名→C罗破僵→哥扳平→默契保持'),
                        ('备选', '哥伦比亚 1-0', '0-0', '葡萄牙围攻无果→哥偷袭'),
                        ('备选', '葡萄牙 2-1', '1-0', 'C罗+B费破局→哥扳回一球')],
    },
    {
        'num': 4, 'time': '07:30', 'group': 'K',
        'home': '刚果金', 'away': '乌兹别克',
        'home_flag': '🇨🇩', 'away_flag': '🇺🇿',
        'situation': '刚果金1分(可争第3出线) / 乌兹别克0分已淘汰  |  刚果金赢球=4分第3大概率晋级',
        'venue': '—',
        'factors': [
            ('刚果金防守验证(仅2失球对葡+哥)', '刚果金 🔥🔥', '姆本巴+图安泽贝=世界杯级防线'),
            ('乌兹别克-7净胜球+0分已淘汰', '刚果金 🔥', '已无动力→荣誉战积极性存疑'),
            ('刚果金赢球可能出线(4分第3)', '刚果金 🔥', '有明确目标→动力碾压'),
            ('乌兹别克对葡萄牙5球惨败', '刚果金', '防线薄弱→刚果金可轻松破门'),
        ],
        'strength': '刚果金: 防守型(对比葡仅丢2球)  |  乌兹别克: 不适用',
        'africa_resilience': [('低位防守', '✅✅ 对比葡仅丢2球'), ('速度反击', '✅ 维萨'), ('前30分钟', '✅ 均未丢球')],
        'upset_risk': '低 — 刚果金防守顶级+有明确目标',
        'predictions': [('首选', '刚果金 2-0', '1-0', '防守稳+对手无动力→刚果金有目标'),
                        ('备选', '刚果金 1-0', '0-0', '小胜即可→保守拿分'),
                        ('备选', '刚果金 3-0', '1-0', '乌兹别克防线崩溃→大胜')],
    },
    {
        'num': 5, 'time': '10:00', 'group': 'J',
        'home': '阿尔及利亚', 'away': '奥地利',
        'home_flag': '🇩🇿', 'away_flag': '🇦🇹',
        'situation': '双方均3分 — 争J组第2  |  🚨 J2 = 西班牙(€1.27B) = 死亡路径! 双方可能都不想赢!',
        'venue': '—',
        'factors': [
            ('🚨 J2=西班牙(€1.27B)=死刑路径', '双方都不想第2! 🔥🔥🔥', 'J2对西班牙=提前出局——类似挪威放弃争头名'),
            ('奥地利净胜球+0 vs -2', '奥地利', '打平=奥地利第2→但没人想要这个"奖"'),
            ('双方打平都出线', '双方', '奥J2+阿J3(4分)→都大概率晋级'),
            ('可参考挪威轮换放弃争头名', '—', '挪威哈兰德+厄德高替补→法国碾压→先例'),
        ],
        'strength': '双方均体系型  |  🚨 与挪威vs法国相同的"头名奖励是死刑"场景',
        'africa_resilience': [('低位防守', '✅ 对阿根廷仅丢3球'), ('速度反击', '✅ 马赫雷斯')],
        'upset_risk': '中高 — J2=西班牙=双方都不想拿第2→0-0是最大公约数',
        'predictions': [('首选', '0-0', '0-0', '🚨 J2=西班牙=双方都不想赢→0-0默契球'),
                        ('备选', '奥地利 1-0', '0-0', '奥地利无法容忍故意输→勉强小胜'),
                        ('备选', '阿尔及利亚 1-0', '0-0', '阿尔及利亚偷赢→"奖励"是面对西班牙💀')],
    },
    {
        'num': 6, 'time': '10:00', 'group': 'J',
        'home': '约旦', 'away': '阿根廷',
        'home_flag': '🇯🇴', 'away_flag': '🇦🇷',
        'situation': '约旦0分已淘汰 / 阿根廷6分已锁J1  |  身价 阿根廷€816.5M vs 约旦€20M ≈ 41:1',
        'venue': '—',
        'factors': [
            ('身价比41:1', '阿根廷 🔥🔥🔥', '极端差距—强队方不做小比分预测'),
            ('约旦已淘汰+3场丢5球', '阿根廷 🔥', '防线无抵抗能力'),
            ('阿根廷可能轮换', '约旦 微弱', '但替补仍是碾压级'),
            ('梅西冲刺金靴(5球)', '阿根廷 🔥', '首发或替补均可刷数据'),
        ],
        'strength': '阿根廷: 超级巨星型(梅西+迪马利亚+劳塔罗≥3超巨)  |  约旦: 不适用',
        'africa_resilience': None,
        'upset_risk': '极低 — 身价41:1+对手已淘汰',
        'predictions': [('首选', '阿根廷 3-0', '2-0', '身价碾压→梅西替补/首发均可刷金靴'),
                        ('备选', '阿根廷 2-0', '1-0', '轮换→小胜'),
                        ('备选', '阿根廷 4-0', '2-0', '梅西首发→帽子戏法冲刺金靴')],
    },
]

# ── Render each match ──
for m in matches:
    add_heading_styled(f'比赛{m["num"]}: {m["home_flag"]} {m["home"]} vs {m["away_flag"]} {m["away"]}  ({m["group"]}组 {m["time"]} BJT)', level=2)

    # info row
    add_para(m['situation'], size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(6))

    # ─ Factors table ─
    add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
    ft_data = [['因素', '有利方', '理由']]
    ft_data.extend(m['factors'])
    ft = doc.add_table(rows=len(ft_data), cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, text in enumerate(ft_data[0]):
        set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    # data
    for i, row in enumerate(ft_data[1:], 1):
        set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8.5))
        clr = None
        if '🔥🔥🔥' in row[1] or '🔥🔥' in row[1]:
            clr = ACCENT
        set_cell(ft.rows[i].cells[1], row[1], size=Pt(8.5), color=clr)
        set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(8.5))
    # column widths
    for row in ft.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(1.8)
        row.cells[2].width = Inches(5.7)

    doc.add_paragraph()

    # ─ Strength classification ─
    add_para(f'强队分类: {m["strength"]}', size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(2))

    # ─ Africa resilience ─
    if m['africa_resilience']:
        add_para('非洲/亚非韧性评估:', bold=True, size=Pt(8.5), space_after=Pt(1))
        for label, val in m['africa_resilience']:
            add_para(f'  {label}: {val}', size=Pt(8.5), color=RGBColor(0x66,0x66,0x66), space_after=Pt(1))

    # ─ Upset risk ─
    add_para(f'冷门风险: {m["upset_risk"]}', bold=True, size=Pt(9), color=ACCENT if '高' in m['upset_risk'] else DARK, space_after=Pt(6))

    # ─ Predictions table ─
    add_para('比分预测', bold=True, size=Pt(10), color=DARK)
    pred_data = [['类型', '比分', '半场', '说明']]
    pred_data.extend(m['predictions'])
    pt = doc.add_table(rows=len(pred_data), cols=4)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(pred_data[0]):
        set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(pred_data[1:], 1):
        bg = None
        fg = None
        if i == 1:  # primary prediction row
            bg = RGBColor(0xE8, 0xF8, 0xF5)  # light green
        set_cell(pt.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg, color=GREEN_BG if i==1 else None)
        set_cell(pt.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
        set_cell(pt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
        set_cell_left(pt.rows[i].cells[3], row[3], size=Pt(8.5))
    for row in pt.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(7.8)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: FIFA API (api.fifa.com) + ESPN API + Transfermarkt via Mundo Deportivo  |  框架: CLAUDE.md v17 (身价量化+强队三类+淘汰赛路径+10项必填清单)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('⚠️ 首发阵容尚未官方确认(距开球>4小时)  |  预测仅供分析参考  |  生成工具: python-docx',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

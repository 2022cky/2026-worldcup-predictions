# -*- coding: utf-8 -*-
"""Generate 6月29日 1场预测 as .docx — 沿用 gen_docx.py 预测v7同款样式."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年6月29日_1场预测.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

DARK  = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT= RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0x27, 0xAE, 0x60)

def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)[2:]}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None, bg=None):
    set_cell(cell, text, bold=bold, size=size, color=color, bg=bg, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════
# COVER
# ═══════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('6月29日 一场预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('淘汰赛32强开幕  ·  南非 vs 加拿大  ·  Round of 32', bold=False, size=Pt(11), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(18))
add_para('生成时间: 2026年6月29日 北京时间 01:30 (首发确认版)  |  数据源: FIFA API + ESPN API + Standard/Khelnow/Al Jazeera/Goal  |  框架: CLAUDE.md v17',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# ═══════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '时间', '阶段', '比赛', '形势', '首选比分', '半场', '备选比分', '冷门\n风险'],
    ['1', '03:00', 'R32', '南非 vs 加拿大', '两队首次淘汰赛\n加拿大€280M vs 南非€45M≈6:1\n戴维斯替补,拉林替补', '1-1\n(加时加拿大晋级)', '0-0', '0-0 / 0-1\n南非 1-0 / 0-2', '中高'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 0.45, 0.35, 2.10, 1.80, 0.85, 0.55, 1.30, 0.55]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        if j == 3:
            set_cell_left(table.rows[i].cells[j], text, size=Pt(8.5))
        elif j == 5:
            set_cell(table.rows[i].cells[j], text, bold=True, size=Pt(8.5), bg=bg)
        elif j == 8:
            clr = ACCENT if '高' in text else None
            set_cell(table.rows[i].cells[j], text, size=Pt(8.5), color=clr, bg=bg)
        else:
            c = 'center'
            set_cell(table.rows[i].cells[j], text, size=Pt(8.5), bg=bg)

doc.add_paragraph()
add_para('6月29日仅1场比赛 — Round of 32淘汰赛开幕战。两队均为队史首次世界杯淘汰赛。',
         size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(12))

# ═══════════════════════════════════
# 2. GROUP STAGE REVIEW
# ═══════════════════════════════════
add_heading_styled('小组赛回顾', level=2)

group_review = [
    ['', '南非 (A组第2, 4分)', '加拿大 (B组第2, 4分)'],
    ['G1', '0-2 墨西哥 (兹瓦内+西索尔红牌)', '1-1 波黑 (拉林扳平)'],
    ['G2', '1-1 捷克 (莫科纳83\'点球)', '6-0 卡塔尔 (戴维帽子戏法,队史最大胜)'],
    ['G3', '1-0 韩国 (马塞科63\'制胜!大巴成功)', '1-2 瑞士 (输掉头名之争)'],
    ['进/失', '2球 / 3球', '8球 / 3球'],
    ['FIFA', '#60', '#31'],
]
gr = doc.add_table(rows=len(group_review), cols=3)
gr.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(group_review[0]):
    set_cell(gr.rows[0].cells[j], text, bold=True, size=Pt(8.5), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(group_review[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell(gr.rows[i].cells[0], row[0], bold=True, size=Pt(8.5), bg=bg)
    set_cell_left(gr.rows[i].cells[1], row[1], size=Pt(8.5), bg=bg)
    set_cell_left(gr.rows[i].cells[2], row[2], size=Pt(8.5), bg=bg)
for row in gr.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(3.3)
    row.cells[2].width = Inches(3.3)
doc.add_paragraph()

# ═══════════════════════════════════
# 3. MATCH DETAIL
# ═══════════════════════════════════
add_heading_styled('比赛: 南非 vs 加拿大  (Round of 32  03:00 BJT)', level=2)

# Info
add_para('SoFi体育场, 洛杉矶, 美国  |  身价比: 加拿大~€280M vs 南非~€45M ≈ 6:1  |  FIFA: #31 vs #60  |  历史: 仅1次友谊赛(2007南非2-0)',
         size=Pt(8.5), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(8))

# ─ Injuries ─
add_para('伤病/停赛', bold=True, size=Pt(10), color=DARK)
idata = [['球队', '球员', '状态', '影响']]
idata.extend([
    ['南非', '兹瓦内(Themba Zwane)', '停赛', '进攻组织核心缺阵'],
    ['南非', '莫科纳(Teboho Mokoena)', '复出首发(C)', '中场核心回归! 队长'],
    ['加拿大', '科内(Ismael Kone)', '报销', '中场核心腿骨折——缺少第二进攻维度'],
    ['加拿大', '阿方索·戴维斯(Alphonso Davies)', '替补!', "本届首次进入名单——65'后可变招"],
    ['加拿大', '欧斯塔基奥(Stephen Eustaquio)', '首发(C)', '肌肉问题无碍——组织核心在位'],
    ['加拿大', '拉林(Cyle Larin)', '替补!', '历史射手王→板凳——马施重大用人决策'],
])
it = doc.add_table(rows=len(idata), cols=4)
it.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(idata[0]):
    set_cell(it.rows[0].cells[j], text, bold=True, size=Pt(8.5), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(idata[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell(it.rows[i].cells[0], row[0], size=Pt(8.5), bg=bg)
    set_cell_left(it.rows[i].cells[1], row[1], size=Pt(8.5), bg=bg)
    clr = ACCENT if row[2] in ['停赛', '报销'] else (GREEN_BG if row[2] == '复出' else None)
    set_cell(it.rows[i].cells[2], row[2], bold=True, size=Pt(8.5), color=clr, bg=bg)
    set_cell_left(it.rows[i].cells[3], row[3], size=Pt(8.5), bg=bg)
for row in it.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(2.2)
    row.cells[2].width = Inches(1.2)
    row.cells[3].width = Inches(3.2)
doc.add_paragraph()

# ─ Factors ─
add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
fdata = [['因素', '有利方', '理由']]
fdata.extend([
    ['身价比6:1(加拿大~€280M)', '加拿大 ★★★', '阵容深度碾压——淘汰赛板凳更重要'],
    ['★乔纳森·戴维状态(3球)', '加拿大 ★★★', '32强中状态最火前锋之一——单兵可破大巴'],
    ['★南非大巴已验证(1-0韩国零封)', '南非 ★★', '32强中最可能拖入加时的防守体系'],
    ['★兹瓦内停赛+西索尔首发→创造力更降', '加拿大 ★★★', '西索尔纯防守型→南非反击发起点减少'],
    ['★戴维斯进名单+拉林替补=超级板凳!', '加拿大 ★★★', "65'后戴维斯+拉林双替补登场→后程碾压"],
    ['★拉林替补→首发终结力下降', '南非 ★', '奥卢瓦塞伊国际大赛先发经验不足'],
    ['★邦比托换德富热罗勒=防线微调', '南非 ★', '中卫搭档变动→默契度不足，马塞科机会'],
    ['★科内报销+艾哈迈德替补=中场换三人', '南非 ★', '与小组赛中场完全不同——默契存疑'],
    ['南非速度反击(马塞科速度是32强顶级)', '南非 ★★', '加拿大压上→身后空间大→反击机会多'],
    ['淘汰赛经验: 双方均为零', '双方', '均首次淘汰赛——都可能紧张,上半场试探为主'],
    ['加拿大东道主+洛杉矶半个主场', '加拿大 ★', 'SoFi球场氛围有利——但不是决定性因素'],
    ['欧斯塔基奥确认首发(C)=组织核心在位', '加拿大 ★★', '肌肉问题无碍→中场控制力稳固'],
])
ft = doc.add_table(rows=len(fdata), cols=3)
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(fdata[0]):
    set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(fdata[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8))
    clr = ACCENT if '★★★' in row[1] else None
    set_cell(ft.rows[i].cells[1], row[1], size=Pt(8), color=clr, bg=bg)
    set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(8), bg=bg)
for row in ft.rows:
    row.cells[0].width = Inches(3.5)
    row.cells[1].width = Inches(1.6)
    row.cells[2].width = Inches(5.4)
doc.add_paragraph()

# ─ Strength class ─
add_para('强队分类:', bold=True, size=Pt(9), color=DARK)
add_para('加拿大: 介于超级巨星型与体系型之间 — 戴维(超巨,€50M+)具备单兵破局能力。首发缺少戴维斯+科内→第二维度不足，但戴维斯+拉林替补待命→65\'后升级双超巨阵容。',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(2))
add_para('南非: 大巴型 — 防守体系稳固,但进攻创造力完全依赖反击速度(马塞科/莫福肯)。兹瓦内停赛=唯一10号被移除。',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(6))

# ─ Africa resilience ─
add_para('非洲韧性评估(南非):', bold=True, size=Pt(9), color=DARK)
rdata = [['韧性维度', '评分', '说明']]
rdata.extend([
    ['低位防守', '[对][对]', '对韩国零封=大巴满分。对捷克仅丢1球。'],
    ['速度反击', '[对][对]', '马塞科速度32强顶级,莫福肯+阿波利斯双翼齐飞。'],
    ['定位球高点', '[对]', '姆博卡齐+奥孔双塔,莫科纳回归增加威胁。'],
    ['前30分钟专注度', '[对][对]', '3场均未在前30\'丢球——开场纪律满分。'],
    ['被压制不崩盘', '[对][对]', '对韩国68%控球围攻不崩盘。少2人打墨西哥例外。'],
])
rt = doc.add_table(rows=len(rdata), cols=3)
rt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(rdata[0]):
    set_cell(rt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(rdata[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell(rt.rows[i].cells[0], row[0], size=Pt(8.5), bg=bg)
    set_cell(rt.rows[i].cells[1], row[1], size=Pt(8.5), color=GREEN_BG, bg=bg)
    set_cell_left(rt.rows[i].cells[2], row[2], size=Pt(8.5), bg=bg)
for row in rt.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(1.2)
    row.cells[2].width = Inches(4.0)
doc.add_paragraph()

add_para('南非韧性 5/5 — 除少2人打墨西哥外,防守体系完整且已验证。', bold=True, size=Pt(9), color=GREEN_BG, space_after=Pt(8))

# ─ Set pieces + Coach ─
add_para('定位球攻防:', bold=True, size=Pt(9), color=DARK)
add_para('南非进攻定位球: 莫科纳回归(对捷克点球得分),姆博卡齐+奥孔双塔。加拿大防守定位球:对瑞士丢过定位球。加拿大进攻定位球:科内报销失去最佳头球点,戴维/科尼利厄斯仍有威胁。',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(4))

add_para('教练博弈:', bold=True, size=Pt(9), color=DARK)
add_para('南非(布鲁斯,67岁,本届后退休): 西索尔+莫科纳双防守后腰=大巴信号明确。低位防守+反击。落后→60\'上姆巴塔/福斯特/雷纳斯搏命。加拿大(马施): 拉林替补=本届最大胆用人——奥卢瓦塞伊首发冲防线,拉林+戴维斯65\'后替补收割。高位压迫+米勒/布坎南纯速度两翼。落后→戴维斯+拉林双超巨进场火力全开。\n关键换人博弈: 65\'仍是0-0→戴维斯+拉林替补登场=加拿大后25分钟火力碾压。本场最大胜负手!',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(6))

# ─ Upset risk ─
add_para('冷门风险: 中高', bold=True, size=Pt(10), color=ACCENT)
add_para('南非大巴已验证(1-0韩国)+西索尔首发=大巴硬度升级。加拿大首发少了拉林终结力+戴维斯破局能力→破大巴不如纸面。关键变量:南非能否撑65\'不丢球(在戴维斯+拉林进场前守住);奥卢瓦塞伊首发处子级表现;邦比托中卫新搭档是否被马塞科速度惩罚。',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(4))
add_para('冷门路径: 南非大巴守75\'→马施赌博式上戴维斯+拉林→南非反击偷一个(马塞科vs邦比托)→1-0。或0-0入加时→点球→威廉姆斯(AFCON扑4个)决胜。但戴维斯+拉林替补=加拿大后程火力太强,南非想守120分钟几乎不可能。',
         size=Pt(8.5), color=ACCENT, space_after=Pt(8))

# ─ Predictions ─
add_para('比分预测', bold=True, size=Pt(10), color=DARK)
pred_data = [['类型', '比分(常规)', '半场', '说明']]
pred_data.extend([
    ['首选', '1-1', '0-0', '南非大巴+西索尔死守→戴维下半场破僵→马塞科反击扳平→加时戴维斯+拉林决胜'],
    ['备选', '0-0', '0-0', '双方均保守→戴维斯/拉林登场仍无建树→点球(南非优势)'],
    ['备选', '0-1', '0-0', '戴维个人能力破大巴→南非反击无果'],
    ['备选', '南非 1-0', '0-0', '马塞科反击偷邦比托身后→大巴死守→冷门!'],
    ['备选', '0-2', '0-0', '戴维斯65\'替补助攻→加拿大75\'后连入2球'],
])
pt = doc.add_table(rows=len(pred_data), cols=4)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(pred_data[0]):
    set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(pred_data[1:], 1):
    bg = RGBColor(0xE8, 0xF8, 0xF5) if i == 1 else (LIGHT_GRAY if i % 2 == 0 else None)
    set_cell(pt.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg)
    set_cell(pt.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
    set_cell(pt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
    set_cell_left(pt.rows[i].cells[3], row[3], size=Pt(8.5), bg=bg)
for row in pt.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(1.2)
    row.cells[2].width = Inches(0.7)
    row.cells[3].width = Inches(7.8)
doc.add_paragraph()

# ─ Advancement ─
add_para('晋级预测', bold=True, size=Pt(10), color=DARK)
adata = [['类型', '晋级方', '方式']]
adata.extend([
    ['首选', '加拿大', '加时赛后'],
    ['备选', '加拿大', '常规时间0-1'],
    ['冷门', '南非', '常规时间1-0 或 点球决胜'],
])
at = doc.add_table(rows=len(adata), cols=3)
at.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(adata[0]):
    set_cell(at.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(adata[1:], 1):
    bg = RGBColor(0xE8, 0xF8, 0xF5) if i == 1 else (LIGHT_GRAY if i % 2 == 0 else None)
    set_cell(at.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg)
    set_cell(at.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
    set_cell_left(at.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
for row in at.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(3.5)
doc.add_paragraph()

# ─ Bracket ─
add_heading_styled('淘汰赛路径', level=2)
add_para('本场胜者 → Round of 16 → 对阵 荷兰 vs 摩洛哥 的胜者', bold=True, size=Pt(10), color=ACCENT)
add_para('如果加拿大晋级: 大概率面对荷兰(€837M)——几乎无晋级希望。如果南非晋级: 若摩洛哥爆冷胜荷兰→非洲内战机会增。',
         size=Pt(8.5), color=RGBColor(0x44,0x44,0x44), space_after=Pt(8))

# ═══════════════════════════════════
# FOOTER
# ═══════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: FIFA API + ESPN API + Standard / Khelnow / Al Jazeera / Goal (多方首发确认)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('身价数据: Transfermarkt 估算(加拿大~€280M / 南非~€45M)  |  框架: CLAUDE.md v17 (身价量化+强队三类+淘汰赛路径+10项必填清单)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('[注意] 淘汰赛无平局可能进入加时+点球  |  预测仅供分析参考  |  生成工具: python-docx',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

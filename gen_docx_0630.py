# -*- coding: utf-8 -*-
"""Generate 6月30日 3场预测 as .docx — v6美化规范."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年6月30日_3场预测.docx")
doc = Document()

# Page setup — landscape
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

DARK   = RGBColor(0x1A, 0x1A, 0x2E)
RED    = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
HDR_BG = RGBColor(0x1A, 0x1A, 0x2E)
GREEN  = RGBColor(0x27, 0xAE, 0x60)

def sc(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color: run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # vertical center
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def add_header_row(table, row_idx, texts, size=Pt(9)):
    for i, t in enumerate(texts):
        sc(table.cell(row_idx, i), t, bold=True, size=size, color=WHITE, bg=HDR_BG)

def add_data_row(table, row_idx, texts, bold_cols=None, bg=None, size=Pt(9)):
    for i, t in enumerate(texts):
        b = bold_cols and i in bold_cols
        sc(table.cell(row_idx, i), t, bold=b, size=size, bg=bg)

def add_heading(text, size=Pt(14), color=DARK):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_text(text, size=Pt(9), bold=False, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color: run.font.color.rgb = color
    return p

# ═══════════════════════════════════════
# COVER
# ═══════════════════════════════════════
add_heading('2026世界杯 — 6月30日淘汰赛 R32 三场预测', Pt(16), DARK)
add_text('生成时间: 2026年6月29日 北京时间 18:00', Pt(8), color=RGBColor(0x66,0x66,0x66))
add_text('数据来源: ESPN API + FIFA API + Sports Mole + SI + Transfermarkt (Planet Football)', Pt(8), color=RGBColor(0x66,0x66,0x66))
add_text('分析框架: CLAUDE.md v17 (身价量化+强队三类+淘汰赛路径+10项必填清单)', Pt(8), color=RGBColor(0x66,0x66,0x66))
doc.add_paragraph()

# ═══════════════════════════════════════
# SUMMARY TABLE
# ═══════════════════════════════════════
add_heading('预测汇总', Pt(12), DARK)
t = doc.add_table(rows=5, cols=8)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(t, 0, ['#', '时间 (BJT)', '阶段', '比赛', '身价比', '首选比分', '半场', '冷门风险'])
summary = [
    ['1', '01:00', 'R32', '巴西 vs 日本', '3.4:1', '巴西 2-0', '1-0', '中'],
    ['2', '04:30', 'R32', '德国 vs 巴拉圭', '6.2:1', '德国 3-0', '1-0', '低中'],
    ['3', '09:00', 'R32', '荷兰 vs 摩洛哥', '1.7:1', '1-1 (荷兰加时晋级)', '0-0', '中高'],
]
for i, row in enumerate(summary):
    add_data_row(t, i+1, row, bg=GRAY if i%2==0 else None)

# ═══════════════════════════════════════
# MATCH DETAILS FUNCTION
# ═══════════════════════════════════════
def write_match(match_data):
    """match_data: dict with all match info fields"""

    # Match title
    add_heading(f'{match_data["title"]}', Pt(12), RED)

    # Basic info table
    info = match_data['info']
    t = doc.add_table(rows=len(info)+1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, 0, ['项目', '内容'])
    for i, (k, v) in enumerate(info.items()):
        add_data_row(t, i+1, [k, v], bold_cols={0}, bg=GRAY if i%2==0 else None)
    doc.add_paragraph()

    # Group stage table
    add_text('小组赛回顾', Pt(10), bold=True)
    gs = match_data['group_stage']
    t2 = doc.add_table(rows=len(gs['home'])+len(gs['away'])+5, cols=4)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t2, 0, ['球队', '对手', '比分', '关键'])

    row = 1
    # Home team header — merged single cell spanning all 4 columns
    home_hdr = t2.cell(row, 0)
    home_hdr.merge(t2.cell(row, 3))
    sc(home_hdr, gs['home_name'], bold=True, size=Pt(9), color=DARK, align='left')
    row += 1
    for g in gs['home']:
        add_data_row(t2, row, [''] + g, bg=GRAY if (row%2==0) else None)
        row += 1

    # Away team header — merged single cell spanning all 4 columns
    away_hdr = t2.cell(row, 0)
    away_hdr.merge(t2.cell(row, 3))
    sc(away_hdr, gs['away_name'], bold=True, size=Pt(9), color=DARK, align='left')
    row += 1
    for g in gs['away']:
        add_data_row(t2, row, [''] + g, bg=GRAY if (row%2==0) else None)
        row += 1

    add_text(gs.get('home_note', ''), Pt(8), color=RGBColor(0x66,0x66,0x66))
    add_text(gs.get('away_note', ''), Pt(8), color=RGBColor(0x66,0x66,0x66))
    doc.add_paragraph()

    # 10-item analysis (condensed for DOCX)
    items = match_data.get('items', {})
    for item_title, item_data in items.items():
        add_text(item_title, Pt(10), bold=True, color=DARK)
        if isinstance(item_data, str):
            add_text(item_data, Pt(8.5))
        elif isinstance(item_data, list):
            for line in item_data:
                add_text(line, Pt(8.5))

        # If there's a table
        if isinstance(item_data, dict) and 'table' in item_data:
            tbl_data = item_data['table']
            t3 = doc.add_table(rows=len(tbl_data)+1, cols=len(tbl_data[0]))
            t3.style = 'Table Grid'
            t3.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_header_row(t3, 0, tbl_data[0])
            for j, r in enumerate(tbl_data[1:]):
                add_data_row(t3, j+1, r, bg=GRAY if j%2==0 else None)
        doc.add_paragraph()

    # Score prediction table
    add_text('比分预测', Pt(10), bold=True, color=DARK)
    preds = match_data['predictions']
    tp = doc.add_table(rows=len(preds)+1, cols=4)
    tp.style = 'Table Grid'
    tp.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tp, 0, ['类型', '比分', '半场', '说明'])
    for j, pr in enumerate(preds):
        bg = GREEN if pr[0] == '首选' else (GRAY if j%2==0 else None)
        add_data_row(tp, j+1, pr, bold_cols={0} if pr[0]=='首选' else None, bg=bg)
    doc.add_paragraph()

# ═══════════════════════════════════════
# MATCH 1: Brazil vs Japan
# ═══════════════════════════════════════
m1 = {
    'title': '比赛1: 巴西 vs 日本 — 01:00 BJT, NRG体育场, 休斯顿',
    'info': {
        'FIFA排名': '巴西 #6 vs 日本 #18',
        '身价比': '巴西 €928M vs 日本 €271M ≈ 3.4:1',
        '历史交锋': '巴西9胜4平1负, 但日本最近一次3-2获胜(2025年10月友谊赛)',
        '晋级奖励': '16强对阵 科特迪瓦 vs 挪威 的胜者',
        '强队分类': '巴西 超级巨星型 | 日本 大巴+反击型',
        '冷门风险': '中 — 日本2026年已击败巴西+英格兰+战平荷兰',
    },
    'group_stage': {
        'home_name': '巴西 (C组第1, 7分, +6GD)',
        'home': [
            ['摩洛哥', '1-1', '维尼修斯进球'],
            ['海地', '3-0', '库尼亚+维尼修斯主导'],
            ['苏格兰', '3-0', '维尼修斯3.06xG创世界杯单场纪录'],
        ],
        'home_note': '> 维尼修斯三场全部进球。最后两场零封。',
        'away_name': '日本 (F组第2, 5分, +4GD)',
        'away': [
            ['荷兰', '2-2', '两次落后两次扳平'],
            ['突尼斯', '胜', '小组赛第二场'],
            ['瑞典', '1-1', '铃木彩艳关键扑救保住第二'],
        ],
        'away_note': '> 最近540分钟仅丢3球。3-4-2-1大巴运转流畅。',
    },
    'items': {
        '因素导向表': {
            'table': [
                ['因素', '有利方', '理由'],
                ['维尼修斯状态(3场全进球+历史级xG)', '巴西 ★★★', '日本尚未面对过此级别单兵破局者'],
                ['身价比3.4:1', '巴西 ★★', '明显差距, 巴西正常发挥即胜'],
                ['日本3-4-2-1大巴已验证(540分钟仅丢3球)', '日本 ★★★', '已对荷兰+英格兰+巴西验证'],
                ['日本2025年10月3-2击败巴西', '日本 ★★★', '但不是同一防线(当时无马尔基尼奥斯+加布里埃尔)'],
                ['巴西对摩洛哥大巴1-1', '日本 ★★', '巴西对5后卫大巴仍有破局困难'],
                ['拉菲尼亚伤疑', '日本 ★', '巴西右路攻击力下降'],
                ['巴西防线vs10月完全不同', '巴西 ★★', '马尔基尼奥斯+加布里埃尔=更稳固'],
            ],
        },
        '伤病/停赛': {
            'table': [
                ['球队', '球员', '状态', '影响'],
                ['巴西', '拉菲尼亚(Raphinha)', '伤疑', '腿筋问题, 大概率替补'],
                ['巴西', '内马尔(Neymar)', '未入选', '已非巴西核心'],
                ['日本', '—', '全员可用', '无伤停'],
            ],
        },
        '亚非韧性评估(日本)': '5/5 — 低位防守[对][对] | 速度反击[对][对] | 定位球高点[对] | 前30分钟[对][对] | 被压制不崩盘[对][对]。本届世界杯防守纪律最好的非欧洲球队之一。',
        '教练博弈': '巴西(安切洛蒂): 4次欧冠冠军, 领先=控制节奏+防守中场; 落后=拉菲尼亚替补+边路传中。日本(森保一): 3-4-2-1大巴已成熟, 领先=532死守, 落后=堂安律+前田大然冲吊。关键: 若65分仍是0-0, 巴西是否急躁?',
        '定位球攻防': '巴西: 加布里埃尔+马尔基尼奥斯双中卫头球威胁。日本: 板仓滉+富安健洋高点, 但三中卫制空能力不如巴西双中卫。定位球是巴西对日本最可能得分路径之一。',
        '冷门路径': '日本3-4-2-1大巴守住前70分钟→巴西急躁压上→伊东纯也反击1v1过掉丹尼洛→铃木彩艳扑出所有射门→1-0或1-1(加时/点球)。最可能冷门比分: 1-1(加时)',
    },
    'predictions': [
        ['首选', '巴西 2-0', '1-0', '维尼修斯破局→库尼亚锁定'],
        ['备选', '巴西 1-0', '0-0', '定位球或维尼修斯晚破门'],
        ['备选', '巴西 2-1', '1-1', '日本反击偷一球→巴西逆转'],
        ['备选', '1-1', '0-0', '日本大巴满分→加时决胜'],
    ],
}

write_match(m1)

# ═══════════════════════════════════════
# MATCH 2: Germany vs Paraguay
# ═══════════════════════════════════════
m2 = {
    'title': '比赛2: 德国 vs 巴拉圭 — 04:30 BJT, 吉列体育场, 福克斯堡(波士顿)',
    'info': {
        'FIFA排名': '德国 #10 vs 巴拉圭 #48(估)',
        '身价比': '德国 €947M vs 巴拉圭 €154M ≈ 6.2:1',
        '历史交锋': '德国2胜1平。世界杯: 2002 R16 德国1-0。最近: 2013年3-3',
        '晋级奖励': '16强对阵 法国 vs 瑞典 的胜者 (大概率法国!)',
        '强队分类': '德国 超级巨星型 | 巴拉圭 大巴+阿尔米隆反击型',
        '冷门风险': '低中 — 德国对厄瓜多尔失利暴露防线速度短板, 但巴拉圭3场仅1球',
    },
    'group_stage': {
        'home_name': '德国 (E组第1, 6分, +6GD)',
        'home': [
            ['库拉索', '7-1', '穆西亚拉+维尔茨双双建功'],
            ['科特迪瓦', '2-1', '险胜, 防线暴露问题'],
            ['厄瓜多尔', '1-2', '失利! 防线混乱, 进攻哑火'],
        ],
        'home_note': '> 穆西亚拉(€100M)+维尔茨(€100M)双核。但对厄瓜多尔失利是警报。',
        'away_name': '巴拉圭 (D组第3, 4分, -2GD)',
        'away': [
            ['美国', '负', '小组赛揭幕战'],
            ['土耳其', '1-0', '关键胜利'],
            ['澳大利亚', '0-0', '闷平锁定晋级'],
        ],
        'away_note': '> 3场仅进1球但仅丢2球。大巴防守已验证。',
    },
    'items': {
        '因素导向表': {
            'table': [
                ['因素', '有利方', '理由'],
                ['身价比6.2:1', '德国 ★★★', '大差距, 德国正常发挥=稳胜'],
                ['德国对厄瓜多尔1-2失利→状态存疑', '巴拉圭 ★★', '防线被速度反击惩罚→巴拉圭有阿尔米隆'],
                ['施洛特贝克伤缺+布朗肌肉问题', '巴拉圭 ★', '德国防线深度减弱'],
                ['穆西亚拉+维尔茨双核(€200M组合)', '德国 ★★★', '破大巴能力在纸面上极强'],
                ['巴拉圭3场仅进1球', '德国 ★★★', '攻击力极弱, 即使德国防线犯错也难惩罚'],
                ['巴拉圭大巴已验证(3场仅丢2球)', '巴拉圭 ★★', '零封土耳其+澳大利亚'],
            ],
        },
        '伤病/停赛': {
            'table': [
                ['球队', '球员', '状态', '影响'],
                ['德国', '施洛特贝克(Schlotterbeck)', '伤缺', '左中卫重要轮换'],
                ['德国', '布朗(Nathaniel Brown)', '伤疑', '肌肉问题, 左后卫存疑'],
                ['巴拉圭', '—', '全员可用', '无重大伤停'],
            ],
        },
        '防守韧性评估(巴拉圭)': '4/5 — 低位防守[对][对] | 速度反击[对] (仅阿尔米隆一人) | 定位球高点[对] | 前30分钟[对] | 被压制不崩盘[对]。防守纪律好但攻击力太弱是致命短板。',
        '教练博弈': '德国(纳格尔斯曼): 战术灵活但有时过度复杂, 对厄瓜多尔失利暴露高位防线脆弱。巴拉圭(阿尔法罗): 防守优先, 领先=532死守, 落后=恩西索+定位球。关键: 德国能否尽早(前45分)破门?',
        '定位球攻防': '德国: 吕迪格+哈弗茨高点, 穆西亚拉/维尔茨罚球。巴拉圭: 南美传统定位球威胁, 阿尔米隆任意球。德国对厄瓜多尔定位球防守存疑。',
        '冷门路径': '德国久攻不下→纳格尔斯曼过度压上→阿尔米隆反击1v1过掉吕迪格→1-0。或0-0加时→点球。但巴拉圭3场1球的攻击力=德国大概率仍能进2+球。最可能冷门比分: 1-1(德国加时胜)',
    },
    'predictions': [
        ['首选', '德国 3-0', '1-0', '上半场破大巴→下半场双核收割'],
        ['备选', '德国 2-0', '1-0', '标准比分, 德国控制全场'],
        ['备选', '德国 3-1', '2-0', '阿尔米隆反击偷一个→德国仍稳赢'],
        ['备选', '德国 2-1', '1-1', '若巴拉圭先偷一个→德国逆转'],
    ],
}

write_match(m2)

# ═══════════════════════════════════════
# MATCH 3: Netherlands vs Morocco
# ═══════════════════════════════════════
m3 = {
    'title': '比赛3: 荷兰 vs 摩洛哥 — 09:00 BJT, 西班牙对外银行体育场, 蒙特雷, 墨西哥',
    'info': {
        'FIFA排名': '荷兰 #8 vs 摩洛哥 #7',
        '身价比': '荷兰 €754M vs 摩洛哥 €448M ≈ 1.7:1',
        '历史交锋': '荷兰2胜1负。世界杯: 1994小组赛 荷兰2-1',
        '晋级奖励': '16强对阵 加拿大',
        '强队分类': '荷兰 体系型倾向 | 摩洛哥 大巴+高质量反击型',
        '冷门风险': '中高 — 本日最难预测比赛。博彩市场定价摩洛哥略优(+130)。',
    },
    'group_stage': {
        'home_name': '荷兰 (F组第1, 7分, +6GD)',
        'home': [
            ['日本', '2-2', '两次领先两次被扳平'],
            ['瑞典', '胜', '小组赛第二场'],
            ['突尼斯', '胜', '小组赛第三场'],
        ],
        'home_note': '> 3场10球=本届攻击力最强之一。但防守对日本丢2球。',
        'away_name': '摩洛哥 (C组第2, 7分, +3GD)',
        'away': [
            ['巴西', '1-1', '大巴逼平巴西! 防守经典案例'],
            ['苏格兰', '1-0', '小胜'],
            ['海地', '4-2', '放开打攻击力展现, 赛巴里3场3球'],
        ],
        'away_note': '> 与巴西同组仅以净胜球屈居第2。2022四强+2026再进淘汰赛。',
    },
    'items': {
        '因素导向表': {
            'table': [
                ['因素', '有利方', '理由'],
                ['身价比1.7:1 (荷兰€754M vs 摩洛哥€448M)', '荷兰 ★', '接近, 但荷兰略有优势'],
                ['摩洛哥1-1逼平巴西(已验证大巴对超巨级)', '摩洛哥 ★★★', '巴西级别攻击力也被压制'],
                ['赛巴里状态(3场3球, 即将加盟拜仁)', '摩洛哥 ★★', '攻击线有破局者'],
                ['荷兰10球攻击力', '荷兰 ★★', '但仅面对一个强队(日本)'],
                ['荷兰对日本2-2防守不稳', '摩洛哥 ★★', '摩洛哥反击质量高于日本'],
                ['摩洛哥2022四强+跨届经验', '摩洛哥 ★★', '淘汰赛经验远超普通球队'],
                ['🚨 身价比1.7:1 < 3:1 → 任何结果都可能', '双方', '不做强队方向强硬断言'],
            ],
        },
        '伤病/停赛': {
            'table': [
                ['球队', '球员', '状态', '影响'],
                ['荷兰', '邓弗里斯(Dumfries)', '确认可用', '此前存疑已通过体检'],
                ['荷兰', '布罗比(Brobbey)', '确认可用', '轻伤无碍'],
                ['荷兰', '范德芬(Van de Ven)', '回归首发', '左后卫替换阿克, 速度提升'],
                ['摩洛哥', '马兹拉维+迪奥普+布阿迪+乌纳希', '回归首发', '上轮轮休主力全回'],
            ],
        },
        '非洲韧性评估(摩洛哥)': '5/5 — 低位防守[对][对] (对巴西1-1=大巴经典) | 速度反击[对][对] (阿什拉夫+迪亚斯+赛巴里) | 定位球高点[对][对] | 前30分钟[对][对] | 被压制不崩盘[对][对] (37%控球逼平巴西)。2022四强非偶然。',
        '教练博弈': '荷兰(科曼): 战术偏保守, 领先=防守中场控节奏, 落后=德佩+克鲁伊维特加强攻击。摩洛哥(瓦赫比): 大巴+阿什拉夫右路高速反击, 领先=大巴升级, 落后=75分后上第二个前锋。关键: 谁会先破门?',
        '定位球攻防': '荷兰: 范戴克+布罗比高点(空中优势)。摩洛哥: 迪奥普+赛巴里高点, 对海地定位球得分过。关键交锋: 范戴克防空 vs 赛巴里(186cm)头球。',
        '冷门路径': '荷兰控球但不进球→摩洛哥60-70分反击(阿什拉夫传中/赛巴里头球)→1-0→摩洛哥大巴升级→荷兰全线压上无建树→常规时间摩洛哥胜。或0-0加时→点球→布努扑点。最可能冷门比分: 摩洛哥1-0(常规时间)',
    },
    'predictions': [
        ['首选', '1-1 (荷兰加时晋级)', '0-0', '互有进球→加时荷兰深度优势胜'],
        ['备选', '荷兰 1-0', '0-0', '哈克波个人能力破大巴'],
        ['备选', '0-0 (点球)', '0-0', '双方保守→点球决胜'],
        ['备选', '摩洛哥 1-0', '0-0', '赛巴里反击/定位球→大巴死守'],
    ],
}

write_match(m3)

# ═══════════════════════════════════════
# KNOCKOUT PATH
# ═══════════════════════════════════════
add_heading('淘汰赛路径', Pt(12), DARK)
add_text('巴西/日本胜者 → R16 → 科特迪瓦 vs 挪威 的胜者 → QF → 可能碰德国区')
add_text('德国/巴拉圭胜者 → R16 → 法国 vs 瑞典 的胜者 (大概率法国! €1.52B)')
add_text('荷兰/摩洛哥胜者 → R16 → 加拿大')

pt = doc.add_table(rows=4, cols=2)
pt.style = 'Table Grid'
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(pt, 0, ['路径', '难度'])
add_data_row(pt, 1, ['巴西方向 → 科特迪瓦/挪威 → QF碰德国区', '中'])
add_data_row(pt, 2, ['德国方向 → 大概率法国(€1.52B)', '极高!'])
add_data_row(pt, 3, ['荷兰/摩洛哥方向 → 加拿大', '中高'])

doc.add_paragraph()

# ═══════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════
add_text('—', Pt(8), color=RGBColor(0xAA,0xAA,0xAA))
add_text('数据: ESPN API + FIFA API + Sports Mole + SI + Planet Football + Transfermarkt + Wikipedia', Pt(7), color=RGBColor(0x99,0x99,0x99))
add_text('身价: Transfermarkt via Planet Football (2026年6月) | 分析框架: CLAUDE.md v17', Pt(7), color=RGBColor(0x99,0x99,0x99))
add_text('生成时间: 2026年6月29日 北京时间 18:00 | Co-Authored-By: Claude Opus 4.8', Pt(7), color=RGBColor(0x99,0x99,0x99))

doc.save(OUT)
print(f'[对] DOCX saved to: {OUT}')

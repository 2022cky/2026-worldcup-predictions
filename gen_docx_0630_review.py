# -*- coding: utf-8 -*-
"""Generate 6月30日 复盘 as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年6月30日_复盘.docx")
doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

DARK  = RGBColor(0x1A,0x1A,0x2E)
RED   = RGBColor(0xC0,0x39,0x2B)
WHITE = RGBColor(0xFF,0xFF,0xFF)
GRAY  = RGBColor(0xF2,0xF2,0xF2)
HDR_BG= RGBColor(0x1A,0x1A,0x2E)
GREEN = RGBColor(0x27,0xAE,0x60)
YELLOW= RGBColor(0xF3,0x9C,0x12)

def sc(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT,'right':WD_ALIGN_PARAGRAPH.RIGHT}.get(align,WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold; run.font.size = size; run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if color: run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def hdr_row(t, row, texts, size=Pt(9)):
    for i,txt in enumerate(texts): sc(t.cell(row,i),txt,bold=True,size=size,color=WHITE,bg=HDR_BG)

def dat_row(t, row, texts, bold_cols=None, bg=None, size=Pt(9)):
    for i,txt in enumerate(texts):
        b = bold_cols and i in bold_cols
        sc(t.cell(row,i),txt,bold=b,size=size,bg=bg)

def heading(text, size=Pt(14), color=DARK):
    p=doc.add_paragraph();p.space_before=Pt(8);p.space_after=Pt(4)
    r=p.add_run(text);r.bold=True;r.font.size=size;r.font.color.rgb=color
    r.font.name='微软雅黑';r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def text(text, size=Pt(9), bold=False, color=None):
    p=doc.add_paragraph();p.space_before=Pt(1);p.space_after=Pt(1)
    r=p.add_run(text);r.bold=bold;r.font.size=size;r.font.name='微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if color: r.font.color.rgb = color

# ═══════════════ COVER ═══════════════
heading('2026世界杯 — 6月30日 R32 三场复盘',Pt(16),DARK)
text('复盘时间: 2026年6月30日 北京时间 14:00 | 数据: ESPN API + AnySearch 交叉验证',Pt(8),color=RGBColor(0x66,0x66,0x66))
text('框架: CLAUDE.md v17 + 6月30日复盘教训',Pt(8),color=RGBColor(0x66,0x66,0x66))
doc.add_paragraph()

# ═══════════════ SUMMARY ═══════════════
heading('复盘汇总',Pt(12),DARK)
t=doc.add_table(rows=4,cols=8);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t,0,['#','比赛','预测','半场预测','实际','半场实际','方向','比分'])
for i,row in enumerate([
    ['1','巴西 vs 日本','巴西 2-0','1-0','巴西 2-1','0-1','[对]','[错]-备选[对]'],
    ['2','德国 vs 巴拉圭','德国 3-0','1-0','1-1(巴拉圭点球4-3晋级)','0-1','[错]','[错]'],
    ['3','荷兰 vs 摩洛哥','1-1(荷兰加时晋级)','0-0','1-1(摩洛哥点球3-2晋级)','0-0','[错]','[对]'],
]): dat_row(t,i+1,row,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

t2=doc.add_table(rows=4,cols=2);t2.style='Table Grid';t2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t2,0,['指标','数值'])
for i,row in enumerate([['方向正确','1/3 (33%)'],['比分命中(含备选)','2/3 (67%)'],['冷门漏报','2场 (德国+荷兰出局)']]):
    bg = RED if '冷门' in row[0] else (GRAY if i%2==0 else None)
    dat_row(t2,i+1,row,bg=bg,size=Pt(10))
doc.add_paragraph()

# ═══════════════ MATCH 1 ═══════════════
heading('比赛1: 巴西 2-1 日本 — 补时绝杀掩盖苦战',Pt(12),RED)

text('预测: 巴西 2-0 (半场 1-0) | 实际: 巴西 2-1 (半场 0-1) | 判定: 方向[对] 备选2-1[对]',Pt(9),bold=True)

t3=doc.add_table(rows=4,cols=5);t3.style='Table Grid';t3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t3,0,['项目','预测','实际','判定','说明'])
dat_row(t3,1,['比分','巴西 2-0','巴西 2-1','[错]','备选2-1[对] — "日本反击偷袭→巴西逆转"完美命中'])
dat_row(t3,2,['半场','1-0','0-1','[错]','日本主动前压拦截+远射破门, 非纯大巴'])
dat_row(t3,3,['进球者','维尼修斯+库尼亚','卡塞米罗+马丁内利','[错]','后腰+替补 vs 超巨前锋'])
doc.add_paragraph()

heading('进球',Pt(10),DARK)
t4=doc.add_table(rows=4,cols=3);t4.style='Table Grid';t4.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t4,0,['时间','球员','描述'])
dat_row(t4,1,['29\'','佐野海舟(Kaishu Sano)','拦截达尼洛传球→推进→远射(首粒世界杯进球)'])
dat_row(t4,2,['56\'','卡塞米罗(Casemiro)','加布里埃尔传中→后点头球'])
dat_row(t4,3,['90+5\'','马丁内利(Gabriel Martinelli)','布鲁诺·吉马良斯直塞→单刀绝杀'])
doc.add_paragraph()

heading('关键统计',Pt(10),DARK)
text('控球率: 巴西 69% vs 31% 日本 | 射门: 19 vs 5 | 射正: 7 vs 2 | xG: 1.69 vs 0.23',Pt(9))
doc.add_paragraph()

heading('复盘分析',Pt(10),DARK)
text('[对]项: 1)方向正确 2)备选2-1精准命中走势 3)日本韧性5/5全部兑现 4)冷门风险"中"合理',Pt(9))
text('[错]项: 1)半场1-0→0-1, 低估日本主动攻击能力 2)进球者全错 3)达尼洛39岁右后卫失误未预判',Pt(9))
text('教训: ①大巴≠不进攻 ②淘汰赛替补>首发预测 ③超巨可被包夹锁死',Pt(9),bold=True,color=RED)
doc.add_paragraph()

# ═══════════════ MATCH 2 ═══════════════
heading('比赛2: 德国 1-1 巴拉圭 (巴拉圭点球4-3晋级) — 超级冷门!',Pt(12),RED)

text('预测: 德国 3-0 (半场 1-0) | 实际: 1-1 (半场 0-1) 巴拉圭点球4-3 | 判定: 全错',Pt(9),bold=True,color=RED)

heading('进球',Pt(10),DARK)
t5=doc.add_table(rows=4,cols=3);t5.style='Table Grid';t5.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t5,0,['时间','球员','描述'])
dat_row(t5,1,['42\'','恩西索(Julio Enciso)','加拉尔萨传中→头球破门'])
dat_row(t5,2,['54\'','哈弗茨(Kai Havertz)','维尔茨助攻→头球扳平'])
dat_row(t5,3,['102\'','塔(Jonathan Tah)','VAR取消 — 安东犯规门将吉尔'])
doc.add_paragraph()

heading('点球大战 (巴拉圭 4-3)',Pt(10),DARK)
t6=doc.add_table(rows=7,cols=5);t6.style='Table Grid';t6.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t6,0,['轮','德国','结果','巴拉圭','结果'])
for i,row in enumerate([
    ['1','哈弗茨','[错]被扑','毛里西奥','[对]'],
    ['2','基米希','[对]','戈麦斯','[对]'],
    ['3','穆西亚拉','[对]','加拉尔萨','[对]'],
    ['4','沃尔特马德','[错]被扑','萨纳布里亚','[错]'],
    ['5','阿米里','[对]','巴尔布埃纳','[错]被扑'],
    ['6','塔','[错]打飞','卡纳莱','[对]胜!'],
]): dat_row(t6,i+1,row,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

text('德国队史首次世界杯点球失利(此前4/4全胜)。连续三届大赛首轮出局: 2018小组→2022小组→2026 R32。',Pt(9),bold=True,color=RED)
doc.add_paragraph()

heading('核心教训 (严重)',Pt(10),RED)
for s in [
    '1. 淘汰赛身价比阈值不够 — 6.2:1被大巴完全抵消',
    '2. 超级巨星型需验证"不同进攻方式" — Musiala+Wirtz同质化=伪双核',
    '3. 前一场失利是结构警报 — 厄瓜多尔1-2不是"状态波动"',
    '4. 备选4个比分全预测德国胜 — 必须覆盖冷平',
    '5. 德国淘汰赛心理阴影 — 连续三届首轮出局是系统性崩溃',
]: text(s,Pt(9),bold=True)
doc.add_paragraph()

# ═══════════════ MATCH 3 ═══════════════
heading('比赛3: 荷兰 1-1 摩洛哥 (摩洛哥点球3-2晋级)',Pt(12),RED)

text('预测: 1-1 (荷兰加时晋级) 半场0-0 | 实际: 1-1 (摩洛哥点球3-2晋级) 半场0-0 | 比分[对] 晋级方[错]',Pt(9),bold=True)

heading('进球',Pt(10),DARK)
t7=doc.add_table(rows=3,cols=3);t7.style='Table Grid';t7.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t7,0,['时间','球员','描述'])
dat_row(t7,1,['72\'','哈克波(Cody Gakpo)','韦格霍斯特摆渡→萨默维尔传中→右脚破门(妻子流产仍出战)'])
dat_row(t7,2,['90+1\'','迪奥普(Issa Diop)','塔尔比传中→头球绝平!'])
doc.add_paragraph()

heading('点球大战 (摩洛哥 3-2)',Pt(10),DARK)
t8=doc.add_table(rows=6,cols=5);t8.style='Table Grid';t8.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t8,0,['轮','荷兰','结果','摩洛哥','结果'])
for i,row in enumerate([
    ['1','科普梅纳斯','[对]','埃尔艾纳乌伊','[错]横梁'],
    ['2','克鲁伊维特','[错]立柱','拉希米','[对]'],
    ['3','韦格霍斯特','[对]','塔尔比','[对]'],
    ['4','廷伯','[错]偏出','阿什拉夫','[错]立柱'],
    ['5','萨默维尔','[错]布努扑出','赛巴里','[对]胜!'],
]): dat_row(t8,i+1,row,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('关键统计',Pt(10),DARK)
text('控球率: 摩洛哥 65% vs 35% 荷兰 | 射门: 10 vs 6 | xG: 1.40 vs 0.23 | 扑救: 费尔布鲁根5次(全场最高)',Pt(9))
doc.add_paragraph()

heading('复盘分析',Pt(10),DARK)
text('[对]: 1)首选1-1命中 2)半场0-0命中 3)冷门风险中高 4)定位球进球预判 5)比赛走势预判(僵持→先进→追平→加时)',Pt(9))
text('[错]: 1)晋级方(荷兰→摩洛哥) 2)低估布努点球战力 3)低估荷兰点球诅咒(连续3届点球出局) 4)韦格霍斯特关键变量未专门分析',Pt(9))
text('核心教训: ①布努=世界杯最强点球门将之一 ②荷兰无冕之王心理不是★权重是★★★ ③1.4:1身价比=方向不应倾斜',Pt(9),bold=True,color=RED)
doc.add_paragraph()

# ═══════════════ RULE IMPACT ═══════════════
heading('6月30日对规则的影响',Pt(12),DARK)
t9=doc.add_table(rows=7,cols=3);t9.style='Table Grid';t9.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t9,0,['#','规则调整','来源'])
for i,row in enumerate([
    ['1','淘汰赛身价比阈值提高 — R32需>20:1才不做平局首选',  '德国出局'],
    ['2','超级巨星型需验证"不同进攻方式" — 同质化降为体系型','德国出局'],
    ['3','点球门将加权 — 布努/马丁内斯等已证明的点球门将额外加权','荷兰出局'],
    ['4','前一场暴露防守缺陷 = 淘汰赛结构警报','德国出局'],
    ['5','备选必须覆盖冷平 — 4个备选全预测强队胜=不合格','德国出局'],
    ['6','"大巴已验证"权重 > 身价优势','荷兰出局'],
]): dat_row(t9,i+1,row,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

# ═══════════════ BRACKET ═══════════════
heading('R32已完赛汇总',Pt(12),DARK)
t10=doc.add_table(rows=5,cols=4);t10.style='Table Grid';t10.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr_row(t10,0,['比赛','结果','晋级','R16对手'])
for i,row in enumerate([
    ['巴西 vs 日本','2-1','巴西','科特迪瓦/挪威 胜者'],
    ['德国 vs 巴拉圭','1-1 (点球3-4)','巴拉圭 [火]','法国/瑞典 胜者'],
    ['荷兰 vs 摩洛哥','1-1 (点球2-3)','摩洛哥 [火]','加拿大'],
    ['加拿大 vs 南非','1-0','加拿大','摩洛哥'],
]): dat_row(t10,i+1,row,bg=GRAY if i%2==0 else None)

doc.add_paragraph()
text('R32前4场: 2场点球冷门! 德国+荷兰同天出局, 欧洲两大豪门一天内被淘汰。',Pt(9),bold=True,color=RED)

text('—',Pt(8),color=RGBColor(0xAA,0xAA,0xAA))
text('数据: ESPN API + AnySearch + Al Jazeera + Sporting News + ekantipur | 分析: CLAUDE.md v17',Pt(7),color=RGBColor(0x99,0x99,0x99))
text('复盘时间: 2026年6月30日 北京时间 14:00 | Co-Authored-By: Claude Opus 4.8',Pt(7),color=RGBColor(0x99,0x99,0x99))

doc.save(OUT)
print(f'[OK] DOCX saved to: {OUT}')

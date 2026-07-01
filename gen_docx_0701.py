# -*- coding: utf-8 -*-
"""Generate 7月1日 3场预测 as .docx — updated with lineup confirmations & odds"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月1日_3场预测.docx")
doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7); section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2); section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']; font = style.font
font.name = '微软雅黑'; font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

DARK=RGBColor(0x1A,0x1A,0x2E);RED=RGBColor(0xC0,0x39,0x2B);WHITE=RGBColor(0xFF,0xFF,0xFF)
GRAY=RGBColor(0xF2,0xF2,0xF2);HDR_BG=RGBColor(0x1A,0x1A,0x2E);GREEN_BG=RGBColor(0xE8,0xF8,0xF5)
GRAY_TEXT=RGBColor(0x7F,0x8C,0x8D)

def sc(cell,text,bold=False,size=Pt(9),color=None,bg=None,align='center'):
    cell.text='';p=cell.paragraphs[0]
    p.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT,'right':WD_ALIGN_PARAGRAPH.RIGHT}.get(align,WD_ALIGN_PARAGRAPH.CENTER)
    r=p.add_run(str(text));r.bold=bold;r.font.size=size;r.font.name='微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if color:r.font.color.rgb=color
    if bg:
        shading=parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc=cell._tc;tcPr=tc.get_or_add_tcPr()
    vAlign=parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>');tcPr.append(vAlign)

def hdr(t,row,texts,size=Pt(9)):
    for i,txt in enumerate(texts):sc(t.cell(row,i),txt,bold=True,size=size,color=WHITE,bg=HDR_BG)

def dat(t,row,texts,bold_cols=None,bg=None,size=Pt(9)):
    for i,txt in enumerate(texts):
        b=bold_cols and i in bold_cols;sc(t.cell(row,i),txt,bold=b,size=size,bg=bg)

def heading(text,size=Pt(14),color=DARK):
    p=doc.add_paragraph();p.space_before=Pt(8);p.space_after=Pt(4)
    r=p.add_run(text);r.bold=True;r.font.size=size;r.font.color.rgb=color
    r.font.name='微软雅黑';r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def txt(text,size=Pt(9),bold=False,color=None):
    p=doc.add_paragraph();p.space_before=Pt(1);p.space_after=Pt(1)
    r=p.add_run(text);r.bold=bold;r.font.size=size;r.font.name='微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if color:r.font.color.rgb=color

# ═══════════════ COVER ═══════════════
heading('2026世界杯 — 7月1日淘汰赛 R32 三场预测',Pt(16),DARK)
txt('生成时间: 2026年6月30日 北京时间 16:00  |  更新: 2026年7月1日 08:05 BJT (官方首发确认+赔率)',Pt(8),color=GRAY_TEXT)
txt('数据来源: ESPN API + FIFA API + Sports Mole + SI + RotoWire + Transfermarkt + DraftKings',Pt(8),color=GRAY_TEXT)
txt('分析框架: CLAUDE.md v17 + 6月30日复盘教训(淘汰赛身价比阈值提高+备选覆盖冷平+点球门将加权)',Pt(8),color=GRAY_TEXT)
doc.add_paragraph()

# ═══════════════ SUMMARY ═══════════════
heading('预测汇总',Pt(12),DARK)
t=doc.add_table(rows=4,cols=8);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t,0,['#','时间 (BJT)','阶段','比赛','身价比','首选比分','半场','冷门风险'])
for i,row in enumerate([
    ['1','01:00','R32','科特迪瓦 vs 挪威','2.1:1 (挪威占优)','挪威 2-1','1-0','中高'],
    ['2','05:00','R32','法国 vs 瑞典','5.9:1','法国 3-0','2-0','低'],
    ['3','09:00','R32','墨西哥 vs 厄瓜多尔','1.4:1','墨西哥 1-0','0-0','中→中偏高'],
]): dat(t,i+1,row,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

# ═══════════════ BRACKET ═══════════════
heading('淘汰赛路径 (含6/30冷门影响)',Pt(12),DARK)
tb=doc.add_table(rows=4,cols=2);tb.style='Table Grid';tb.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tb,0,['路径','难度'])
dat(tb,1,['科特迪瓦/挪威 → R16 → 巴西 → 八强可能碰法国区','高'])
dat(tb,2,['法国/瑞典 → R16 → 巴拉圭(非德国!) → 八强','大幅降低!'],bold_cols={1})
dat(tb,3,['墨西哥/厄瓜多尔 → R16 → 英格兰/刚果民主共和国 胜者','中高'])
doc.add_paragraph()
txt('已确定: 巴西淘汰日本等待R16 / 巴拉圭点球淘汰德国等待R16 / 摩洛哥淘汰荷兰等待加拿大',Pt(9),bold=True,color=RED)
txt('德国出局→法国R16仅需过巴拉圭。八强避开德国。法国夺冠路径大幅简化。',Pt(9),color=RED)
doc.add_paragraph()

# ═══════════════ MATCH 1 ═══════════════
heading('比赛1: 科特迪瓦 vs 挪威 — 01:00 BJT, 达拉斯体育场',Pt(12),RED)

info1=[['项目','内容'],['FIFA排名','科特迪瓦 #31 vs 挪威 #17'],['身价比','科特迪瓦~€280M vs 挪威€601M ≈ 1:2.1 (挪威优)'],
       ['历史交锋','首次世界杯交锋'],['晋级奖励','16强对阵 巴西 (巴西2-1淘汰日本已在等待)'],
       ['强队分类','挪威: 超级巨星型(偏单一)—哈兰德世界最佳9号 | 科特迪瓦: 大巴+双爆点反击型'],
       ['冷门风险','中高 — 科特迪瓦队史首次淘汰赛, 但挪威防线10场仅1零封']]
ti=doc.add_table(rows=len(info1),cols=2);ti.style='Table Grid';ti.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(ti,0,info1[0])
for i,r in enumerate(info1[1:]):dat(ti,i+1,r,bold_cols={0},bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('小组赛回顾',Pt(10),DARK)
tg=doc.add_table(rows=3+3+3,cols=4);tg.style='Table Grid';tg.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tg,0,['球队','对手','比分','关键'])
tg.cell(1,0).merge(tg.cell(1,3));sc(tg.cell(1,0),'科特迪瓦 (E组第2, 6分) 队史首次进入淘汰赛',bold=True,color=DARK,align='left')
for i,r in enumerate([['厄瓜多尔','1-0','迪亚洛90\'绝杀, 开门红'],['德国','1-2','凯西先破门→昂达夫替补双响逆转'],['库拉索','2-0','佩佩梅开二度, 零封']]):
    dat(tg,i+2,['']+r,bg=GRAY if i%2==0 else None)
tg.cell(5,0).merge(tg.cell(5,3));sc(tg.cell(5,0),'挪威 (I组第2, 6分) 哈兰德依赖症严重',bold=True,color=DARK,align='left')
for i,r in enumerate([['伊拉克','4-1','哈兰德2球, 上半场曾被紧逼'],['塞内加尔','3-2','哈兰德再2球, 防线失2球拉响警报'],['法国(轮换)','4-1','轮换哈兰德→全队无还手之力, 登贝莱帽子戏法']]):
    dat(tg,i+6,['']+r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('首轮球员评分',Pt(10),DARK)
txt('科特迪瓦: 佩佩 8.0(对库拉索2球)/ 迪奥曼德 7.8(边路爆点)/ 凯西 7.5(中场核心)/ 桑加雷 7.3/ 科索努 7.4(3场2零封) | 最低分: 进攻创造力不够稳定',Pt(8.5))
txt('挪威: 哈兰德 9.2 MOTM(4球,52场59球)/ 厄德高 7.8(本届0球0助)/ 努萨 7.5/ 贝格 7.2/ 阿耶尔 6.8 | 最低分: 挪威防线=5.5(10场仅1零封)',Pt(8.5))
doc.add_paragraph()

heading('因素导向表',Pt(10),DARK)
tf=doc.add_table(rows=10,cols=3);tf.style='Table Grid';tf.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tf,0,['因素','有利方','理由'])
for i,r in enumerate([
    ['哈兰德个人能力(52场59球)','挪威 ★★★','科特迪瓦防线从未面对过此级别终结者'],
    ['挪威防线脆弱(10场仅1零封)','科特迪瓦 ★★','佩佩+迪奥曼德双爆点可击穿'],
    ['雷尔森伤缺→右后卫漏洞','科特迪瓦 ★★','挪威RB被迫用中场奥尔斯内斯客串'],
    ['辛戈伤疑→科特迪瓦右中卫缺口','挪威 ★★','科特迪瓦最强防线一环缺失'],
    ['科特迪瓦3场2零封(防守纪律)','科特迪瓦 ★★','对比挪威防守, 非洲大象更稳'],
    ['哈兰德缺席时挪威=另一支队(对法国1-4)','科特迪瓦 ★★','如果哈兰德被锁死, 挪威无B计划'],
    ['厄德高本届0球0助(创造核心哑火)','科特迪瓦 ★','挪威进攻仅靠哈兰德终结'],
    ['科特迪瓦首次淘汰赛(经验不足)','挪威 ★','心理层面, 但挪威也是28年来首次淘汰赛'],
    ['身价比1:2.1(3:1~10:1区间偏下)','挪威 ★','挪威占优但不足以碾压'],
]): dat(tf,i+1,r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('伤病/停赛',Pt(10),DARK)
tj=doc.add_table(rows=5,cols=4);tj.style='Table Grid';tj.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tj,0,['球队','球员','状态','影响'])
for i,r in enumerate([
    ['科特迪瓦','辛戈(Singo)','[注意]伤疑','右中卫, 腿筋问题。缺阵=科索努搭档乌斯曼·迪奥曼德'],
    ['挪威','雷尔森(Ryerson)','[缺]确认缺阵','多特右后卫, 大腿伤→奥尔斯内斯(中场)客串→右路防守削弱'],
    ['挪威','哈兰德','[可]','对法国轮休→满血出战'],
    ['挪威','厄德高','[可]','需找回状态(本届0球0助)'],
]): dat(tj,i+1,r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('非洲韧性评估(科特迪瓦)',Pt(10),DARK)
txt('4/5 — 低位防守:良好 | 速度反击:特优(佩佩+迪奥曼德=本届非洲最快双翼) | 定位球:良好 | 前30分钟:良好 | 被压制不崩盘:良好',Pt(9))
txt('唯一瑕疵: 对德国的下半场被逆转暴露被持续压制时的脆弱。',Pt(9),color=GRAY_TEXT)
doc.add_paragraph()

heading('比分预测',Pt(10),DARK)
tp=doc.add_table(rows=5,cols=4);tp.style='Table Grid';tp.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tp,0,['类型','比分','半场','说明'])
for i,r in enumerate([
    ['首选','挪威 2-1','1-0','哈兰德上半场破门→科特迪瓦反击扳平→哈兰德/索尔洛特锁定'],
    ['备选','挪威 1-0','0-0','僵持大半场→哈兰德个人能力解决'],
    ['备选','挪威 3-1','1-0','科特迪瓦防线被哈兰德连续击穿'],
    ['备选','1-1(挪威加时晋级)','0-0','科特迪瓦大巴成功→哈兰德被限制→加时挪威深度优势'],
]): dat(tp,i+1,r,bold_cols={0} if r[0]=='首选' else None,bg=GREEN_BG if r[0]=='首选' else (GRAY if i%2==0 else None))
doc.add_paragraph()

heading('冷门路径',Pt(10),RED)
txt('科特迪瓦大巴守住65分钟→挪威急躁→佩佩反击1v1过掉客串RB奥尔斯内斯→1-0→升级大巴→科索努+迪奥曼德双人锁死哈兰德→1-0或1-1(加时/点球)。最可能冷门比分: 1-1(科特迪瓦加时或点球胜)',Pt(9))
doc.add_paragraph()

# ═══════════════ MATCH 2 ═══════════════
heading('比赛2: 法国 vs 瑞典 — 05:00 BJT, 纽约新泽西体育场 (决赛场地!)',Pt(12),RED)

info2=[['项目','内容'],['FIFA排名','法国 #2 vs 瑞典 #36'],['身价比','法国€1.476B vs 瑞典~€250M ≈ 5.9:1'],
       ['历史交锋','法国近6战5胜1负, 近4胜中3场仅赢1球'],
       ['晋级奖励','16强对阵 巴拉圭! (德国出局后的巨大奖励)'],
       ['强队分类','法国: 超级巨星型(超级版)—姆巴佩+登贝莱=7球 | 瑞典: 非大巴的攻击型弱队'],
       ['冷门风险','低 — 瑞典对荷兰丢5球, 不具备大巴能力']]
ti2=doc.add_table(rows=len(info2),cols=2);ti2.style='Table Grid';ti2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(ti2,0,info2[0])
for i,r in enumerate(info2[1:]):dat(ti2,i+1,r,bold_cols={0},bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('小组赛回顾',Pt(10),DARK)
tg2=doc.add_table(rows=3+3+3,cols=4);tg2.style='Table Grid';tg2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tg2,0,['球队','对手','比分','关键'])
tg2.cell(1,0).merge(tg2.cell(1,3));sc(tg2.cell(1,0),'法国 (I组第1, 9分, +8GD) 3战全胜! 进10球仅丢2球',bold=True,color=DARK,align='left')
for i,r in enumerate([['伊拉克','3-0','姆巴佩2球, 轻松开局'],['塞内加尔','3-1','姆巴佩+登贝莱+奥利塞各1球'],['挪威(轮换)','4-1','登贝莱帽子戏法! 无哈兰德的挪威被打穿']]):
    dat(tg2,i+2,['']+r,bg=GRAY if i%2==0 else None)
tg2.cell(5,0).merge(tg2.cell(5,3));sc(tg2.cell(5,0),'瑞典 (F组第3, 4分, 0GD) 最佳第三名勉强晋级',bold=True,color=DARK,align='left')
for i,r in enumerate([['突尼斯','2-1','小胜'],['荷兰','1-5','防线被荷兰完全打穿!'],['日本','1-1','埃兰加远射得分, 勉强守住平局']]):
    dat(tg2,i+6,['']+r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('首轮球员评分',Pt(10),DARK)
txt('法国: 姆巴佩 9.0 MOTM(4球,金靴领跑)/ 登贝莱 8.8(帽子戏法)/ 奥利塞 7.8(3助)/ 楚阿梅尼 7.5/ 于帕梅卡诺 7.2/ 萨利巴[注意](背伤休战对挪威,大概率带伤出战)',Pt(8.5))
txt('瑞典: 伊萨克 7.5(本届1球)/ 埃兰加 7.5(对日本世界波)/ 哲凯赖什 7.3/ 贝里瓦尔 7.0 | 最低分: 瑞典防线=4.5(对荷兰丢5球,希恩伤缺使林德洛夫被迫回撤)',Pt(8.5))
doc.add_paragraph()

heading('因素导向表',Pt(10),DARK)
tf2=doc.add_table(rows=10,cols=3);tf2.style='Table Grid';tf2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tf2,0,['因素','有利方','理由'])
for i,r in enumerate([
    ['身价比5.9:1(法国1.476B vs 瑞典250M)','法国 ★★★','大差距。瑞典不是大巴型(对荷兰丢5球)'],
    ['姆巴佩+登贝莱状态(合计7球)','法国 ★★★','瑞典防线无法同时应对两个爆点'],
    ['伊萨克·希恩伤缺(瑞典防线核心)','法国 ★★★','林德洛夫被迫从中场回撤→防区混乱'],
    ['瑞典三叉戟(伊萨克+哲凯赖什+埃兰加)','瑞典 ★★','有能力攻破任何防线, 法国被动防守可能被惩罚'],
    ['法国无球时被动(小组赛未暴露但历史存在)','瑞典 ★★','如果瑞典先进球, 比赛走势可能不同'],
    ['瑞典对荷兰1-5(非大巴型防守)','法国 ★★','瑞典不是摩洛哥/巴拉圭——他们会给法国空间'],
    ['法国R16打巴拉圭(非德国)','法国 ★','路径大幅简化, 心态更放松'],
    ['决赛场地打首场淘汰赛(法国心理象征)','法国 ★','提前熟悉+积极心理暗示'],
    ['德尚赛末离任→"最后一舞"动力','法国 ★','全队为教练而战的额外动力'],
]): dat(tf2,i+1,r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('伤病/停赛',Pt(10),DARK)
tj2=doc.add_table(rows=3,cols=4);tj2.style='Table Grid';tj2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tj2,0,['球队','球员','状态','影响'])
dat(tj2,1,['法国','萨利巴(Saliba)','[可]带伤','背伤, 大概率首发。阿森纳级别中卫'])
dat(tj2,2,['瑞典','伊萨克·希恩(Hien)','[缺]腿筋','防线核心缺阵! 林德洛夫被迫从中场回撤→防守体系重组'])
doc.add_paragraph()

heading('瑞典韧性评估',Pt(10),DARK)
txt('2/5 — 低位防守:差(对荷兰丢5球) | 速度反击:良好 | 定位球:良好 | 前30分钟:差(17分钟被布罗比双响) | 被压制不崩盘:差',Pt(9))
txt('瑞典不擅长也不愿意摆大巴。对荷兰1-5证明了这一点。这正是法国大胜的底层逻辑。',Pt(9),color=RED)
doc.add_paragraph()

heading('比分预测',Pt(10),DARK)
tp2=doc.add_table(rows=5,cols=4);tp2.style='Table Grid';tp2.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tp2,0,['类型','比分','半场','说明'])
for i,r in enumerate([
    ['首选','法国 3-0','2-0','上半场解决战斗→下半场控制节奏'],
    ['备选','法国 2-0','1-0','瑞典开场保守→姆巴佩下半场破局'],
    ['备选','法国 3-1','2-1','瑞典反击偷一个→法国实力碾压'],
    ['备选','法国 2-1','1-1','最接近冷门: 伊萨克先破门→法国逆转'],
]): dat(tp2,i+1,r,bold_cols={0} if r[0]=='首选' else None,bg=GREEN_BG if r[0]=='首选' else (GRAY if i%2==0 else None))
doc.add_paragraph()

heading('冷门路径',Pt(10),RED)
txt('法国开场太放松→瑞典三叉戟高速反击→伊萨克/哲凯赖什先破门→法国陷入"攻大巴"困境→瑞典龟缩30分钟→1-0。需要: 法国极度轻敌+瑞典防线超常发挥(从未展示)+迈尼昂失误。三者同时发生概率极低。最可能"冷门": 法国2-1(瑞典先进→法国逆转),这不是真正的冷门。',Pt(9))
doc.add_paragraph()

# ═══════════════ MATCH 3 ═══════════════
heading('比赛3: 墨西哥 vs 厄瓜多尔 — 09:00 BJT, 墨西哥城阿兹台克体育场',Pt(12),RED)
txt('赛前首发确认: 2026年7月1日 08:05 BJT — FIFA API + ESPN API 双重确认 | DraftKings: 墨西哥晋级64% / 厄瓜多尔35%',Pt(8.5),bold=True,color=RED)
doc.add_paragraph()

info3=[['项目','内容'],['FIFA排名','墨西哥 #14 vs 厄瓜多尔 #30'],['身价比','墨西哥~€250M vs 厄瓜多尔~€180M ≈ 1.4:1'],
       ['历史交锋','墨西哥17胜7平4负。世界杯: 2002小组赛墨西哥2-1'],
       ['晋级奖励','16强对阵 英格兰 vs 刚果民主共和国 的胜者(大概率英格兰)'],
       ['强队分类','墨西哥: 体系型(主场加成版)—阿兹台克2240m+80000人 | 厄瓜多尔: 不稳定爆冷型'],
       ['冷门风险','中→中偏高 — 17岁莫拉淘汰赛首秀vs凯塞多=最大中场经验差']]
ti3=doc.add_table(rows=len(info3),cols=2);ti3.style='Table Grid';ti3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(ti3,0,info3[0])
for i,r in enumerate(info3[1:]):dat(ti3,i+1,r,bold_cols={0},bg=GRAY if i%2==0 else None)
doc.add_paragraph()

# ── Official Lineups ──
heading('官方首发阵容 (FIFA+ESPN双重确认)',Pt(10),RED)
txt('墨西哥 4-3-3: 兰赫尔(GK,3场零封) | 桑切斯(RB) | 蒙特斯(CB,C) | 巴斯克斯(CB) | 加利亚多(LB) | 利拉(CDM) | 罗莫(CM) | 莫拉(CM,17岁淘汰赛首秀!) | 阿尔瓦拉多(RW) | 劳尔·希门尼斯(ST,2球) | 基尼奥内斯(LW)',Pt(8.5))
txt('替补关键: 埃德松·阿尔瓦雷斯(费内巴切铁腰,下半场核武!) / 圣地亚哥·希门尼斯(费耶诺德) / 奥乔亚(40岁6届)',Pt(8),color=GRAY_TEXT)
txt('厄瓜多尔 4-4-2: 加林德斯(GK) | 佛朗哥(RB) | 奥尔多涅斯(CB) | 帕乔(CB,PSG) | 欣卡皮耶(LB,阿森纳) | 凯塞多(CM,C,切尔西!) | 维特(CM) | 耶博阿(RM) | 安古洛(LM,对德国扳平球) | 瓦伦西亚(ST,36岁距50球差1球) | 普拉塔(ST,对德国77\'制胜球!)',Pt(8.5))
txt('替补关键: 肯德里·派斯(17岁天才,下半场核武) / 埃斯图皮尼安(布莱顿)',Pt(8),color=GRAY_TEXT)
doc.add_paragraph()

heading('小组赛回顾',Pt(10),DARK)
tg3=doc.add_table(rows=3+3+3,cols=4);tg3.style='Table Grid';tg3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tg3,0,['球队','对手','比分','关键'])
tg3.cell(1,0).merge(tg3.cell(1,3));sc(tg3.cell(1,0),'墨西哥 (A组第1, 9分, +6GD) 3战全胜+0丢球! 唯一小组赛满零封球队',bold=True,color=DARK,align='left')
for i,r in enumerate([['南非','2-0','基尼奥内斯+希门尼斯, 3红牌混战'],['韩国','1-0','小胜, 防线再次零封'],['捷克','3-0','碾压, 希门尼斯+阿尔瓦拉多建功']]):
    dat(tg3,i+2,['']+r,bg=GRAY if i%2==0 else None)
tg3.cell(5,0).merge(tg3.cell(5,3));sc(tg3.cell(5,0),'厄瓜多尔 (E组第3, 4分, 0GD) 击败德国惊天逆转晋级',bold=True,color=DARK,align='left')
for i,r in enumerate([['科特迪瓦','0-1','90\'被绝杀, 0射正'],['库拉索','0-0','27射0球! 攻击效率灾难'],['德国','2-1','安古洛扳平→普拉塔77\'制胜! 惊天逆转']]):
    dat(tg3,i+6,['']+r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('首轮球员评分 + 首发调整分析',Pt(10),DARK)
txt('墨西哥: 希门尼斯 7.8/ 基尼奥内斯 7.5/ 阿尔瓦拉多 7.3/ 利拉 7.2/ 蒙特斯(C)7.4/ 巴斯克斯 7.3/ 兰赫尔 7.5(3场零封)/ 莫拉 7.0(新,17岁首秀X因素) | 埃德松·阿尔瓦雷斯 7.6(替补席! 下半场核武)',Pt(8.5))
txt('最低分: 中场创造力—莫拉17岁淘汰赛首发是赌博。但如果0-0僵持, 阿尔瓦雷斯下半场入场可瞬间改变中场上限。',Pt(8.5),color=GRAY_TEXT)
txt('厄瓜多尔: 凯塞多(C)7.8/ 普拉塔 7.8(改打ST双前锋!)/ 安古洛 7.5/ 欣卡皮耶 7.2/ 帕乔 7.0/ 瓦伦西亚 7.0/ 加林德斯 6.5/ 维特 6.8(新) | 派斯 7.2(替补席! 17岁天才下半场核武)',Pt(8.5))
txt('凯塞多(23岁切尔西队长)vs 莫拉(17岁墨超新秀)=本届淘汰赛最大的中场经验鸿沟。',Pt(8.5),bold=True,color=RED)
doc.add_paragraph()

heading('因素导向表',Pt(10),DARK)
tf3=doc.add_table(rows=13,cols=3);tf3.style='Table Grid';tf3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tf3,0,['因素','有利方','理由'])
for i,r in enumerate([
    ['墨西哥主场(阿兹台克80000+2240m海拔)','墨西哥 ★★★','世界杯主场9场不败! 海拔体能加成'],
    ['墨西哥3场零封(唯一全零封球队)','墨西哥 ★★★','防线默契+主场防守信心'],
    ['凯塞多(C) vs 17岁莫拉—经验鸿沟','厄瓜多尔 ★★★','凯塞多(切尔西)vs莫拉(墨超新秀)=本届最大中场经验差!'],
    ['埃德松·阿尔瓦雷斯替补(核武后手)','墨西哥 ★★','0-0僵持时上场→瞬间提升中场上限'],
    ['厄瓜多尔击败德国(信心爆棚)','厄瓜多尔 ★★','证明能击败任何对手'],
    ['厄瓜多尔4-4-2双前锋(非大巴心态)','双方进球↑ ★★','普拉塔+瓦伦西亚→比预期更具攻击性→降低0-0概率'],
    ['厄瓜多尔极度不稳定(27射0球vs2-1胜德国)','墨西哥 ★★','无法预测哪个版本会出现'],
    ['身价比1.4:1(接近比赛)','墨西哥 ★','均势比赛, 任何结果都可能(6/30教训)'],
    ['历史交锋(墨西哥17胜vs厄瓜多尔4胜)','墨西哥 ★','心理优势明确'],
    ['肯德里·派斯替补(厄瓜多尔后手)','厄瓜多尔 ★','17岁天才 vs 17岁莫拉=下半场年轻人对决'],
    ['蒙特斯戴上队长袖标主场作战','墨西哥 ★','防线领袖+主场加成=防守信心MAX'],
    ['瓦伦西亚距50球1球之遥(个人里程碑)','厄瓜多尔 ★','36岁老将最后一舞动力'],
]): dat(tf3,i+1,r,bg=GRAY if i%2==0 else None)
doc.add_paragraph()

heading('伤病/停赛',Pt(10),DARK)
tj3=doc.add_table(rows=7,cols=4);tj3.style='Table Grid';tj3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tj3,0,['球队','球员','状态','影响'])
for i,r in enumerate([
    ['墨西哥','全主力','[可]','11人全健康。无伤病困扰'],
    ['厄瓜多尔','全主力','[可]','击败德国的首发11人均可用'],
    ['墨西哥','埃德松·阿尔瓦雷斯','[替补]','战术决定非伤病。费内巴切赛季核心→下半场可入场'],
    ['厄瓜多尔','肯德里·派斯','[替补]','17岁天才→下半场核武。如果落后80分钟入场'],
    ['墨西哥','圣地亚哥·希门尼斯','[替补]','费耶诺德前锋, 下半场火力补充'],
    ['墨西哥','奥乔亚','[替补]','40岁6届老将, 如果进入点球大战是武器'],
]): dat(tj3,i+1,r,bg=GRAY if i%2==0 else None)
txt('关键观察: 双方都没有伤病困扰, 所有替补席上的关键球员都是战术选择而非被迫。双方教练都有完整手牌。',Pt(8.5),color=GRAY_TEXT)
doc.add_paragraph()

heading('厄瓜多尔韧性评估',Pt(10),DARK)
txt('3.5/5 — 低位防守:良好 | 速度反击:良好(普拉塔+安古洛,凯塞多推进) | 定位球:良好(帕乔+欣卡皮耶) | 前30分钟:一般(对德国早期丢球但迅速扳平) | 被压制不崩盘:良好(对德国被压制→逆转)',Pt(9))
txt('介于"不稳定"和"能爆冷"之间的灰色地带。对库拉索0-0(27射0球)和对德国2-1逆转的差距是世界杯最大的方差。',Pt(9),color=GRAY_TEXT)
doc.add_paragraph()

heading('比分预测',Pt(10),DARK)
tp3=doc.add_table(rows=5,cols=4);tp3.style='Table Grid';tp3.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(tp3,0,['类型','比分','半场','说明'])
for i,r in enumerate([
    ['首选','墨西哥 1-0','0-0','主场僵持→希门尼斯定位球/头球破局→零封保持'],
    ['备选','1-1(墨加时晋级)','0-0','[上调]厄瓜多尔双前锋→凯塞多反击助攻→墨西哥追平→加时主场优势'],
    ['备选','墨西哥 2-0','1-0','上半场破门→下半场厄瓜多尔压上→阿尔瓦拉多反击锁定'],
    ['备选','0-0(墨点球晋级)','0-0','[下调]厄瓜多尔双前锋不摆大巴→双方都有进攻意图→0-0概率低于1-1'],
]): dat(tp3,i+1,r,bold_cols={0} if r[0]=='首选' else None,bg=GREEN_BG if r[0]=='首选' else (GRAY if i%2==0 else None))
txt('DraftKings赔率验证: 墨西哥晋级64% vs 厄瓜多尔35%。大小球1.5线大球略占优→与1-0首选一致。',Pt(8.5),bold=True,color=DARK)
txt('首发后调整: 1-1上调(厄瓜多尔攻击意图明确) / 0-0下调(双方都不摆大巴)',Pt(8.5),color=GRAY_TEXT)
doc.add_paragraph()

heading('教练博弈(首发后更新)',Pt(10),DARK)
txt('墨西哥(阿吉雷): 用17岁莫拉而非阿尔瓦雷斯首发→想在开场阶段用创造力而非硬度。领先:442双后腰→可能换下莫拉上阿尔瓦雷斯。落后:阿尔瓦雷斯入场+圣地亚哥·希门尼斯双前锋+加强边路传中。',Pt(8.5))
txt('厄瓜多尔(贝卡塞塞): 普拉塔+瓦伦西亚双前锋(而非保守5后卫)→说明厄瓜多尔想进球。领先:5后卫死守+凯塞多拖后。落后:上派斯+埃斯图皮尼安→加速边路+天才创造力。',Pt(8.5))
txt('关键博弈: 阿吉雷赌莫拉—如果前30分钟莫拉被凯塞多完全压制→墨西哥中场失控→阿尔瓦雷斯提前入场(30-45分钟)。贝卡塞塞的4-4-2双前锋意味着后场只有2个中场(凯塞多+维特)→墨西哥中场人数3v2占优。',Pt(8.5),bold=True,color=RED)
doc.add_paragraph()

heading('定位球攻防',Pt(10),DARK)
txt('墨西哥进攻: 蒙特斯+巴斯克斯头球, 希门尼斯支点。角球在主场球迷助威下是重要武器。厄瓜多尔进攻: 帕乔(PSG)+欣卡皮耶(阿森纳)头球。凯塞多任意球。关键: 墨西哥1-0的路径最可能在定位球—希门尼斯或蒙特斯头球在阿兹台克高压氛围中会被放大。',Pt(8.5))
doc.add_paragraph()

heading('冷门路径',Pt(10),RED)
txt('墨西哥久攻不下→主场焦虑上升→凯塞多中场抢断→普拉塔反击1v1→0-1→墨西哥全线压上→厄瓜多尔大巴升级(5-4-1)→厄瓜多尔1-0或1-1(加时/点球胜)。最可能冷门比分: 1-1(厄瓜多尔加时或点球晋级)—这种模式与荷兰vs摩洛哥高度相似: 主场热门控球但不进球, 客场爆冷队大巴+反击+点球。',Pt(9))
doc.add_paragraph()

# ═══════════════ RISK WARNINGS ═══════════════
heading('风险提示',Pt(12),RED)
txt('1. 科特迪瓦vs挪威: "一个人的球队"vs"没有超巨但更均衡"的经典对决。哈兰德是否被锁死决定一切。',Pt(9),bold=True)
txt('2. 法国vs瑞典: 本日最没悬念, 但不是0悬念。瑞典三叉戟(伊萨克+哲凯赖什+埃兰加)有能力进球。法国R16对手是巴拉圭而非德国=巨大奖励。',Pt(9),bold=True)
txt('3. 墨西哥vs厄瓜多尔: [首发后不确定性上升] 17岁莫拉vs凯塞多=最大经验鸿沟, 但墨西哥替补席有阿尔瓦雷斯这个后手核武。厄瓜多尔双前锋表示他们来进球而非守平。核心悬念: 莫拉前30分钟活下来就是墨西哥的天才, 活不下来就是厄瓜多尔的突破口。DraftKings墨西哥64%晋级→与我们的方向一致。',Pt(9),bold=True)

doc.add_paragraph()
txt('─',Pt(8),color=RGBColor(0xAA,0xAA,0xAA))
txt('数据: ESPN API + FIFA API + Sports Mole + SI + RotoWire + Transfermarkt + DraftKings | 身价: Transfermarkt via Mundo Deportivo (2026年6月)',Pt(7),color=RGBColor(0x99,0x99,0x99))
txt('分析框架: CLAUDE.md v17 + 6月30日复盘教训 | 预测时间: 2026年6月30日 16:00 BJT | 更新: 2026年7月1日 08:05 BJT (首发确认+赔率)',Pt(7),color=RGBColor(0x99,0x99,0x99))
txt('Co-Authored-By: Claude Opus 4.8',Pt(7),color=RGBColor(0x99,0x99,0x99))

doc.save(OUT)
print(f'[OK] DOCX saved to: {OUT}')

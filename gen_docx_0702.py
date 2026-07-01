# -*- coding: utf-8 -*-
"""Generate 2026-07-02 World Cup predictions as a formatted .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月2日_3场预测.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# colour helpers
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GOLD   = RGBColor(0xD4, 0xA0, 0x17)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0xE8, 0xF8, 0xF5)
RED_BG     = RGBColor(0xE7, 0x4C, 0x3C)
GRAY_TEXT  = RGBColor(0x7F, 0x8C, 0x8D)

# ─── helper functions ───
def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('7月2日 三场预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('淘汰赛 Round of 32  |  ENG vs COD / BEL vs SEN / USA vs BIH', bold=False, size=Pt(11), color=GRAY_TEXT, align='center', space_after=Pt(18))

add_para('生成时间: 2026年7月1日 22:00 北京时间  |  数据源: ESPN API + Sporting News + Yahoo Sports + SI  |  框架: CLAUDE.md v17',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# ═══════════════════════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════════════════════
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '时间', '阶段', '比赛', '身价比', '首选', '半场', '备选', '冷门\n风险'],
    ['1', '00:00', 'R32', 'ENG 英格兰 vs COD 刚果民主共和国', '~16:1', '英格兰 2-0', '1-0', '3-0 / 2-1 / 1-1(加时)', '中低'],
    ['2', '04:00', 'R32', 'BEL 比利时 vs SEN 塞内加尔', '~1.9:1', '比利时 2-1', '1-0', '1-0 / 3-1 / 2-0', '中'],
    ['3', '08:00', 'R32', 'USA 美国 vs BIH 波黑', '~2.9:1', '美国 2-0', '1-0', '3-0 / 2-1 / 1-1(加时)', '中低'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 0.50, 0.42, 2.55, 0.70, 0.85, 0.60, 1.80, 0.55]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

# header row
for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

# data rows
for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        align = 'left' if j == 3 else 'center'
        bold = (j == 5)
        if j == 3:
            set_cell_left(table.rows[i].cells[j], text, bold=False, size=Pt(8.5))
        else:
            set_cell(table.rows[i].cells[j], text, bold=bold, size=Pt(8.5), bg=bg)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. KNOCKOUT BRACKET PATHS
# ═══════════════════════════════════════════════════
add_heading_styled('淘汰赛路径（7月1日收束后）', level=2)

bracket_lines = [
    '巴西 → R16 → 挪威 (淘汰科特迪瓦2-1)',
    '巴拉圭 (淘汰德国) → R16 → 法国 (淘汰瑞典3-0)',
    '摩洛哥 (淘汰荷兰) → R16 → 加拿大',
    '墨西哥 (2-0厄瓜多尔) → R16 → 英格兰/刚果民主共和国 胜者',
    '',
    '7月2日路径:',
    '  英格兰/刚果民主共和国 → R16 → 墨西哥 → 八强',
    '  比利时/塞内加尔 → R16 → 美国/波黑 胜者',
    '  美国/波黑 → R16 → 比利时/塞内加尔 胜者',
]
for line in bracket_lines:
    add_para(line, size=Pt(8.5), color=DARK if line.startswith('7月') else RGBColor(0x55,0x55,0x55), space_after=Pt(2))

# Path difficulty table
add_para('路径难度分析', bold=True, size=Pt(10), color=DARK, space_after=Pt(4))
path_data = [
    ['路径', '难度', '说明'],
    ['英格兰方向 → 墨西哥 → 八强', '中', '墨西哥主场是硬核考验，但路径避开法国区'],
    ['比利时方向 → 可能美国', '中', '美国主场=硬仗，比利时不稳定'],
    ['塞内加尔方向 → 可能波黑', '可控', '波黑是本日最弱热门，塞内加尔有真实机会进R16'],
    ['美国方向 → 比利时/塞内加尔', '中', '如果过了波黑，R16无论哪个对手都可打'],
]
pt = doc.add_table(rows=len(path_data), cols=3)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(path_data[0]):
    set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(path_data[1:], 1):
    set_cell_left(pt.rows[i].cells[0], row[0], size=Pt(8.5))
    set_cell(pt.rows[i].cells[1], row[1], size=Pt(8.5))
    set_cell_left(pt.rows[i].cells[2], row[2], size=Pt(8.5))
for row in pt.rows:
    row.cells[0].width = Inches(4.0)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(5.5)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. MATCH DETAILS
# ═══════════════════════════════════════════════════

matches = [
    {
        'num': 1, 'time': '00:00', 'stage': 'R32',
        'home': '英格兰', 'away': '刚果民主共和国',
        'home_flag': 'ENG', 'away_flag': 'COD',
        'venue': '梅赛德斯-奔驰体育场, 亚特兰大, 美国',
        'fifa_rank': '英格兰 #4 vs 刚果民主共和国 ~#61',
        'value_ratio': '英格兰 1.31B vs 刚果民主共和国 ~80M ≈ ~16:1',
        'history': '首次世界杯交锋',
        'reward': '16强对阵 墨西哥 (墨西哥2-0淘汰厄瓜多尔已在等待)',

        'group_review_home': '英格兰 (L组第1, 7分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['克罗地亚', '4-2', '凯恩2球(12\'点球+42\'头球)，贝林厄姆+拉什福德各1球'],
            ['加纳', '0-0', '英格兰第二场综合征再次发作！凯恩97分钟头球中横梁'],
            ['巴拿马', '2-0', '贝林厄姆62\'+凯恩67\'，下半场才破局'],
        ],
        'group_review_away': '刚果民主共和国 (K组第3, 4分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['葡萄牙', '1-1', '大巴5-3-2成功逼平葡萄牙！维萨(Wissa)进球'],
            ['哥伦比亚', '0-1', '惜败，防线持续稳固'],
            ['乌兹别克斯坦', '3-1', '变阵4后卫进攻→维萨2球，巴坎布进球'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['哈里·凯恩(Harry Kane) [火] MOTM', '8.5', '3球，英格兰世界杯历史射手王(11球)，点球+头球+跑位全能'],
            ['祖德·贝林厄姆(Jude Bellingham)', '8.2', '2球1助，决定性关键传球，图赫尔体系中前场自由人'],
            ['马库斯·拉什福德(Marcus Rashford)', '7.5', '对克罗地亚进球，左路速度+内切'],
            ['德克兰·赖斯(Declan Rice)', '7.8', '对巴拿马轮休后回归，中场屏障无可替代'],
            ['布卡约·萨卡(Bukayo Saka)', '7.2', '跟腱问题管理出场，但仍是右路最大威胁'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['约安·维萨(Yoane Wissa) [火]', '8.0', '3球！纽卡斯尔前锋，刚果民主共和国唯一进攻火力点'],
            ['塞德里克·巴坎布(Cedric Bakambu)', '7.3', '57球国家队，对乌兹别克斯坦进球，经验领袖'],
            ['阿隆·万-比萨卡(Aaron Wan-Bissaka)', '7.5', '右翼卫，前英格兰青年国脚，英超防守能力顶级'],
            ['阿克塞尔·图安泽贝(Axel Tuanzebe)', '7.0', '中卫，前英格兰U21，知道英格兰球员的踢法'],
            ['尚塞尔·姆本巴(Chancel Mbemba)', '7.2', '中卫，马赛队长，防守领袖'],
        ],

        'factors': [
            ('身价比~16:1(英格兰1.31B vs 刚果民主共和国~80M)', '英格兰 [火][火][火]', '大差距。强队方可做≥2球预测'),
            ('英格兰右后卫危机(詹姆斯+宽萨双伤)', '刚果民主共和国 [火][火]', '斯彭斯淘汰赛首秀→维萨将从这侧冲击'),
            ('刚果民主共和国5-3-2大巴已验证(逼平葡萄牙)', '刚果民主共和国 [火][火]', '本届已证明能逼平超级强队'),
            ('凯恩+贝林厄姆的破大巴组合', '英格兰 [火][火][火]', '凯恩禁区终结+贝林厄姆禁区边缘远射=破大巴最佳二人组'),
            ('赖斯回归(对巴拿马轮休)', '英格兰 [火][火]', '中场屏障恢复→刚果民主共和国反击起始被切断'),
            ('万-比萨卡+图安泽贝=前英格兰青年国脚', '刚果民主共和国 [火]', '心理层面: 证明自己的额外动力。维萨也在英超'),
            ('图赫尔10战11零封(含对巴拿马零封)', '英格兰 [火][火]', '但被克罗地亚2球打穿→防线非铁板'),
            ('刚果民主共和国对葡萄牙1-1=心理已验证', '刚果民主共和国 [火][火]', '不会因对手牌面产生畏惧'),
        ],
        'strength': '英格兰: 超级巨星型 — 凯恩(100M+)+贝林厄姆(180M+)+萨卡(150M+)+赖斯(110M+) = 不止2名超巨  |  刚果民主共和国: 大巴型+单点反击',
        'resilience_label': '非洲韧性评估(刚果民主共和国)',
        'resilience': [
            ('低位防守', '特优', '对葡萄牙1-1 = 本届非洲球队最佳大巴展示之一'),
            ('速度反击', '良好', '维萨(纽卡斯尔英超速度)，巴坎布经验'),
            ('定位球高点', '良好', '姆本巴+图安泽贝头球，但未在本届转化为进球'),
            ('前30分钟', '良好', '对葡萄牙0-0开场，对哥伦比亚0-0开场'),
            ('被压制不崩盘', '良好', '对哥伦比亚0-1惜败而非崩盘'),
        ],
        'upset_risk': '中低',
        'upset_detail': '英格兰是超级大热门，但存在两个真实裂缝:[注意] 右后卫危机: 詹姆斯(切尔西)腿筋伤→可能报销。宽萨(利物浦)脚踝伤→已离场。斯彭斯(热刺替补)被迫在淘汰赛首发。[注意] 刚果民主共和国的5-3-2已验证: 对葡萄牙1-1的战术纪律可直接复制。但是: 刚果民主共和国进攻只有维萨一人，16:1身价比=不做0-0首选。',

        'predictions': [
            ('首选', '英格兰 2-0', '1-0', '上半场凯恩破大巴→下半场贝林厄姆锁定→零封保持'),
            ('备选', '英格兰 3-0', '1-0', '如果斯彭斯端稳固→刚果民主共和国完全无反击→多点开花'),
            ('备选', '英格兰 2-1', '1-1', '维萨反击偷一个→英格兰实力碾压逆转'),
            ('备选', '1-1 (英格兰加时晋级)', '0-0', '刚果民主共和国复制葡萄牙大巴→斯彭斯侧被靶向→加时英格兰深度制胜'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['英格兰', '里斯·詹姆斯(Reece James)', '缺阵', '腿筋，可能报销。切尔西右后卫→防线最大漏洞'],
            ['英格兰', '贾雷尔·宽萨(Jarell Quansah)', '缺阵', '脚踝。利物浦，RB第二选择也伤了'],
            ['英格兰', '德克兰·赖斯(Declan Rice)', '可用', '小腿恢复。对巴拿马轮休后满血'],
            ['英格兰', '约翰·斯通斯(John Stones)', '可用', '带球能力对破大巴关键'],
            ['刚果民主共和国', '—', '可用', '德萨布尔有完整的26人大名单可选'],
        ],

        'coach': [
            '英格兰(图赫尔): 4-2-3-1凯恩支点+贝林厄姆自由人+萨卡拉什福德双边爆破。领先: 赖斯+安德森双后腰。落后: 凯恩+托尼双中锋冲吊。',
            '刚果民主共和国(德萨布尔): 法国教练，5-3-2大巴忠实信徒。对葡萄牙已验证→对英格兰必然复制。领先: 5-4-1收缩。落后: 变阵4后卫但可能不敢。',
            '关键博弈: 如果0-0到60分钟→德萨布尔可能换第二前锋→图赫尔上托尼双塔。但英格兰右后卫是未知数。',
        ],

        'set_piece': '英格兰进攻: 凯恩头球+贝林厄姆禁区边缘+斯通斯定位球跑位。刚果民主共和国进攻: 姆本巴+图安泽贝头球。关键: 图赫尔治下的英格兰定位球效率极高——凯恩对克罗地亚头球+对巴拿马头球都是定位球/传中。',

        'upset_path': '如果翻车: 英格兰久攻不下→斯彭斯侧被维萨单人反击击穿→0-1→刚果民主共和国升级至5-4-1→英格兰全线压上但大巴已验证逼平葡萄牙→0-1或1-1(加时/点球)。最可能冷门比分: 1-1 (刚果民主共和国加时或点球胜)——但需要: 斯彭斯严重失常+凯恩+贝林厄姆都哑火+刚果民主共和国大巴维持120分钟。三者同时发生概率不高。',
    },
    {
        'num': 2, 'time': '04:00', 'stage': 'R32',
        'home': '比利时', 'away': '塞内加尔',
        'home_flag': 'BEL', 'away_flag': 'SEN',
        'venue': '流明球场, 西雅图, 美国',
        'fifa_rank': '比利时 #5 vs 塞内加尔 #20',
        'value_ratio': '比利时 530M vs 塞内加尔 ~280M ≈ 1.9:1',
        'history': '首次世界杯交锋',
        'reward': '16强对阵 美国 vs 波黑 的胜者',

        'group_review_home': '比利时 (G组第1, 5分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['埃及', '1-1', '被动——靠卢卡库上场20秒逼出乌龙才扳平。德布劳内被萨拉赫压制'],
            ['伊朗', '0-0', '恩戈伊67\'红牌→10人顶住→伊朗门将MVP。比利时进攻完全哑火'],
            ['新西兰', '5-1', '最后一场才爆发！德布劳内+特罗萨德2球+卢卡库+萨勒马科尔斯'],
        ],
        'group_review_away': '塞内加尔 (I组第3, 3分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['法国', '1-3', '被法国全面碾压，萨尔任意球挽回颜面'],
            ['挪威', '2-3', '萨尔2球，门将门迪膝盖重伤离队！防线失3球'],
            ['伊拉克', '5-0', '萨尔再1球+马内+杰克逊+迪亚拉，非洲球队单场5球纪录'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['凯文·德布劳内(Kevin De Bruyne) [火]', '8.0', '创造核心，但对埃及和伊朗被成功限制。对新西兰1球+统治中场'],
            ['莱安德罗·特罗萨德(Leandro Trossard)', '7.8', '2球，本届创造机会最多(13次)'],
            ['罗梅卢·卢卡库(Romelu Lukaku)', '7.2', '对新西兰进球，前两场替补。预计继续替补→德凯特拉雷伪9号首发'],
            ['蒂博·库尔图瓦(Thibaut Courtois)', '7.5', '世界级门将，对伊朗多次关键扑救保零封'],
            ['查尔斯·德凯特拉雷(Charles De Ketelaere)', '7.0', '伪9号首发，对阵新西兰的表现说服了教练'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['伊斯梅拉·萨尔(Ismaila Sarr) [火] MOTM', '8.5', '3球(对挪威2球+对伊拉克1球)，速度+终结，本届非洲最佳攻击手'],
            ['萨迪奥·马内(Sadio Mane)', '7.2', '对伊拉克进球，利雅得胜利赛季验证'],
            ['尼古拉·雅克松(Nicolas Jackson)', '7.0', '对伊拉克进球，切尔西前锋'],
            ['阿卜杜拉耶·塞克(Abdoulaye Seck)', '7.2', '对伊拉克零封，替代库利巴利位置'],
            ['莫里·迪奥(Mory Diaw)', '6.5', '替补门将! 门迪伤退后接手，对伊拉克零封但对法国+挪威丢6球'],
        ],

        'factors': [
            ('门迪重伤缺阵(塞内加尔世界级门将)', '比利时 [火][火][火]', '迪奥(克莱蒙)vs德布劳内/特罗萨德射门→能力差至少1档'),
            ('比利时前两场0运动战进球(面对大巴瘫痪)', '塞内加尔 [火][火]', '如果塞内加尔摆大巴→比利时可能再次哑火'),
            ('萨尔3球(本届非洲最佳攻击手)', '塞内加尔 [火][火][火]', '速度+终结，比利时防线(默尔+特亚特)非世界级'),
            ('德布劳内创造能力(对新西兰证明)', '比利时 [火][火][火]', '塞内加尔非大巴型防守(对法国丢3球)→比利时会有空间'),
            ('身价比1.9:1(<3:1接近)', '持平', '不是压倒性差距，任何结果都可能'),
            ('塞内加尔攻击力非洲最强(8球)', '塞内加尔 [火][火]', '这是他们最大的武器——比利时防线不是大巴型'),
        ],
        'strength': '比利时: 体系型(偏低效热门)——德布劳内(55M) = 唯一超巨。前两场0运动战进球=面对大巴体系完全无效  |  塞内加尔: 攻击型非洲劲旅——萨尔(60M+)+马内+杰克逊=本届非洲最强攻击群(8球)',
        'resilience_label': '非洲韧性评估(塞内加尔)',
        'resilience': [
            ('低位防守', '一般', '对法国丢3球+对挪威丢3球=欧洲球队可以击穿'),
            ('速度反击', '特优', '萨尔(3球)+马内+杰克逊=本届非洲最快的反击三叉戟'),
            ('定位球高点', '良好', '塞克+尼亚卡特+库利巴利头球，但门迪缺阵削弱防空'),
            ('前30分钟', '差', '对法国和挪威都是早期丢球'),
            ('被压制不崩盘', '良好', '0-3落后法国没有崩(最终1-3)，对挪威2-3惜败'),
        ],
        'upset_risk': '中',
        'upset_detail': '本日最不确定的比赛——比利时不稳定(前两场0运动战进球) vs 塞内加尔攻击力(非洲最强8球) + 门迪缺阵。核心矛盾: 比利时面对大巴时进攻瘫痪。塞内加尔会不会摆大巴？如果不会(给比利时空间)→比利时5-1新西兰证明他们能得分。但萨尔3球+马内的反击速度→比利时防线有真实被打穿风险。',

        'predictions': [
            ('首选', '比利时 2-1', '1-0', '上半场德布劳内/特罗萨德破门→萨尔反击扳平→下半场比利时深度制胜'),
            ('备选', '比利时 1-0', '0-0', '极度僵持→德布劳内个人能力→零封守住'),
            ('备选', '比利时 3-1', '1-1', '双方对攻→比利时火力优势→塞内加尔门将差距暴露'),
            ('备选', '比利时 2-0', '1-0', '塞内加尔大巴开局→比利时早破门→反击第二球锁定'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['塞内加尔', '爱德华·门迪(Edouard Mendy)', '缺阵', '膝盖重伤已离队！世界级门将→迪奥(克莱蒙)=最大降级'],
            ['比利时', '泽诺·德巴斯特(Zeno Debast)', '疑', '腿伤，预计不首发'],
            ['比利时', '纳坦·恩戈伊(Nathan Ngoy)', '可用', '红牌停赛结束，可用但非首选'],
            ['比利时', '杰雷米·多库(Jeremy Doku)', '可用', '孩子出生后归队，边路爆点回归'],
            ['比利时', '罗梅卢·卢卡库(Romelu Lukaku)', '可用', '满血替补，下半场核武后手'],
            ['塞内加尔', '卡利杜·库利巴利(Kalidou Koulibaly)', '疑', '对挪威表现不佳，塞克可能顶替'],
        ],

        'coach': [
            '比利时(特德斯科): 战术一直在调整。前两场4-2-3-1无效→第三场变阵德凯特拉雷伪9号+激活德布劳内。领先: 蒂勒曼斯+瓦纳肯双后腰。落后: 卢卡库入场+双塔冲吊+多库边路爆破。',
            '塞内加尔(西塞): 非洲杯冠军教练。战术: 4-3-3萨尔的左路速度反击为核心。领先: 5-4-1收缩。落后: 马内→前锋+增加第二攻击手。门迪缺阵=他必须用4后卫保护迪奥(门将)。',
            '关键博弈: 如果塞内加尔对攻→比利时有空间(像对新西兰)。如果摆大巴→比利时可能重演对埃及/伊朗的哑火。',
        ],

        'set_piece': '比利时进攻: 德布劳内任意球+特罗萨德远射+卢卡库(如果上场)头球。塞内加尔进攻: 萨尔任意球+塞克头球+马内跑位。关键: 门迪缺阵对定位球防守的影响——迪奥的扑救范围和控制禁区能力远不如门迪。比利时定位球是破大巴的最佳途径之一。',

        'upset_path': '如果翻车: 塞内加尔对攻→萨尔速度打穿比利时慢防线→1-0→比利时恐慌→德布劳内被锁→塞内加尔大巴升级→1-0或2-1。塞内加尔赢这场不是超级冷——两队身价比仅1.9:1，门迪缺阵是比利时的幸运但萨尔状态是塞内加尔的武器。最可能冷门比分: 塞内加尔 2-1 (萨尔+马内各1球)。',
    },
    {
        'num': 3, 'time': '08:00', 'stage': 'R32',
        'home': '美国', 'away': '波黑',
        'home_flag': 'USA', 'away_flag': 'BIH',
        'venue': '李维斯体育场, 圣克拉拉, 加利福尼亚, 美国',
        'fifa_rank': '美国 #15 vs 波黑 #45',
        'value_ratio': '美国 ~350M vs 波黑 ~120M ≈ 2.9:1',
        'history': '首次世界杯交锋',
        'reward': '16强对阵 比利时 vs 塞内加尔 的胜者',

        'group_review_home': '美国 (D组第1, 6分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['巴拉圭', '4-1', '雷纳世界波，主场气势完全碾压'],
            ['澳大利亚', '2-0', '伯吉斯乌龙+弗里曼头球，零封'],
            ['土耳其', '2-3', '轮换首发(已锁定第1)→2-0领先被3-2逆转！阿伊汉90+8\'绝杀'],
        ],
        'group_review_away': '波黑 (B组第3, 4分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['加拿大', '1-1', '平局'],
            ['瑞士', '1-4', '防线被瑞士进攻完全打穿'],
            ['卡塔尔', '3-1', '0-1落后→历史性逆转→淘汰赛门票'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['克里斯蒂安·普利西奇(Christian Pulisic) [火]', '8.0', '对土耳其替补出场(小腿恢复)→满血首发'],
            ['福拉林·巴洛贡(Folarin Balogun)', '7.5', '摩纳哥前锋，对澳大利亚制造乌龙，支点作用'],
            ['乔瓦尼·雷纳(Gio Reyna)', '7.8', '对巴拉圭世界波，多特框架验证'],
            ['韦斯顿·麦肯尼(Weston McKennie)', '7.3', '前场自由人，尤文图斯'],
            ['泰勒·亚当斯(Tyler Adams)', '7.5', '防守中场，伯恩茅斯，对土耳其轮休'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['埃丁·哲科(Edin Dzeko) [火]', '7.3', '38岁费内巴切，国家队传奇'],
            ['埃斯米尔·巴伊拉科塔雷维奇(Esmir Bajraktarevic)', '7.5', '21岁威斯康星出生→选择波黑→对意大利资格赛罚进制胜点球'],
            ['埃尔梅丁·德米罗维奇(Ermedin Demirovic)', '7.0', '奥格斯堡前锋'],
            ['尼古拉·瓦西利(Nikola Vasilj)', '7.0', '门将，圣保利'],
            ['塔里克·穆哈雷莫维奇(Tarik Muharemovic)', '7.0', '中卫，对瑞士停赛后回归=防线升级'],
        ],

        'factors': [
            ('美国主场(圣克拉拉+全美支持)', '美国 [火][火][火]', '2002年以来首个淘汰赛胜利的机会，10万人级别的国家支持'),
            ('普利西奇满血回归(小腿恢复)', '美国 [火][火][火]', '队长+进攻核心，对土耳其替补出场已恢复'),
            ('波黑是最高球队(188cm平均)+46次犯规', '波黑 [火][火]', '定位球+身体对抗=波黑唯二武器。美国需避免不必要定位球'),
            ('巴伊拉科塔雷维奇"回家"(威斯康星出生)', '波黑 [火][火]', '21岁前美国青年国脚→对意大利点球制胜的心理素质'),
            ('美国全主力休息了半场(对土耳其轮换)', '美国 [火][火]', '亚当斯+巴洛贡+罗宾逊+理查兹+普利西奇皆满血'),
            ('波黑对瑞士1-4(面对强队防线崩溃)', '美国 [火][火]', '如果美国早破门→波黑可能重演对瑞士的崩溃'),
            ('美国从未击败欧洲队(2021年至今)', '波黑 [火][火]', '尴尬纪录——对美国是真实心理负担'),
            ('身价比2.9:1(接近3:1但<3:1)', '美国 [火]', '美国占优但不碾压'),
        ],
        'strength': '美国: 体系型(主场加成版)——普利西奇(50M+)最接近超巨，无≥2名超巨。但主场+淘汰赛+2002年以来首胜机会=额外加成  |  波黑: 大巴+定位球型——本届最高球队(188cm)，战术: 低姿态+定位球=唯二进攻路径',
        'resilience_label': '欧洲球队特性分析(波黑)',
        'resilience': [
            ('低位防守', '一般', '对瑞士1-4=面对高质量进攻时防线脆弱'),
            ('速度反击', '差', '波黑不是速度型——他们是身体对抗型'),
            ('定位球高点', '特优', '本届最高球队+哲科头球=定位球是最可能得分路径'),
            ('前30分钟', '一般', '对加拿大早期丢球，对瑞士1-4'),
            ('被压制不崩盘', '差', '对瑞士连续丢球——但3-1逆转卡塔尔展示部分韧性'),
        ],
        'upset_risk': '中低',
        'upset_detail': '美国是明显热门——主场+全主力休息。但两个隐患: 1. 对欧洲队的诅咒(2021年至今未胜)。2. 波黑身体优势(188cm+46次犯规)。但是: 波黑对瑞士1-4证明面对高质量进攻的脆弱。普利西奇+巴洛贡+雷纳+麦肯尼应该创造大量机会。',

        'predictions': [
            ('首选', '美国 2-0', '1-0', '上半场破大巴→下半场反击锁定→零封'),
            ('备选', '美国 3-0', '2-0', '早破门→波黑被迫攻出→像对瑞士1-4般崩溃'),
            ('备选', '美国 2-1', '1-1', '哲科定位球头球偷一个→美国实力逆转'),
            ('备选', '1-1 (美国加时晋级)', '0-0', '波黑大巴成功→但120分钟波黑体能(高龄哲科)是问题'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['美国', '克里斯蒂安·普利西奇(Christian Pulisic)', '可用', '小腿恢复满血。队长首发，对土耳其替补已验证'],
            ['美国', '马克·麦肯齐(Mark McKenzie)', '疑', '脚伤，可能不首发。中卫轮换'],
            ['美国', '克里斯蒂安·罗尔丹(Cristian Roldan)', '疑', '部分训练。中场轮换'],
            ['美国', '奥斯顿·特拉斯蒂(Auston Trusty)', '疑', '脚踝。中卫轮换，对土耳其进球但受伤'],
            ['波黑', '塔里克·穆哈雷莫维奇(Tarik Muharemovic)', '可用', '停赛结束回归。防线领袖→波黑防线升级'],
            ['波黑', '阿马尔·德迪奇(Amar Dedic)', '疑', '恢复训练。RB可能轮换'],
        ],

        'coach': [
            '美国(波切蒂诺): 对土耳其失利后向媒体道歉——知道淘汰赛的期望值。战术: 4-2-3-1普利西奇左路自由+巴洛贡支点+雷纳/麦肯尼创造力。领先: 亚当斯+蒂尔曼双后腰。落后: 替补席深度(阿伦森+佩皮+卡多索)是美国最大优势之一。',
            '波黑(哈吉贝吉奇): 防守优先+定位球=一切。领先: 5-4-1极致收缩+全员回防。落后: 哲科+德米罗维奇双塔+增加边路传中。',
            '关键博弈: 波切蒂诺需要早进球。如果0-0到60分钟→美国球迷焦虑+波黑大巴升级→可能重演荷兰vs摩洛哥模式。但美国应该有意识在前30分钟高强度施压。',
        ],

        'set_piece': '美国进攻: 理查兹+雷纳头球，普利西奇任意球精确度。角落球是对波黑大巴的重要破局手段。波黑进攻: 哲科(193cm)+德米罗维奇+卡蒂奇=波黑最危险的武器。关键: 美国需避免在防守三区给不必要犯规→每个定位球都是波黑的进球机会。',

        'upset_path': '如果翻车: 美国久攻不下→波黑角球→哲科头球→0-1→美国全线压上→波黑升级5-4-1→美国对欧洲队的魔咒心理加重→0-1或1-1(加时/点球)。最可能冷门比分: 1-1 (波黑加时或点球晋级)——但需要: 美国进攻完全哑火+波黑定位球效率100%+美国对欧洲队魔咒持续。概率偏低但不是零。',
    },
]

# ── Render each match ──
for m in matches:
    add_heading_styled(f'比赛{m["num"]}: {m["home_flag"]} {m["home"]} vs {m["away_flag"]} {m["away"]}  ({m["stage"]} {m["time"]} BJT)', level=2)

    # ── Basic info ──
    add_para('基本信息', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))
    info_lines = [
        f'场地: {m["venue"]}',
        f'FIFA排名: {m["fifa_rank"]}  |  身价比: {m["value_ratio"]}',
        f'历史交锋: {m["history"]}  |  晋级奖励: {m["reward"]}',
    ]
    for line in info_lines:
        add_para(line, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # ── Group stage review ──
    add_para('小组赛回顾', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    # Home team review
    add_para(m['group_review_home'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_home_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_home_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_home_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    # Away team review
    add_para(m['group_review_away'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_away_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_away_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_away_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    # ── Player ratings ──
    add_para('首轮球员评分', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    add_para(f'{m["home"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_home']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_home'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_home'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    add_para(f'{m["away"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_away']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_away'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_away'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    # ── Factors table ──
    add_para('因素导向表', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    ft_data = [['因素', '有利方', '理由']]
    ft_data.extend(m['factors'])
    ft = doc.add_table(rows=len(ft_data), cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(ft_data[0]):
        set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(8.5), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(ft_data[1:], 1):
        set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8))
        clr = None
        if '[火][火][火]' in row[1] or '[火][火]' in row[1]:
            clr = ACCENT
        set_cell(ft.rows[i].cells[1], row[1], size=Pt(8), color=clr)
        set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in ft.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(5.0)
    doc.add_paragraph()

    # ── Strength + Resilience ──
    add_para(f'强队分类: {m["strength"]}', size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))

    add_para(m['resilience_label'], bold=True, size=Pt(9), color=DARK, space_after=Pt(2))
    res_data = [['维度', '评分', '说明']]
    res_data.extend(m['resilience'])
    rt = doc.add_table(rows=len(res_data), cols=3)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(res_data[0]):
        set_cell(rt.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(res_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(rt.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(rt.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(rt.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in rt.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(7.5)
    doc.add_paragraph()

    # ── Upset risk ──
    risk_color = ACCENT if m['upset_risk'] in ('中', '中高', '高') else DARK
    add_para(f'冷门风险: {m["upset_risk"]}', bold=True, size=Pt(10), color=risk_color, space_after=Pt(2))
    add_para(m['upset_detail'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # ── Predictions table ──
    add_para('比分预测', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    pred_data = [['类型', '比分', '半场', '说明']]
    pred_data.extend(m['predictions'])
    pt = doc.add_table(rows=len(pred_data), cols=4)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(pred_data[0]):
        set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(pred_data[1:], 1):
        bg = GREEN_BG if i == 1 else None
        set_cell(pt.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg)
        set_cell(pt.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
        set_cell(pt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
        set_cell_left(pt.rows[i].cells[3], row[3], size=Pt(8))
    for row in pt.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(7.4)
    doc.add_paragraph()

    # ── Injuries ──
    add_para('伤病/停赛', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    inj_data = m['injuries']
    it = doc.add_table(rows=len(inj_data), cols=4)
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(inj_data[0]):
        set_cell(it.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(inj_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(it.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[1], row[1], size=Pt(7.5))
        set_cell(it.rows[i].cells[2], row[2], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[3], row[3], size=Pt(7.5))
    for row in it.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(5.2)
    doc.add_paragraph()

    # ── Coach tactics ──
    add_para('教练博弈', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    for c in m['coach']:
        add_para(c, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # ── Set piece ──
    add_para('定位球攻防', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    add_para(m['set_piece'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))
    doc.add_paragraph()

    # ── Upset path ──
    add_para('冷门路径', bold=True, size=Pt(10), color=ACCENT, space_after=Pt(2))
    add_para(m['upset_path'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # separator between matches
    add_para('─' * 60, size=Pt(6), color=RGBColor(0xCC,0xCC,0xCC), align='center', space_after=Pt(8))

# ═══════════════════════════════════════════════════
# RISK SUMMARY
# ═══════════════════════════════════════════════════
add_heading_styled('风险提示', level=2)

risk_notes = [
    ('英格兰 vs 刚果民主共和国', '身价比16:1=本日最悬殊。但英格兰右后卫双伤(詹姆斯+宽萨)是真实漏洞。刚果民主共和国的维萨(纽卡斯尔3球)会精准靶向斯彭斯侧。如果英格兰不早进球，刚果民主共和国的大巴(已验证逼平葡萄牙)可能导致1-1加时。'),
    ('比利时 vs 塞内加尔', '本日最可能翻车的比赛。比利时前两场0运动战进球+面对大巴完全哑火，塞内加尔是非洲攻击力最强(8球)——但门迪(门将)缺阵。逻辑矛盾: 塞内加尔不是大巴队(给比利时空间)但攻击力足够打穿比利时防线。2-1首选承认双方都能进球。'),
    ('美国 vs 波黑', '美国主场是压倒性优势。但美国2021年至今未赢欧洲队=尴尬纪录。波黑最高的球队+定位球是唯一威胁。如果美国前30分钟进球→可能3-0。如果0-0到60分钟→警惕荷兰vs摩洛哥的重演。'),
]

for title, detail in risk_notes:
    add_para(title, bold=True, size=Pt(9), color=ACCENT, space_after=Pt(1))
    add_para(detail, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: ESPN API + Sporting News + Yahoo Sports + SI + OneFootball + Transfermarkt  |  框架: CLAUDE.md v17 (身价量化+强队三类+10项必填清单+7月1日复盘教训)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('预测时间: 2026年7月1日 22:00 北京时间  |  教训加权: 主场强队不做0-0半场+边路可绕过中场+2-0领先≠防守换人+一场爆冷≠能力质变',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('预测仅供分析参考  |  生成工具: python-docx',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

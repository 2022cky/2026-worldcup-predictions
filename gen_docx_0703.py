# -*- coding: utf-8 -*-
"""Generate 2026-07-03 World Cup predictions as a formatted .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月3日_3场预测_v3.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# colour helpers
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GOLD   = RGBColor(0xD4, 0xA0, 0x17)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0xE8, 0xF8, 0xF5)
RED_BG     = RGBColor(0xE7, 0x4C, 0x3C)
GRAY_TEXT  = RGBColor(0x7F, 0x8C, 0x8D)

# ─── helper functions ───
def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('7月3日 三场预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('淘汰赛 Round of 32  |  ESP vs AUT / POR vs CRO / SUI vs ALG', bold=False, size=Pt(11), color=GRAY_TEXT, align='center', space_after=Pt(18))

add_para('生成时间: 2026年7月2日 14:00 北京时间  |  数据源: ESPN API + Rotowire + Sports Mole + Yahoo Sports + Transfermarkt  |  框架: CLAUDE.md v16',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# ═══════════════════════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════════════════════
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '时间', '阶段', '比赛', '身价比', '首选', '半场', '备选', '冷门\n风险'],
    ['1', '03:00', 'R32', 'ESP 西班牙 vs AUT 奥地利', '~5:1', '西班牙 2-0', '1-0', '2-1 / 3-1 / 1-0', '中低'],
    ['2', '07:00', 'R32', 'POR 葡萄牙 vs CRO 克罗地亚', '~2.6:1', '葡萄牙 2-1', '1-0', '3-1 / 2-0 / 1-1(加时)', '低'],
    ['3', '11:00', 'R32', 'SUI 瑞士 vs ALG 阿尔及利亚', '~1.3:1', '瑞士 2-1', '1-0', '1-0 / 2-0 / 1-1(加时)', '中'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 0.50, 0.42, 2.55, 0.70, 0.85, 0.60, 1.80, 0.55]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

# header row
for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

# data rows
for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        align = 'left' if j == 3 else 'center'
        bold = (j == 5)
        if j == 3:
            set_cell_left(table.rows[i].cells[j], text, bold=False, size=Pt(8.5))
        else:
            set_cell(table.rows[i].cells[j], text, bold=bold, size=Pt(8.5), bg=bg)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. KNOCKOUT BRACKET PATHS
# ═══════════════════════════════════════════════════
add_heading_styled('淘汰赛路径（7月2日收束后更新）', level=2)

bracket_lines = [
    '英格兰 (2-1刚果民主共和国) → R16 → 墨西哥',
    '比利时 (3-2塞内加尔 AET) → R16 → 美国',
    '美国 (2-0波黑) → R16 → 比利时',
    '',
    '7月3日路径:',
    '  西班牙/奥地利 → R16 → 葡萄牙/克罗地亚 胜者',
    '  葡萄牙/克罗地亚 → R16 → 西班牙/奥地利 胜者',
    '  瑞士/阿尔及利亚 → R16 → 哥伦比亚/加纳 胜者',
]
for line in bracket_lines:
    add_para(line, size=Pt(8.5), color=DARK if line.startswith('7月') else RGBColor(0x55,0x55,0x55), space_after=Pt(2))

# Path difficulty table
add_para('路径难度分析', bold=True, size=Pt(10), color=DARK, space_after=Pt(4))
path_data = [
    ['路径', '难度', '说明'],
    ['西班牙方向 → 可能葡萄牙', '高', '伊比利亚半岛德比潜在八强——双方都想避开'],
    ['葡萄牙方向 → 可能西班牙', '高', '葡萄牙比西班牙更不稳定，但如过关则面对奥地利(更弱)'],
    ['瑞士方向 → 哥伦比亚/加纳', '中', '瑞士纸面热门但佩特科维奇知道瑞士的一切'],
]
pt = doc.add_table(rows=len(path_data), cols=3)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(path_data[0]):
    set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(path_data[1:], 1):
    set_cell_left(pt.rows[i].cells[0], row[0], size=Pt(8.5))
    set_cell(pt.rows[i].cells[1], row[1], size=Pt(8.5))
    set_cell_left(pt.rows[i].cells[2], row[2], size=Pt(8.5))
for row in pt.rows:
    row.cells[0].width = Inches(4.0)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(5.5)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. MATCH DETAILS
# ═══════════════════════════════════════════════════

matches = [
    {
        'num': 1, 'time': '03:00', 'stage': 'R32',
        'home': '西班牙', 'away': '奥地利',
        'home_flag': 'ESP', 'away_flag': 'AUT',
        'venue': 'SoFi体育场, 洛杉矶, 美国',
        'fifa_rank': '西班牙 #3 vs 奥地利 #22',
        'value_ratio': '西班牙 1.22B vs 奥地利 245M ≈ 5:1',
        'history': '西班牙2胜1平0负 (最近: 2009年友谊赛5-1)',
        'reward': '16强对阵 葡萄牙 vs 克罗地亚 的胜者',

        'group_review_home': '西班牙 (H组第1, 7分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['佛得角', '0-0', '74%控球→0进球。佛得角大巴成为本届防守教科书。亚马尔缺阵'],
            ['沙特阿拉伯', '4-0', '尼科·威廉姆斯2球+亚马尔1球+佩德里1球。西班牙火力全开'],
            ['乌拉圭', '1-0', '巴埃纳(Baena)进球，但代价惨重——威廉姆斯+皮诺双双受伤离场'],
        ],
        'group_review_away': '奥地利 (J组第2, 4分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['约旦', '3-1', '朗尼克的压迫体系见效'],
            ['阿根廷', '0-2', '面对超级豪门完全被压制'],
            ['阿尔及利亚', '3-3', '96分钟卡拉季奇绝平→戏剧性晋级！自1954年首次进入淘汰赛'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['拉明·亚马尔(Lamine Yamal) [火] MOTM', '8.5', '18岁200M，对沙特进球，对乌拉圭76分钟。本届最佳年轻球员候选人'],
            ['佩德里(Pedri)', '8.0', '对沙特进球，中场创造核心，150M'],
            ['罗德里(Rodri) [C]', '8.0', '世界最佳后腰，西班牙防线前不可逾越的屏障'],
            ['米克尔·奥亚萨瓦尔(Mikel Oyarzabal)', '7.8', '3次进球贡献(进球+助攻)，本届西班牙最稳定的攻击点'],
            ['亚历克斯·巴埃纳(Alex Baena)', '7.5', '对乌拉圭制胜球→将在威廉姆斯缺阵时首发左路'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['马塞尔·萨比策(Marcel Sabitzer) [火]', '7.8', '100场国家队出场，多特蒙德，奥地利进攻引擎'],
            ['马尔科·阿瑙托维奇(Marko Arnautovic)', '7.3', '经验丰富的支点中锋，对阵阿尔及利亚追平球'],
            ['康拉德·莱默(Konrad Laimer)', '7.5', '拜仁RB/CM，对阿根廷表现最稳定的奥地利球员'],
            ['萨沙·卡拉季奇(Sasa Kalajdzic)', '7.2', '2米高中锋，96分钟绝平英雄'],
            ['大卫·阿拉巴(David Alaba) [C]', '7.0', '轻微伤→预计首发。皇马经验，但本赛季出场时间有限'],
        ],

        'factors': [
            ('西班牙身价5:1(1.22B vs 245M)', '西班牙 [火][火][火]', '明显差距。强队方占优'),
            ('尼科·威廉姆斯缺阵(2球，核心爆破点)', '奥地利 [火][火][火]', '西班牙左路爆破能力消失→巴埃纳是内切型而非下底型→进攻维度减少'),
            ('西班牙0失球(3场小组赛)', '西班牙 [火][火][火]', '罗德里+库巴西+拉波尔特的防线是本届最强。'),
            ('奥地利12场世界杯0零封', '西班牙 [火][火]', '西班牙即使进攻效率低也有足够时间破门'),
            ('朗尼克的gegenpressing体系', '奥地利 [火][火]', '高位压迫可以扰乱西班牙控球体系→制造转换机会'),
            ('亚马尔满血回归(对乌拉圭76分钟)', '西班牙 [火][火][火]', '18岁200M天才，即使威廉姆斯缺阵也有单兵爆破能力'),
            ('波施破碎下巴+定制面罩', '西班牙 [火][火]', '西班牙边路会靶向波施侧→面部是身体对抗中最脆弱的部位'),
            ('卡拉季奇2米空中威胁', '奥地利 [火][火]', '西班牙防线没有真正的空中巨人→定位球是奥地利最可能得分路径'),
        ],
        'strength': '西班牙: 超级巨星型——亚马尔(200M)+佩德里(150M)+罗德里(130M) = 不止2名超巨。但威廉姆斯缺阵让进攻从超级巨星型向体系型偏移半档  |  奥地利: 压迫+定位球型——朗尼克的gegenpressing体系+卡拉季奇2米空中威胁',
        'resilience_label': '欧洲球队特性分析(奥地利)',
        'resilience': [
            ('低位防守', '一般', '对阿根廷0-2→面对顶级进攻无法零封。12场世界杯0零封'),
            ('速度反击', '良好', '萨比策+莱默的转换速度是德甲验证的'),
            ('定位球高点', '特优', '卡拉季奇(2米)+阿瑙托维奇(192cm)+阿拉巴=欧洲最高定位球威胁之一'),
            ('前30分钟', '一般', '对阿根廷早期丢球，对约旦稳住'),
            ('被压制不崩盘', '良好', '对阿根廷0-2未崩盘，对阿尔及利亚1-0领先后被3-3逆转'),
        ],
        'upset_risk': '中低',
        'upset_detail': '西班牙是明显热门——0失球防线+34场不败+亚马尔状态回升。但两个真实裂缝: 1. 威廉姆斯缺阵→进攻宽度消失，巴埃纳替代=更内切更技术性→奥地利可收缩中路防守。2. 佛得角0-0教训——西班牙74%控球但0进球vs大巴。奥地利不会像对阿根廷那样开放。但是: 西班牙0失球防线+奥地利从未零封=大概率西班牙会在某个时间点破门。5:1身价比=不做0-0首选。',

        'predictions': [
            ('首选', '西班牙 2-0', '1-0', '[历史数据]先进球方74%获胜→西班牙更强→更可能先进球→更可能赢。亚马尔+佩德里破大巴→零封'),
            ('强备选', '西班牙 2-1', '1-0', '罗德里重伤+卡拉季奇近10场13球贡献→奥地利定位球得分概率真实(19-22%少数场景)'),
            ('备选', '西班牙 3-1', '1-1', '奥地利先进球→西班牙实力碾压→多点开花'),
            ('备选', '西班牙 1-0', '0-0', '佛得角模式重演→亚马尔最后一刻破局'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['西班牙', '尼科·威廉姆斯(Nico Williams)', '缺阵', '右内收肌伤。80M+2球1助→左路爆点缺失'],
            ['西班牙', '耶雷米·皮诺(Yeremy Pino)', '疑→可用', '肩伤奇迹恢复→训练正常→可替补出场'],
            ['西班牙', '维克托·穆尼奥斯(Victor Munoz)', '疑', '小腿伤。中场轮换，非首发'],
            ['西班牙', '拉明·亚马尔(Lamine Yamal)', '可用', '伤愈满血→76分钟vs乌拉圭验证'],
            ['奥地利', '斯特凡·波施(Stefan Posch)', '可用', '破碎下巴+面罩出战→勇敢但面部脆弱'],
            ['奥地利', '大卫·阿拉巴(David Alaba)', '可用', '轻微伤→预计首发→皇马经验无可替代'],
            ['奥地利', '马尔科·阿瑙托维奇(Marko Arnautovic)', '可用', '轻微伤→预计可战'],
        ],

        'coach': [
            '西班牙(路易斯·德拉富恩特): 4-3-3控球+亚马尔右路爆破+佩德里创造。关键调整: 奥尔莫取代梅里诺首发→增加前场创造力以弥补威廉姆斯缺阵。领先: 罗德里单后腰控节奏。落后: 皮诺(如可出场)+费兰·托雷斯+更多边路传中。',
            '奥地利(拉尔夫·朗尼克): Gegenpressing教父。战术: 4-2-3-1收缩防守+高位压迫转换。领先: 5-4-1极致大巴(卡拉季奇支点)。落后: 卡拉季奇+阿瑙托维奇双塔+边路传中轰炸。',
            '关键博弈: 西班牙需要早进球——0-0到60分钟→奥地利信心增长+定位球偷一个→世界杯最大冷门之一。奥尔莫首发=进攻信号。',
        ],

        'set_piece': '西班牙进攻: 拉波尔特+库巴西头球，亚马尔+佩德里任意球精度。罗德里禁区边缘远射。威廉姆斯缺阵→失去一个定位球接应点。奥地利进攻: 卡拉季奇(2米!)+阿瑙托维奇(192cm)+阿拉巴+林哈特=本届最高定位球进攻群之一。萨比策任意球+角球传中精确度。关键: 西班牙需避免在防守三区给不必要犯规→每个定位球都是奥地利的进球机会。',

        'upset_path': '如果翻车: 西班牙围攻无果→0-0到70分钟→奥地利角球→卡拉季奇2米头球→0-1→西班牙全线压上→奥地利大巴升级5-4-1→威廉姆斯不在→没有边路爆点突破封锁→0-1或1-1(加时/点球)。最可能冷门比分: 1-1 (奥地利加时或点球)——需要: 西班牙进攻转化率问题再次出现+亚马尔被双人夹防锁死+奥地利定位球效率100%。概率中等偏低——但佛得角0-0证明西班牙面对大巴时确实有困难。',
    },
    {
        'num': 2, 'time': '07:00', 'stage': 'R32',
        'home': '葡萄牙', 'away': '克罗地亚',
        'home_flag': 'POR', 'away_flag': 'CRO',
        'venue': 'BMO球场, 多伦多, 加拿大',
        'fifa_rank': '葡萄牙 #6 vs 克罗地亚 #8',
        'value_ratio': '葡萄牙 1.01B vs 克罗地亚 387M ≈ 2.6:1',
        'history': '葡萄牙6胜1平1负 (从未在世界杯交锋)',
        'reward': '16强对阵 西班牙 vs 奥地利 的胜者',

        'group_review_home': '葡萄牙 (K组第2, 5分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['刚果民主共和国', '1-1', '残血版本——C罗3射0正(6.7分)+B席5.8分+Dias缺阵。大巴奏效'],
            ['乌兹别克斯坦', '5-0', '火力全开——C罗帽子戏法+莱昂2球。但乌兹别克不摆大巴'],
            ['哥伦比亚', '0-0', '再次面对大巴哑火——C罗全场0射正，B费被锁死'],
        ],
        'group_review_away': '克罗地亚 (L组第2, 6分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['英格兰', '2-4', '被凯恩+贝林厄姆+拉什福德打穿，但进了2球→证明进攻可击穿强队防线'],
            ['巴拿马', '1-0', '稳扎稳打，莫德里奇控制节奏'],
            ['加纳', '2-1', '佩里西奇+布迪米尔进球→淘汰赛门票'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['布鲁诺·费尔南德斯(Bruno Fernandes) [火] MOTM', '8.5', '30岁曼联，9球21助(英超单赛季助攻纪录!)，30 G+A。葡萄牙真正进攻核心'],
            ['拉斐尔·莱昂(Rafael Leao)', '8.0', 'AC米兰进攻核心，对乌兹别克2球。速度+1v1是葡萄牙破大巴最大武器'],
            ['若昂·内维斯(Joao Neves)', '8.0', '21岁PSG 140M，+1.47/90分钟。PSG欧冠冠军→欧洲最佳中场双核之一'],
            ['维蒂尼亚(Vitinha)', '7.5', 'PSG中场，与内维斯组成技术型双后腰'],
            ['克里斯蒂亚诺·罗纳尔多(Cristiano Ronaldo) [C] (最低分)', '6.5', '41岁沙特联→不是超巨。沙特联28球但金靴第3+6点球。对大巴2场0射正。禁区终结者但不可单兵破局'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['卢卡·莫德里奇(Luka Modric) [C] [火]', '8.0', '40岁皇马中场，对加纳打满全场→工作量管理关键。克罗地亚节拍器'],
            ['约什科·格瓦迪奥尔(Josko Gvardiol)', '8.0', '70-81M曼城中卫→对巴拿马轮休后回归左后卫。世界级'],
            ['马特奥·科瓦契奇(Mateo Kovacic)', '7.5', '曼城中场，与莫德里奇组成技术型中场双核'],
            ['伊万·佩里西奇(Ivan Perisic)', '7.5', '对加纳进球。世界杯淘汰赛专家(2018决赛进球)'],
            ['安特·布迪米尔(Ante Budimir)', '7.3', '奥萨苏纳中锋，对加纳进球，支点作用'],
        ],

        'factors': [
            ('葡萄牙身价2.6:1(1.01B vs 387M)', '葡萄牙 [火][火]', '接近但<3:1。葡萄牙占优但不碾压——任何结果都可能'),
            ('葡萄牙对大巴0运动战进球(刚果+哥伦比亚)', '克罗地亚 [火][火][火]', '如果克罗地亚收缩防守→葡萄牙可能再次哑火——已验证的模式'),
            ('C罗41岁+沙特联赛+对大巴2场0射正', '克罗地亚 [火][火][火]', 'C罗已不是超巨。沙特联金靴第3+6点球+对大巴0射正=面对纪律防线无法单兵破局。葡萄牙9号位是最大短板'),
            ('克罗地亚淘汰赛血统(近3届均≥半决赛)', '克罗地亚 [火][火][火]', '世界杯淘汰赛4连胜+点球4/4全胜。葡萄牙没有这个级别的淘汰赛DNA'),
            ('莱昂+内托的速度 vs 克罗地亚防线速度', '葡萄牙 [火][火][火]', '克罗地亚对英格兰丢4球→英超速度打穿了他们。莱昂(AC米兰)同级别'),
            ('莫德里奇工作量(40岁，对加纳90分钟)', '葡萄牙 [火][火]', '如果比赛进加时→莫德里奇的体能是克罗地亚的定时炸弹'),
            ('迪亚斯回归(对刚果缺阵)', '葡萄牙 [火][火]', '葡萄牙防线升级→更不容易被克罗地亚反击打穿'),
            ('利瓦科维奇点球专家(2022世界杯英雄)', '克罗地亚 [火][火]', '如果进点球大战→克罗地亚是最大热门'),
        ],
        'strength': '葡萄牙: 超级巨星型(大巴前退化半档)——[修正] C罗(41岁沙特联+对大巴0射正)已不是超巨。但B费(90M+21助英超纪录)+莱昂(90M)+内维斯(140M)=3名超巨。破大巴取决于莱昂的速度——如果他被锁=葡萄牙倒退为体系型。C罗作为首发中锋对大巴2场0射正是核心矛盾',
        'resilience_label': '欧洲球队特性分析(克罗地亚)',
        'resilience': [
            ('低位防守', '良好', '对英格兰虽丢4球但前30分钟是稳的。对巴拿马零封'),
            ('速度反击', '良好', '佩里西奇+巴图里纳的反击速度不慢'),
            ('定位球高点', '良好', '布迪米尔(190cm)+格瓦迪奥尔+佩里西奇'),
            ('前30分钟', '良好', '对英格兰半场仅1-2落后→在强队面前不崩盘'),
            ('被压制不崩盘', '特优', '世界杯淘汰赛4连胜证明抗压能力。近三届至少进半决赛'),
        ],
        'upset_risk': '中',
        'upset_detail': '本日最可能爆冷的比赛。葡萄牙对大巴2场0运动战进球+克罗地亚淘汰赛4连胜+点球4/4全胜=克罗地亚赢这不叫冷门。莱昂的速度是克罗地亚防线的最大弱点——但葡萄牙需要早进球。如果0-0到60分钟→克罗地亚节奏控制→莫德里奇+科瓦契奇中场压制→佩里西奇反击1-0→C罗越踢越焦虑→葡萄牙全线压上但无宽度。',

        'predictions': [
            ('首选', '葡萄牙 2-1', '1-0', '莱昂速度爆破→上半场破门。克罗地亚不零封(对英格兰进2球)→定位球追平→下半场葡萄牙深度制胜'),
            ('备选', '葡萄牙 2-0', '1-0', '克罗地亚残血防线被零封→葡萄牙防守专注'),
            ('备选', '葡萄牙 3-1', '1-1', '对攻→克罗地亚追平→葡萄牙多点开花'),
            ('备选', '1-1 (葡萄牙加时晋级)', '0-0', '残血克罗地亚撑到70分钟→板凳差距→加时制胜。注意克罗地亚点球4/4'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['葡萄牙', '—', '全队健康', '罗伯托·马丁内斯有完整的26人名单'],
            ['葡萄牙', '鲁本·迪亚斯(Ruben Dias)', '可用', '已回归。对哥伦比亚首发→防线领袖'],
            ['葡萄牙', '若昂·内维斯(Joao Neves)', '可用', '回归首发。140M中场→控制力升级'],
            ['克罗地亚', '—', '全队健康', '兹拉特科·达利奇有完整的26人名单'],
            ['克罗地亚', '卢卡·莫德里奇(Luka Modric)', '可用', '40岁→对加纳90分钟→短间隔(5天)是挑战'],
            ['克罗地亚', '约什科·格瓦迪奥尔(Josko Gvardiol)', '可用', '对巴拿马轮休→满血回归左后卫'],
        ],

        'coach': [
            '葡萄牙(罗伯托·马丁内斯): 4-2-3-1/4-3-3莱昂+内托双边爆破+C罗禁区终结。关键: 内维斯回归首发→与维蒂尼亚组成PSG技术双后腰→中场控制力升级。领先: 帕利尼亚入场+5-4-1收缩。落后: 若昂·费利克斯+贡萨洛·拉莫斯+更多边路传中。',
            '克罗地亚(兹拉特科·达利奇): 世界杯淘汰赛专家。战术: 4-3-3莫德里奇控制节奏+巴图里纳创造+佩里西奇左路。领先: 极致控球消耗时间。落后: 佩特科维奇入场+双中锋+更多传中。',
            '关键博弈: 达利奇是世界杯淘汰赛最有经验的教练之一——他知道葡萄牙对大巴的困难。如果他选择收缩→葡萄牙可能面对刚果/哥伦比亚的翻版。如果他选择对攻(像对英格兰)→莱昂+内托的速度会杀死克罗地亚的慢防线。达利奇的初始策略决定这场比赛的性质。',
        ],

        'set_piece': '葡萄牙进攻: C罗头球(虽41岁仍是世界级)+迪亚斯跑位+B费任意球。莱昂+内托的角球接应。克罗地亚进攻: 布迪米尔(190cm)+格瓦迪奥尔+佩里西奇头球。莫德里奇+巴图里纳任意球精度。关键: 克罗地亚定位球是对强队的主要武器。葡萄牙防线在定位球防守上有纪律但迪亚斯是唯一的顶级空中威胁防守者。',

        'upset_path': '如果翻车: 克罗地亚收缩防守→葡萄牙对大巴再次哑火→0-0→克罗地亚节奏控制→莫德里奇/科瓦契奇中场压制→佩里西奇反击1-0→克罗地亚大巴升级→C罗越踢越焦虑→0-1或1-1(克罗地亚点球胜)。最可能冷门比分: 克罗地亚 1-0 或 1-1 (点球克罗地亚胜)——这是真实危险。但我们仍首选葡萄牙2-1——莱昂的速度是克罗地亚防线的最大弱点。',
    },
    {
        'num': 3, 'time': '11:00', 'stage': 'R32',
        'home': '瑞士', 'away': '阿尔及利亚',
        'home_flag': 'SUI', 'away_flag': 'ALG',
        'venue': 'BC Place, 温哥华, 加拿大',
        'fifa_rank': '瑞士 #15 vs 阿尔及利亚 ~#30',
        'value_ratio': '瑞士 332M vs 阿尔及利亚 257M ≈ 1.3:1',
        'history': '首次世界杯交锋',
        'reward': '16强对阵 哥伦比亚 vs 加纳 的胜者',

        'group_review_home': '瑞士 (B组第1, 7分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['卡塔尔', '1-1', '慢热——被东道主逼平'],
            ['波黑', '4-1', '瑞士进攻完全爆发——宽度+速度击穿波黑大巴'],
            ['加拿大', '2-1', '击败联合东道主→客场淘汰赛级别的考验通过'],
        ],
        'group_review_away': '阿尔及利亚 (J组第3, 4分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['约旦', '2-1', '开门红'],
            ['阿根廷', '0-3', '被阿根廷完全碾压'],
            ['奥地利', '3-3', '戏剧性——马赫雷斯2球但96分钟被绝平→淘汰赛门票险中求'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['约翰·曼扎姆比(Johan Manzambi) [火] MOTM', '8.5', '3球1助，20岁弗赖堡→本届世界杯突破之星。瑞士全部进攻的催化剂'],
            ['格拉尼特·扎卡(Granit Xhaka) [C]', '8.0', '勒沃库森，瑞士节拍器+更衣室领袖'],
            ['布雷尔·恩博洛(Breel Embolo)', '7.8', '摩纳哥前锋，支点中锋'],
            ['鲁本·巴尔加斯(Ruben Vargas)', '7.5', '奥格斯堡，边路速度'],
            ['曼努埃尔·阿坎吉(Manuel Akanji)', '7.8', '曼城中卫，瑞士防线基石'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['里亚德·马赫雷斯(Riyad Mahrez) [C] [火]', '8.3', '35岁，2球1助→可能是他最后一届世界杯。对奥地利2球证明大赛属性'],
            ['阿明·古伊里(Amine Gouiri)', '7.3', '雷恩前锋'],
            ['易卜拉欣·马扎(Ibrahim Maza)', '7.3', '20岁勒沃库森，52M，阿尔及利亚最有价值的球员'],
            ['希沙姆·布达维(Hicham Boudaoui)', '7.2', '尼斯中场'],
            ['纳比勒·本塔莱布(Nabil Bentaleb)', '7.5', '曾效力热刺，经验中场'],
        ],

        'factors': [
            ('身价比1.3:1(<3:1接近)', '持平', '不是压倒性差距——任何结果都可能'),
            ('瑞士9场不败(含击败加拿大客场)', '瑞士 [火][火]', '状态+信心在最高点'),
            ('阿尔及利亚防线7失球(本届最漏之一)', '瑞士 [火][火][火]', '对阿根廷丢3球→瑞士4-1波黑证明能惩罚脆弱防线'),
            ('马赫雷斯2球1助(大赛属性)', '阿尔及利亚 [火][火][火]', '35岁但仍是阿尔及利亚唯一破局者→定位球+个人能力'),
            ('阿穆拉缺阵(大腿伤)', '瑞士 [火][火]', '阿尔及利亚最快前锋缺失→反击维度减少'),
            ('佩特科维奇=瑞士前主帅(2014-2021)', '阿尔及利亚 [火][火][火]', '知道瑞士的一切: 扎卡的习惯、阿坎吉的弱点、防线结构'),
            ('曼扎姆比3球1助(爆红状态)', '瑞士 [火][火][火]', '20岁弗赖堡→本届突破之星→阿尔及利亚需专门计划限制他'),
            ('阿坎吉+科贝尔=曼城+多特防线', '瑞士 [火][火]', '英超/德甲验证的顶级防守组合'),
        ],
        'strength': '瑞士: 体系型(状态加成版)——332M < 500M，0超巨。曼扎姆比(50-64M)最接近核心级。但9场不败+击败加拿大的客场考验→比纸面强半个级别  |  阿尔及利亚: 马赫雷斯依赖型+防线漏洞——7失球=淘汰赛最差。但佩特科维奇知道瑞士一切+马赫雷斯可创造奇迹',
        'resilience_label': '非洲韧性评估(阿尔及利亚)',
        'resilience': [
            ('低位防守', '差', '7球失球=非洲淘汰赛球队中最差。对阿根廷0-3暴露防线结构问题'),
            ('速度反击', '良好', '马赫雷斯+古伊里的反击速度，但阿穆拉缺阵降低半档'),
            ('定位球高点', '一般', '没有真正的空中巨人。马赫雷斯的任意球精确度是唯一亮点'),
            ('前30分钟', '差', '对阿根廷早期丢球→可能重演'),
            ('被压制不崩盘', '良好', '对阿根廷0-3未崩成1-5，有韧性但不完整'),
        ],
        'upset_risk': '中高',
        'upset_detail': '本日身价比最接近的比赛(1.3:1)——任何结果都可能。佩特科维奇执教瑞士7年(2014-2021)=知道扎卡在压力下的倾向、阿坎吉的防守习惯、瑞士中场的战术触发点。这是教练情报的不对称优势——相当于对方手里有你的战术手册。但是: 阿尔及利亚防线7失球+阿穆拉缺阵+瑞士9场不败+曼扎姆比爆红状态=瑞士仍是更好的一方。只是更好不等于一定赢——本日最可能爆冷的比赛。',

        'predictions': [
            ('首选', '瑞士 2-1', '1-0', '曼扎姆比早破门→马赫雷斯个人能力追平(对奥地利2球)→瑞士整体实力制胜'),
            ('备选', '瑞士 1-0', '0-0', '佩特科维奇精准战术→极度胶着→曼扎姆比最后一刻破局'),
            ('备选', '瑞士 2-0', '1-0', '阿尔及利亚防线崩溃+马赫雷斯被锁→零封(概率低——对卡塔尔都丢球)'),
            ('备选', '1-1 (瑞士加时晋级)', '0-0', '佩特科维奇战术奏效→阿尔及利亚大巴120分钟→板凳制胜'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['瑞士', '西尔万·维德默(Silvan Widmer)', '疑', '臀部不适→缺席周日训练。右后卫→如缺阵需客串'],
            ['瑞士', '米罗·穆海姆(Miro Muheim)', '疑', '小伤→非严重。左后卫轮换'],
            ['阿尔及利亚', '穆罕默德·阿穆拉(Mohamed Amoura)', '很可能缺', '大腿拉伤→缺席最后2场。最快前锋缺失→反击维度减少'],
            ['阿尔及利亚', '纳迪尔·本布阿利(Nadir Benbouali)', '可用', '内收肌恢复→进攻替补选项回归'],
            ['阿尔及利亚', '里亚德·马赫雷斯(Riyad Mahrez)', '可用', '工作量管理→35岁短间隔挑战但首发无疑'],
        ],

        'coach': [
            '瑞士(穆拉特·雅金): 4-2-3-1扎卡+弗罗伊勒双后腰控制中场+曼扎姆比创造力+恩博洛支点+巴尔加斯边路速度。领先: 5-4-1收缩+阿坎吉+埃尔韦迪防空。落后: 更多边路传中+恩博洛+阿姆杜尼双中锋。',
            '阿尔及利亚(弗拉基米尔·佩特科维奇): [火] 核心变量。执教瑞士7年→知道扎卡在压力下的倾向→知道阿坎吉面对速度型前锋时的弱点→知道瑞士4-2-3-1的战术触发点。战术: 预计4-2-3-1收缩防守+马赫雷斯右路反击+古伊里/马扎的跑位。领先: 5-4-1极致收缩。落后: 变阵4-4-2→上第二前锋。',
            '关键博弈: 佩特科维奇vs他的前球队和继任者。他是否知道如何解锁瑞士? 他的战术手册是否比雅金的更深? 这是教练博弈的独特案例——7年的知识不对称。',
        ],

        'set_piece': '瑞士进攻: 阿坎吉+埃尔韦迪头球，扎卡远射+任意球。对波黑4-1包含定位球得分。阿尔及利亚进攻: 马赫雷斯任意球精确度(世界级)+古伊里速度追第二落点。马赫雷斯的左脚任意球是阿尔及利亚最危险的定位球武器。关键: 阿尔及利亚7失球中有部分是定位球防守问题。瑞士4-1波黑证明了能利用定位球惩罚弱防守。',

        'upset_path': '如果翻车: 佩特科维奇精准战术→阿尔及利亚5-4-1封死曼扎姆比→马赫雷斯任意球世界波→0-1→瑞士焦虑→扎卡因身体对抗分散注意力→瑞士进攻变混乱→阿尔及利亚反击+马赫雷斯个人能力锁定。或1-1→加时→阿尔及利亚点球(马赫雷斯是例外)。最可能冷门比分: 阿尔及利亚 1-0 或 2-1 (马赫雷斯致胜)——佩特科维奇因素是真实的外卡。',
    },
]

# ── Render each match ──
for m in matches:
    add_heading_styled(f'比赛{m["num"]}: {m["home_flag"]} {m["home"]} vs {m["away_flag"]} {m["away"]}  ({m["stage"]} {m["time"]} BJT)', level=2)

    # ── Basic info ──
    add_para('基本信息', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))
    info_lines = [
        f'场地: {m["venue"]}',
        f'FIFA排名: {m["fifa_rank"]}  |  身价比: {m["value_ratio"]}',
        f'历史交锋: {m["history"]}  |  晋级奖励: {m["reward"]}',
    ]
    for line in info_lines:
        add_para(line, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # ── Group stage review ──
    add_para('小组赛回顾', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    # Home team review
    add_para(m['group_review_home'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_home_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_home_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_home_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    # Away team review
    add_para(m['group_review_away'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_away_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_away_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_away_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    # ── Player ratings ──
    add_para('首轮球员评分', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    add_para(f'{m["home"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_home']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_home'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_home'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    add_para(f'{m["away"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_away']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_away'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_away'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    # ── Factors table ──
    add_para('因素导向表', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    ft_data = [['因素', '有利方', '理由']]
    ft_data.extend(m['factors'])
    ft = doc.add_table(rows=len(ft_data), cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(ft_data[0]):
        set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(8.5), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(ft_data[1:], 1):
        set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8))
        clr = None
        if '[火][火][火]' in row[1] or '[火][火]' in row[1]:
            clr = ACCENT
        set_cell(ft.rows[i].cells[1], row[1], size=Pt(8), color=clr)
        set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in ft.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(5.0)
    doc.add_paragraph()

    # ── Strength + Resilience ──
    add_para(f'强队分类: {m["strength"]}', size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))

    add_para(m['resilience_label'], bold=True, size=Pt(9), color=DARK, space_after=Pt(2))
    res_data = [['维度', '评分', '说明']]
    res_data.extend(m['resilience'])
    rt = doc.add_table(rows=len(res_data), cols=3)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(res_data[0]):
        set_cell(rt.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(res_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(rt.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(rt.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(rt.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in rt.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(7.5)
    doc.add_paragraph()

    # ── Upset risk ──
    risk_color = ACCENT if m['upset_risk'] in ('中', '中高', '高') else DARK
    add_para(f'冷门风险: {m["upset_risk"]}', bold=True, size=Pt(10), color=risk_color, space_after=Pt(2))
    add_para(m['upset_detail'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # ── Predictions table ──
    add_para('比分预测', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    pred_header = ['类型', '比分', '半场', '说明']
    pred_rows = []
    for p in m['predictions']:
        pred_rows.append(list(p))
    pt = doc.add_table(rows=1 + len(pred_rows), cols=4)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(pred_header):
        set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(pred_rows, 1):
        bg = GREEN_BG if i == 1 else None
        set_cell(pt.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg)
        set_cell(pt.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
        set_cell(pt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
        set_cell_left(pt.rows[i].cells[3], row[3], size=Pt(8))
    for row in pt.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(7.4)
    doc.add_paragraph()

    # ── Injuries ──
    add_para('伤病/停赛', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    inj_data = m['injuries']
    it = doc.add_table(rows=len(inj_data), cols=4)
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(inj_data[0]):
        set_cell(it.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(inj_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(it.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[1], row[1], size=Pt(7.5))
        set_cell(it.rows[i].cells[2], row[2], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[3], row[3], size=Pt(7.5))
    for row in it.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(5.2)
    doc.add_paragraph()

    # ── Coach tactics ──
    add_para('教练博弈', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    for c in m['coach']:
        add_para(c, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # ── Set piece ──
    add_para('定位球攻防', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    add_para(m['set_piece'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))
    doc.add_paragraph()

    # ── Upset path ──
    add_para('冷门路径', bold=True, size=Pt(10), color=ACCENT, space_after=Pt(2))
    add_para(m['upset_path'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # separator between matches
    add_para('─' * 60, size=Pt(6), color=RGBColor(0xCC,0xCC,0xCC), align='center', space_after=Pt(8))

# ═══════════════════════════════════════════════════
# RISK SUMMARY
# ═══════════════════════════════════════════════════
add_heading_styled('风险提示', level=2)

risk_notes = [
    ('西班牙 vs 奥地利', '西班牙0失球防线vs奥地利12场世界杯0零封=基本矛盾决定方向。但威廉姆斯的缺阵让进攻维度减少→如果0-0到60分钟→警惕佛得角重演。卡拉季奇2米头球是奥地利的唯一真实威胁。'),
    ('葡萄牙 vs 克罗地亚 [注意] 本日最可能爆冷', '葡萄牙对大巴2场0运动战进球+克罗地亚淘汰赛4连胜+点球4/4=克罗地亚赢这不叫冷门。莱昂的速度是葡萄牙最好的武器——如果他被限制，葡萄牙可能再次哑火。C罗的世界杯淘汰赛0进球诅咒是额外的心理因素。'),
    ('瑞士 vs 阿尔及利亚 [注意] 本日身价比最接近', '身价比最接近(1.3:1)。佩特科维奇知道瑞士的一切=教练情报不对称。但阿尔及利亚防线7失球+阿穆拉缺阵+曼扎姆比爆红状态=瑞士仍是更好的一方。这是那种瑞士应该赢但没有人会惊讶如果阿尔及利亚赢的比赛。'),
]

for title, detail in risk_notes:
    add_para(title, bold=True, size=Pt(9), color=ACCENT, space_after=Pt(1))
    add_para(detail, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: ESPN API + Rotowire + Sports Mole + Yahoo Sports + Transfermarkt + Xinhua News  |  框架: CLAUDE.md v16 (身价量化+强队三类+10项必填清单+7月2日复盘教训)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('预测时间: 2026年7月2日 14:00 北京时间  |  教训加权: 弱队先破门→强队逆转+非洲75-85分钟防守崩塌+首发确认后不过度修正+防线降级≠冷门风险降低',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('[注意] 首发确认将在开球前约1小时更新  |  预测仅供分析参考  |  生成工具: python-docx',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

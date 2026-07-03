# -*- coding: utf-8 -*-
"""Generate 2026-07-04 World Cup predictions as a formatted .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月4日_3场预测.docx")

doc = Document()

# ── v6 colour palette ──
HEADER_BG  = RGBColor(0x1A, 0x2E, 0x3D)   # deep navy
HEADER_TEXT= RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW    = RGBColor(0xF5, 0xF7, 0xFA)
PREF_ROW   = RGBColor(0xE8, 0xF5, 0xE9)   # green tint
BORDER_CLR = RGBColor(0xD0, 0xD5, 0xDD)
TITLE_RED  = RGBColor(0xC0, 0x39, 0x2B)
SUBTITLE   = RGBColor(0x1A, 0x1A, 0x2E)
META       = RGBColor(0x7F, 0x8C, 0x8D)
META_LIGHT = RGBColor(0x95, 0xA5, 0xA6)
ACCENT_BLUE= RGBColor(0x2E, 0x75, 0xB6)
HIGH_RISK  = RGBColor(0xC0, 0x39, 0x2B)
GREEN_OK   = RGBColor(0x27, 0xAE, 0x60)

# legacy aliases
DARK = SUBTITLE; ACCENT = TITLE_RED; WHITE = HEADER_TEXT
LIGHT_GRAY = ALT_ROW; GREEN_BG = PREF_ROW; GRAY_TEXT = META

# Page setup — landscape A4 (correct v6 approach)
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

def add_para_border(paragraph, color='2E75B6', sz='6'):
    """Add a bottom border to paragraph with proper pPr element ordering.
    Schema order: pStyle, numPr, pBdr, spacing, ind, jc, rPr"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    # Insert before spacing, or before ind, or before jc, or before rPr
    for tag in ['w:spacing', 'w:ind', 'w:jc', 'w:rPr']:
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.insert(list(pPr).index(el), pBdr)
            return
    pPr.append(pBdr)

# ── Header ──
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run('2026 FIFA 世界杯 · 7月4日 三场预测报告')
    run.font.size = Pt(7.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = META_LIGHT
    add_para_border(hp, sz='4')

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    # page number
    run_pg = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_pg._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run_pg._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg._r.append(fldChar2)
    run_pg.font.size = Pt(7.5)
    run_pg.font.name = '微软雅黑'
    run_pg._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_pg.font.color.rgb = META_LIGHT

# Styles
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# helper functions
def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# TITLE PAGE
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('7月4日 三场预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('淘汰赛 Round of 32  |  AUS vs EGY / ARG vs CPV / COL vs GHA', bold=False, size=Pt(11), color=GRAY_TEXT, align='center', space_after=Pt(18))

add_para('生成时间: 2026年7月3日 14:00 北京时间  |  数据源: ESPN API + Sports Mole + BBC Sport + ABC News + FIFA  |  框架: CLAUDE.md v16 + 7月3日教训',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# SUMMARY TABLE
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '时间', '阶段', '比赛', '身价比', '首选', '概率', '备选(按概率)', '冷门'],
    ['1', '02:00', 'R32', 'AUS 澳大利亚 vs EGY 埃及', '~1:1.5', '埃及 2-1', '~35%', '1-1(加时) / 埃及 1-0 / 澳 2-0', '中'],
    ['2', '06:00', 'R32', 'ARG 阿根廷 vs CPV 佛得角', '~16:1', '阿根廷 2-0', '~45%', '1-0 / 0-0(加时) / 佛得角点球胜', '低'],
    ['3', '09:30', 'R32', 'COL 哥伦比亚 vs GHA 加纳', '~2:1', '哥伦比亚 1-0', '~38%', '加纳 1-0(~35%) / 2-0 / 1-1(加时)', '中高'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 0.50, 0.42, 2.55, 0.70, 0.66, 0.44, 2.10, 0.55]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        align = 'left' if j == 3 else 'center'
        bold = (j == 5)
        if j == 3:
            set_cell_left(table.rows[i].cells[j], text, bold=False, size=Pt(8.5))
        else:
            set_cell(table.rows[i].cells[j], text, bold=bold, size=Pt(8.5), bg=bg)

doc.add_paragraph()

# KNOCKOUT BRACKET PATHS
add_heading_styled('淘汰赛路径（7月3日收束后更新）', level=2)

bracket_lines = [
    '已确定:',
    '  西班牙 (3-0奥地利) -> R16 -> 葡萄牙 (2-1克罗地亚)',
    '  瑞士 (2-0阿尔及利亚) -> R16 -> 哥伦比亚/加纳 胜者',
    '',
    '7月4日路径:',
    '  澳大利亚/埃及 -> R16 -> 阿根廷/佛得角 胜者',
    '  阿根廷/佛得角 -> R16 -> 澳大利亚/埃及 胜者',
    '  哥伦比亚/加纳 -> R16 -> 瑞士 (已确定)',
]
for line in bracket_lines:
    add_para(line, size=Pt(8.5), color=DARK if line.startswith('7月') or line.startswith('已') else RGBColor(0x55,0x55,0x55), space_after=Pt(2))

# Path difficulty table
add_para('路径难度分析', bold=True, size=Pt(10), color=DARK, space_after=Pt(4))
path_data = [
    ['路径', '难度', '说明'],
    ['阿根廷方向 -> 澳大利亚或埃及', '低', '阿根廷是R32最大热门——无论对手是谁，身价比>10:1'],
    ['哥伦比亚方向 -> 瑞士(R16) -> 可能阿根廷(QF)', '高', '若晋级将碰瑞士(已验证体系防守)->QF可能碰阿根廷'],
    ['加纳方向 -> 瑞士(R16)', '中', '加纳已验证大巴(0-0英格兰)——但瑞士刚展示了如何锁死非洲对手'],
]
pt = doc.add_table(rows=len(path_data), cols=3)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(path_data[0]):
    set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(path_data[1:], 1):
    set_cell_left(pt.rows[i].cells[0], row[0], size=Pt(8.5))
    set_cell(pt.rows[i].cells[1], row[1], size=Pt(8.5))
    set_cell_left(pt.rows[i].cells[2], row[2], size=Pt(8.5))
for row in pt.rows:
    row.cells[0].width = Inches(4.0)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(5.5)

doc.add_paragraph()

# ============================================================
# MATCH DATA
# ============================================================

matches = [
    {
        'num': 1, 'time': '02:00', 'stage': 'R32',
        'home': '澳大利亚', 'away': '埃及',
        'home_flag': 'AUS', 'away_flag': 'EGY',
        'venue': 'AT&T体育场, 阿灵顿(达拉斯), 美国',
        'fifa_rank': '澳大利亚 #26 vs 埃及 #30',
        'value_ratio': '澳大利亚 ~60M vs 埃及 ~90M (埃及略高)',
        'history': '仅1次交锋——2010年友谊赛埃及3-0澳大利亚',
        'reward': '16强对阵 阿根廷 vs 佛得角 的胜者',

        'group_review_home': '澳大利亚 (D组第2, 4分)',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['土耳其', '2-0', '伊兰昆达进球——20岁拜仁前锋是澳大利亚唯一破局者'],
            ['美国', '0-2', '被东道主完全压制——0射正'],
            ['巴拉圭', '0-0', '死守拿1分→晋级'],
        ],
        'group_review_away': '埃及 (G组第2, 5分)',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['比利时', '1-1', '逼平身价5倍对手——萨拉赫送助攻'],
            ['新西兰', '3-1', '萨拉赫进球——埃及进攻火力全开'],
            ['伊朗', '1-1', '萨拉赫57\'伤退(腿筋)→随后被追平'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['内斯托里·伊兰昆达(Irankunda) [火]', '7.5', '20岁拜仁前锋→6球/18场国家队。本届已进1球。速度是澳大利亚唯一反击武器'],
            ['克里斯蒂安·沃尔帕托(Volpato)', '7.0', '前罗马/萨索洛→创造型中场。澳大利亚技术核心'],
            ['哈里·苏塔尔(Harry Souttar)', '7.3', '莱斯特城1.98m中卫→防空炮塔+定位球进攻支点'],
            ['杰克逊·欧文(Jackson Irvine) [C]', '7.0', '经验中场'],
            ['乔丹·博斯(Jordan Bos)', '7.0', '左翼卫→速度型'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['穆罕默德·萨拉赫(Mohamed Salah) [C] [火]', '8.0', '[注意] 腿筋伤→赛前决定! 本届1球2助=埃及75%的进攻参与。如果缺阵→埃及降半档'],
            ['奥马尔·马穆什(Omar Marmoush)', '7.0', '曼城前锋但本届0G0A。英超8首发/21场→大部分替补。需证明自己'],
            ['马哈茂德·特雷泽盖(Trezeguet)', '7.3', '前维拉边锋→突破+传中'],
            ['齐佐(Zizo)', '7.3', '扎马雷克→右边锋'],
            ['穆斯塔法·舒贝尔(Shobeir)', '7.0', '阿赫利门将'],
        ],

        'factors': [
            ('萨拉赫腿筋伤(赛前决定/非100%)', '澳大利亚 [火][火][火]', '本届1球2助=埃及75%进攻参与。即使勉强出场也不是100%状态'),
            ('埃及防线多人伤疑(3-4人+停赛)', '澳大利亚 [火][火][火]', '阿卜杜勒蒙内姆+法图+法特希+拉欣停赛->防线被迫重组'),
            ('澳大利亚对美/巴180分钟0进球', '埃及 [火][火][火]', '进攻效率R32最差之一。伊兰昆达(20岁)是唯一速度点'),
            ('埃及小组赛不败(5分) vs 澳4分', '埃及 [火][火]', '状态验证——埃及从未输球'),
            ('萨拉赫即使替补也是后手', '埃及 [火][火][火]', '如果60-70分钟0-0->萨拉赫上场15-20分钟=比赛最大变量'),
            ('澳大利亚定位球高度(苏塔尔1.98m)', '澳大利亚 [火][火]', '面对埃及残血防线->角球/任意球是澳最可能进球的路径'),
            ('伊兰昆达(拜仁20岁)速度', '澳大利亚 [火][火]', '唯一可以单兵破防的球员'),
            ('马穆什(曼城)本届0G0A', '澳大利亚 [火][火]', '若萨拉赫缺阵->埃及进攻核心是本届尚未破荒的球员'),
        ],
        'strength': '澳大利亚: 体系型(进攻受限版)——~60M < 500M。0超巨。对美/巴180分钟0进球。定位球(苏塔尔1.98m)是唯一真实得分路径  |  埃及: 萨拉赫依赖型(残血版)——~90M。萨拉赫(34岁仍是超巨)但腿筋伤。非洲杯冠军+小组赛不败=体系>个人。但防线多人伤疑',
        'resilience_label': '非洲韧性评估(埃及)',
        'resilience': [
            ('低位防守', '良好', '对比利时1-1->面对5倍身价守住。非洲杯冠军防守体系'),
            ('速度反击', '良好', '萨拉赫(如果健康)=世界级反击速度。马穆什本届0G0A'),
            ('定位球高点', '一般', '阿卜杜勒蒙内姆(疑)1.85m->若缺阵防空降级'),
            ('前30分钟', '良好', '小组赛3场前30分钟从未落后'),
            ('被压制不崩盘', '良好', '对伊朗1-0领先到57\'(萨拉赫伤退)->被追平但未崩'),
        ],
        'upset_risk': '中',
        'upset_detail': '双方都不是热门——身价比仅1:1.5。真正的变量只有一个: 萨拉赫。如果萨拉赫100%首发->埃及明显占优。如果萨拉赫替补(最可能)->前60分钟澳大利亚有窗口期(定位球+伊兰昆达速度)。如果萨拉赫缺阵->澳大利亚可能2-0。预测基于"萨拉赫替补60分钟后上场"的假设。',

        'predictions': [
            ('埃及 2-1', '~35%', '[首选] 假设萨拉赫替补60分钟上场。上半场埃及(马穆什+特雷泽盖)仍能破门->0-1。澳下半场定位球追平(苏塔尔1.98m)->萨拉赫上场后30分钟制胜'),
            ('1-1 (埃及加时晋级)', '~20%', '萨拉赫替补->前60分钟双方无法破门->萨拉赫上场后加时突破'),
            ('埃及 1-0', '~20%', '萨拉赫首发但非100%->一球小胜->埃及残血防线守住零封'),
            ('澳大利亚 2-0', '~15%', '萨拉赫完全缺阵时->埃及残血防线+无核心=澳定位球(苏塔尔)+伊兰昆达反击'),
            ('澳大利亚 1-0', '~5%', '萨拉赫缺阵+澳防守极致收缩->定位球偷一个->守住'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['澳大利亚', '马修·莱基(Leckie)', '缺阵', '腿筋伤。主力边锋缺失->反击深度减少'],
            ['澳大利亚', '雅各布·伊塔利亚诺(Italiano)', '缺阵', '腹股沟。轮换球员'],
            ['埃及', '萨拉赫(Mohamed Salah) [注意]', '疑/赛前决定', '腿筋。1球2助=埃及进攻75%。即使出场也不是100%'],
            ['埃及', '阿卜杜勒蒙内姆(Abdelmonem)', '疑', '膝伤->对伊朗14\'伤退。主力中卫->残血防线核心缺失'],
            ['埃及', '艾哈迈德·法图(El Fotouh)', '疑/大概率缺', '腿筋。主力左后卫'],
            ['埃及', '莫哈纳德·拉欣(Lasheen)', '停赛', '累积黄牌。中场防守缺失'],
        ],

        'coach': [
            '澳大利亚(托尼·波波维奇): 3-4-2-1/5-4-1收缩+伊兰昆达单前锋反击+沃尔帕托创造力+定位球(苏塔尔)。领先: 5-4-1极致收缩。核心博弈: 前60分钟(萨拉赫上场前)建立领先。',
            '埃及(侯萨姆·哈桑): [注意] 核心变量: 萨拉赫的使用方式。"不冒任何风险"+"100%健康才首发"->暗示替补。领先: 4-5-1收缩+特雷泽盖反击。落后: 提前上萨拉赫。',
            '关键博弈: 哈桑的"萨拉赫使用时间"是比赛最大战术决策。如果萨拉赫首发但半场状态不佳->浪费换人窗口。如果替补但埃及60分钟已0-1落后->逆转压力更大。',
        ],

        'set_piece': '澳大利亚进攻: [注意] 苏塔尔(1.98m莱斯特城)+赫灵顿=R32最高中卫组合之一。面对埃及残血防线->角球/任意球是澳最可能进球的路径。埃及进攻: 萨拉赫任意球精度(世界级)——即使替补也可利用。关键: 澳应该能利用定位球进1球(埃及防线残血)。埃及定位球取决于萨拉赫是否在场。',

        'upset_path': '如果翻车(澳大利亚赢): 萨拉赫完全缺阵->埃及进攻哑火(马穆什0G0A无法担当核心)->澳定位球(苏塔尔)上半场破门->1-0->波波维奇5-4-1收缩->伊兰昆达反击锁定2-0。最可能冷门比分: 澳大利亚 2-0 (仅当萨拉赫完全缺阵)。',
    },
    {
        'num': 2, 'time': '06:00', 'stage': 'R32',
        'home': '阿根廷', 'away': '佛得角',
        'home_flag': 'ARG', 'away_flag': 'CPV',
        'venue': 'Hard Rock体育场, 迈阿密, 美国',
        'fifa_rank': '阿根廷 #1 vs 佛得角 #63',
        'value_ratio': '阿根廷 816M vs 佛得角 ~50M (约16:1)',
        'history': '首次交锋',
        'reward': '16强对阵 澳大利亚 vs 埃及 的胜者',

        'group_review_home': '阿根廷 (C组第1, 9分) — 唯一全胜球队之一',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['约旦', '3-1', '梅西替补->但阿根廷仍轻松取胜'],
            ['奥地利', '3-0', '梅西2球->完全碾压'],
            ['伊朗', '2-0', '再次零封->3场仅失1球'],
        ],
        'group_review_away': '佛得角 (H组第3, 3分) — 世界杯首秀即进淘汰赛',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['西班牙', '0-0', '[注意] 已验证大巴! 面对1.22B西班牙74%控球->佛得角守住了。沃齐尼亚英雄'],
            ['乌拉圭', '2-2', '1-0领先->2-1落后->2-2追平->面对南美强队不崩盘'],
            ['沙特阿拉伯', '0-0', '第二次零封->大巴验证x2'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['利昂内尔·梅西(Lionel Messi) [C] [火] MOTM', '9.5', '39岁迈阿密国际。本届6球=75%的阿根廷进球。10连胜。7年淘汰赛不败'],
            ['劳塔罗·马丁内斯(Lautaro Martinez)', '8.5', '国际米兰->38国家队进球。首发前锋确认'],
            ['恩佐·费尔南德斯(Enzo Fernandez)', '8.0', '切尔西->B2B中场->后排插上+破大巴关键'],
            ['罗德里戈·德保罗(Rodrigo De Paul)', '7.8', '马竞->梅西的保镖+中场发动机'],
            ['亚历克西斯·麦卡利斯特(Mac Allister)', '8.0', '利物浦->技术型中场->小组赛轮休后满血回归'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['沃齐尼亚(Vozinha) [火]', '8.5', '[注意] 本届最佳门将之一! 对西班牙0-0=全场最佳。面对梅西(6球)=地球最难门将挑战'],
            ['瑞安·门德斯(Ryan Mendes) [C]', '7.5', '22国家队进球->佛得角历史最佳射手。反击速度+经验'],
            ['迪内伊·博尔格斯(Diney Borges)', '7.3', '中卫+领袖->对西班牙0-0的防守组织核心'],
            ['加里·罗德里格斯(Garry Rodrigues)', '7.0', '10国家队进球->35岁->经验边锋'],
            ['贾米罗·蒙泰罗(Jamiro Monteiro)', '7.0', '前费城联->MLS经验(了解阿根廷球员风格)'],
        ],

        'factors': [
            ('身价比16:1(816M vs ~50M)', '阿根廷 [火][火][火]', '>10:1=大差距。阿根廷正常发挥即胜。不是"任何结果都可能"的比赛'),
            ('梅西6球(75%的阿根廷进球)', '阿根廷 [火][火][火]', '39岁可决定任何比赛。佛得角无法用一人限制他->即使大巴也需2-3人'),
            ('佛得角已验证大巴(0-0西班牙!)', '佛得角 [火][火][火]', '[注意] 面对1.22B西班牙74%控球->0-0。这是"已验证大巴"不是"被迫大巴"'),
            ('阿根廷7年淘汰赛不败', '阿根廷 [火][火][火]', '2019年以来未在淘汰赛输球。梅西的时代->最大比赛中从不掉链子'),
            ('佛得角终结能力(R32最低2球)', '阿根廷 [火][火][火]', '[7/3教训] 漏洞需对手有终结能力。佛得角进攻无法惩罚阿根廷防守瑕疵'),
            ('西班牙0-0佛得角=已验证大巴(阿根廷0-0)', '持平', '阿根廷(有)梅西。梅西在破大巴方面比任何现役球员都有经验'),
            ('沃齐尼亚(对西班牙全场最佳)', '佛得角 [火][火]', '门将是佛得角大巴最关键的位置。但对梅西的禁区外射门->门将再强也难挡'),
            ('迈阿密=梅西主场(国际迈阿密)', '阿根廷 [火]', 'Hard Rock Stadium=梅西在迈阿密国际的主场城市。气候/语言=舒适区'),
        ],
        'strength': '阿根廷: 超级巨星型(Messi版)——816M >> 500M。>=2名超巨: 梅西(39岁仍是超巨级->6球)。2019年以来淘汰赛不败=历史上最擅长淘汰赛的球队。10连胜=热得发烫  |  佛得角: 大巴专家型(已验证)——~50M << 500M。对西班牙0-0=这不是一般的大巴。沃齐尼亚+迪内伊轴线=可与任何球队防线抗衡。但进攻2球=无法惩罚阿根廷',
        'resilience_label': '非洲韧性评估(佛得角)',
        'resilience': [
            ('低位防守', '世界级', '[注意] R32最佳低位防守。对西班牙0-0(74%控球->0进球)=本届防守教科书'),
            ('速度反击', '一般', '2球=反击速度存在(门德斯+罗德里格斯)但终结不足(利夫拉门托0球)'),
            ('定位球高点', '一般', '迪内伊1.87m->头球接应。但进攻定位球不是佛得角主要武器'),
            ('前30分钟', '良好', '3场小组赛前30分钟从未落后'),
            ('被压制不崩盘', '世界级', '对西班牙被压制74%控球->0失球。对乌拉圭1-0领先->2-1落后->2-2追平'),
        ],
        'upset_risk': '低',
        'upset_detail': '[注意] 冷门风险"低"不是因为佛得角不危险——而是即使佛得角大巴有效，最可能结果仍是0-0加时(阿根廷晋级)而非佛得角赢。佛得角0-0西班牙(有)佛得角0-0阿根廷——西班牙没有梅西(6球)，梅西的破大巴能力比西班牙的佩德里+奥尔莫更强。半场0-0很可能但90分钟0-0概率~25%。身价比16:1->不做0-0首选。',

        'predictions': [
            ('阿根廷 2-0', '~45%', '[首选] 佛得角大巴上半场守住(0-0)->70-80分钟梅西个人能力破局->1-0->佛得角被迫打开->劳塔罗/阿尔瓦雷斯反击锁定2-0'),
            ('阿根廷 1-0', '~25%', '佛得角大巴几乎守住120分钟->梅西任意球/个人突破最后15分钟破门'),
            ('0-0 (阿根廷加时/点球晋级)', '~15%', '佛得角将0-0带入加时->但120分钟不丢球vs梅西=概率极低->加时或点球阿根廷仍大概率晋级'),
            ('阿根廷 3-0', '~8%', '梅西上半场即破门(20分钟内)->佛得角被迫放弃大巴->多点开花'),
            ('佛得角 1-0 或点球胜', '~7%', '对西班牙0-0完整复制->沃齐尼亚世界级->佛得角定位球/反击破门->阿根廷被大巴封死。概率最低但非零'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['阿根廷', '贡萨洛·蒙铁尔(Montiel)', '缺阵', '腿筋。右后卫轮换->莫利纳首发不受影响'],
            ['阿根廷', '克里斯蒂安·罗梅罗(Romero)', '可用', '膝伤康复->回归训练。防线领导力升级'],
            ['阿根廷', '梅西+麦卡利斯特+德保罗', '满血', '小组赛末轮轮休->核心球员得到了休息'],
            ['佛得角', '特奥姆·阿尔坎若(Arcanjo)', '疑/很可能缺', '肌肉伤。中场创造者缺失->反击出球点减少'],
            ['佛得角', '西德尼·卡布拉尔(Sidny Cabral)', '可用', '停赛复出->首发。左后卫回归->防线完整'],
        ],

        'coach': [
            '阿根廷(利昂内尔·斯卡洛尼): 7年淘汰赛不败的教练。4-4-2梅西+劳塔罗双前锋->梅西自由游走找空隙(破大巴关键)。斯卡洛尼面对大巴的经验比任何现役教练都多。领先: 5-4-1收缩+阿尔瓦雷斯反击。',
            '佛得角(布比斯塔): [注意] 大巴专家。5-4-1极致收缩->对西班牙的成功战术无需改变->封锁中路->迫使走外线。面对梅西(6球)面对的不是"另一个西班牙"——是"西班牙+梅西"。',
            '关键博弈: 布比斯塔的成功模式能否对梅西复制? 西班牙没有梅西->佛得角可以封锁佩德里(12助)但无法封锁梅西(6球)。西班牙破大巴受限于体系->梅西破大巴是个人能力。',
        ],

        'set_piece': '阿根廷进攻: 梅西任意球(世界级)+罗梅罗(头球)+利桑德罗头球。梅西左脚任意球是面对大巴时最致命的破局武器。佛得角进攻: 门德斯+罗德里格斯角球/任意球->迪内伊(1.87m)头球。对乌拉圭2-2进球来自定位球混乱。关键: 阿根廷定位球进攻(梅西任意球)是世界级->佛得角大巴的短板(禁区外)。',

        'upset_path': '如果翻车(佛得角赢): 佛得角将0-0带入加时->沃齐尼亚扑出梅西任意球+多次世界级扑救->点球大战->佛得角门将心理优势(0-0西班牙成功经验)->埃米·马丁内斯vs沃齐尼亚。最可能冷门比分: 佛得角 1-0 (加时)或点球胜。概率~10%。',
    },
    {
        'num': 3, 'time': '09:30', 'stage': 'R32',
        'home': '哥伦比亚', 'away': '加纳',
        'home_flag': 'COL', 'away_flag': 'GHA',
        'venue': 'Arrowhead体育场, 堪萨斯城, 美国',
        'fifa_rank': '哥伦比亚 #10 vs 加纳 ~#35',
        'value_ratio': '哥伦比亚 180M vs 加纳 90M (约2:1)',
        'history': '首次交锋',
        'reward': '16强对阵 瑞士 (已确定: 瑞士2-0阿尔及利亚)',

        'group_review_home': '哥伦比亚 (K组第1, 7分) — 力压葡萄牙头名晋级',
        'group_review_home_data': [
            ['对手', '比分', '关键'],
            ['乌兹别克斯坦', '3-1', '迪亚斯+哈梅斯+苏亚雷斯前场三叉戟全面开火'],
            ['刚果民主共和国', '1-0', '稳扎稳打->刚果民主共和国大巴被破'],
            ['葡萄牙', '0-0', '[注意] 锁死C罗(0射正)+B费——哥伦比亚防守让超级巨星型葡萄牙哑火'],
        ],
        'group_review_away': '加纳 (L组第3, 4分) — 最佳第三名晋级',
        'group_review_away_data': [
            ['对手', '比分', '关键'],
            ['巴拿马', '1-0', '乔丹·阿尤制胜——中北美对手'],
            ['英格兰', '0-0', '[注意] 已验证大巴! 凯恩+贝林厄姆+拉什福德=被加纳零封!'],
            ['克罗地亚', '1-2', '被莫德里奇的克罗地亚击败->但进1球(阿尤)'],
        ],

        'player_ratings_home': [
            ['球员', '评分', '表现'],
            ['路易斯·迪亚斯(Luis Diaz) [火] MOTM', '9.0', '利物浦左边锋。本届1球。英超最佳左边锋之一->速度+盘带+射门'],
            ['哈梅斯·罗德里格斯(James Rodriguez)', '8.5', '33岁->世界杯大赛属性世界级。2014金靴。对葡萄牙送出3次关键传球'],
            ['路易斯·苏亚雷斯(Luis Suarez)', '7.8', '葡萄牙体育前锋——不是乌拉圭那个苏亚雷斯! 本届已进球->首发确认'],
            ['杰弗森·莱尔马(Jefferson Lerma)', '7.5', '水晶宫中场->防守屏障+身体对抗'],
            ['戴维森·桑切斯(Davinson Sanchez)', '7.5', '加拉塔萨雷中卫->哥伦比亚防线基石'],
        ],
        'player_ratings_away': [
            ['球员', '评分', '表现'],
            ['托马斯·帕尔特伊(Thomas Partey) [火]', '8.0', '阿森纳中场。对英格兰0-0=全场最佳——他的拦截+屏障是加纳大巴核心'],
            ['乔丹·阿尤(Jordan Ayew) [C]', '7.5', '34国家队进球->本届2球(加纳全部进球)。水晶宫->加纳唯一破局者'],
            ['安托万·塞梅尼奥(Antoine Semenyo)', '7.3', '曼城前锋->但[疑]踝伤->若缺阵加纳进攻只剩阿尤'],
            ['卡马尔丁·苏莱曼纳(Sulemana)', '7.3', '南安普顿->速度型边锋'],
            ['纳撒尼尔·阿杰泰(Adjetey)', '7.0', '中卫->对英格兰0-0防线基石'],
        ],

        'factors': [
            ('身价比2:1(180M vs 90M)', '哥伦比亚 [火][火]', '<3:1=不是碾压优势。哥伦比亚占优但不稳'),
            ('加纳已验证大巴(0-0英格兰!)', '加纳 [火][火][火]', '[注意] 凯恩+贝林厄姆+拉什福德被零封=这不是偶然。帕尔特伊(阿森纳)是核心'),
            ('哥伦比亚1失球(R32最佳之一)', '哥伦比亚 [火][火][火]', '对葡萄牙0-0->防线可有效限制C罗(0射正)。戴维森·桑切斯+卢库米+巴尔加斯'),
            ('加纳进攻2球(阿尤单核)', '哥伦比亚 [火][火][火]', '[7/3教训] 漏洞需对手有终结能力。加纳大巴世界级但阿尤是唯一得分手'),
            ('迪亚斯(利物浦)+哈梅斯破大巴能力', '哥伦比亚 [火][火][火]', '哈梅斯世界杯创造力(2014金靴)+迪亚斯英超速度=比葡萄牙有更多维度破大巴'),
            ('哥伦比亚力压葡萄牙头名晋级', '哥伦比亚 [火][火]', '7分=小组最佳之一。信心+状态在巅峰'),
            ('塞梅尼奥(曼城)踝伤可能缺阵', '哥伦比亚 [火][火][火]', '加纳进攻=阿尤单核->如果再缺塞梅尼奥->加纳只能依靠大巴+定位球+运气'),
            ('帕尔特伊(阿森纳)防守中场', '加纳 [火][火][火]', '对英格兰0-0=全场最佳。哈梅斯能否绕过他是比赛的关键对位'),
        ],
        'strength': '哥伦比亚: 体系型(南美防守版)——180M < 500M。0超巨(迪亚斯是核心级~50M)。但防守1失球=R32最佳之一。力压葡萄牙头名=已证明体系>个人  |  加纳: 大巴专家型(对英格兰已验证)——~90M。对英格兰0-0=世界级防守。帕尔特伊(阿森纳)=中场屏障。阿尤(34国家队进球)=比佛得角的利夫拉门托更好的终结者',
        'resilience_label': '非洲韧性评估(加纳)',
        'resilience': [
            ('低位防守', '世界级', '[注意] 对英格兰0-0=已验证。凯恩+贝林厄姆+拉什福德=被零封。帕尔特伊拦截+阿杰泰防空'),
            ('速度反击', '良好', '阿尤(34球)>佛得角的利夫拉门托(0球)。苏莱曼纳速度+塞梅尼奥(疑)'),
            ('定位球高点', '一般', '阿尤头球(经验+定位)，阿杰泰身高。帕尔特伊头球也是武器'),
            ('前30分钟', '良好', '对英格兰0-0前30分钟守住->对巴拿马1-0前30分钟领先'),
            ('被压制不崩盘', '世界级', '对英格兰被压制60分钟(60%控球)->0失球。对克罗地亚1-2落后仍有反击威胁'),
        ],
        'upset_risk': '中高',
        'upset_detail': '[注意] 本日最不确定——没有之一。加纳对英格兰0-0=已验证大巴。哥伦比亚对葡萄牙0-0=哥伦比亚自己面对大巴也可被限制。这场是两个防守强队的对决——0-0或1-0的比赛。哥伦比亚1-0首选但加纳1-0备选与首选概率最接近(40%-30%差距)。如果这场比赛进超过2球->将是非常意外的结果。',

        'predictions': [
            ('哥伦比亚 1-0', '~38%', '[首选——但优势仅3%] 加纳大巴(0-0英格兰)上半场守住->70-80分钟哈梅斯任意球/迪亚斯突破破局->加纳无法反扑->1-0守住'),
            ('加纳 1-0', '~35%', '[注意] 差距仅3%! 球员数据库发现:伊尼亚基·威廉姆斯(毕尔巴鄂竞技)在加纳名单->反击=阿尤+伊尼亚基双核。帕尔特伊(阿森纳)拦截->加纳角球->头球->1-0->5-4-1极致收缩->哥伦比亚0进球'),
            ('哥伦比亚 2-0', '~15%', '加纳大巴75-80分钟第一次失球后被迫打开->迪亚斯反击锁定2-0'),
            ('1-1 (加时决胜)', '~10%', '双方各进1球->加纳定位球破门(阿尤)+哥伦比亚追平(哈梅斯)->加时板凳深度决定胜负'),
            ('0-0 (点球决胜)', '~5%', '双方大巴互锁120分钟->点球->谁门将强谁赢'),
        ],

        'injuries': [
            ['球队', '球员', '状态', '影响'],
            ['哥伦比亚', '—', '全队健康', '内斯托尔·洛伦佐有完整阵容'],
            ['哥伦比亚', '路易斯·苏亚雷斯(Suarez)', '满血', '对葡萄牙轮休->满血回归'],
            ['加纳', '安托万·塞梅尼奥(Semenyo)', '疑', '曼城前锋——踝伤。加纳仅有的第二得分点'],
            ['加纳', '劳伦斯·阿蒂-齐吉(Ati Zigi)', '疑', '主力门将身体不适->阿萨雷(替补)对英格兰0-0表现稳健'],
            ['加纳', '—', '无停赛', '—'],
        ],

        'coach': [
            '哥伦比亚(内斯托尔·洛伦佐): 4-2-3-1哈梅斯AM+迪亚斯LW+苏亚雷斯CF->普埃尔塔+莱尔马防守双后腰。对葡萄牙0-0展示了他可锁死超级巨星(C罗0射正)->但这次需展示他可破大巴。领先: 5-4-1收缩+迪亚斯反击。',
            '加纳(卡洛斯·奎罗斯): [注意] 大巴专家+前皇马教练。对英格兰战术->全队在后防+帕尔特伊屏障+阿尤单前锋反击。奎罗斯对南美足球的了解(前伊朗教练)->对哥伦比亚无知识盲区。领先: 5-4-1极致收缩+阿尤反击。',
            '关键博弈: 奎罗斯是非洲版"佩特科维奇"——但比佩特科维奇更危险。佩特科维奇知道瑞士的一切但无法执行->奎罗斯有可以执行的球员(帕尔特伊拦截+阿尤34球)。奎罗斯的大巴不是"被迫大巴"——是"主动选择的已验证大巴"。',
        ],

        'set_piece': '哥伦比亚进攻: 哈梅斯任意球(世界级->2014世界杯金靴一部分就是定位球)。桑切斯+卢库米头球。哈梅斯左脚任意球是哥伦比亚破大巴最直接武器->比梅西任意球略低一个级别但仍是世界级。加纳进攻: 阿尤头球(34球中很多是头球)+帕尔特伊定位球射门。加纳定位球是他们在0-0英格兰中唯一接近进球的方式。',

        'upset_path': '如果翻车(加纳赢): 帕尔特伊(阿森纳)全场拦截->哈梅斯被完全限制->0-0到70分钟->加纳角球->阿尤头球->1-0->加纳5-4-1极致收缩(对英格兰0-0已验证120分钟都可守住)->哥伦比亚全线压上->迪亚斯被金斯利·门萨限制->哈梅斯任意球被阿萨雷扑出->0-1。最可能冷门比分: 加纳 1-0 (阿尤头球致胜)。这不是"低概率冷门"——这是"中概率冷门"。',
    },
]

# Render each match
for m in matches:
    add_heading_styled(f'比赛{m["num"]}: {m["home_flag"]} {m["home"]} vs {m["away_flag"]} {m["away"]}  ({m["stage"]} {m["time"]} BJT)', level=2)

    # Basic info
    add_para('基本信息', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))
    info_lines = [
        f'场地: {m["venue"]}',
        f'FIFA排名: {m["fifa_rank"]}  |  身价比: {m["value_ratio"]}',
        f'历史交锋: {m["history"]}  |  晋级奖励: {m["reward"]}',
    ]
    for line in info_lines:
        add_para(line, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # Group stage review
    add_para('小组赛回顾', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    add_para(m['group_review_home'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_home_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_home_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_home_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    add_para(m['group_review_away'], bold=True, size=Pt(9), space_after=Pt(2))
    t = doc.add_table(rows=len(m['group_review_away_data']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['group_review_away_data'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['group_review_away_data'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(t.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    doc.add_paragraph()

    # Player ratings
    add_para('首轮球员评分', bold=True, size=Pt(10), color=DARK, space_after=Pt(3))

    add_para(f'{m["home"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_home']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_home'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_home'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    add_para(f'{m["away"]}', bold=True, size=Pt(9), space_after=Pt(1))
    t = doc.add_table(rows=len(m['player_ratings_away']), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(m['player_ratings_away'][0]):
        set_cell(t.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(m['player_ratings_away'][1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(t.rows[i].cells[0], row[0], size=Pt(7.5))
        set_cell(t.rows[i].cells[1], row[1], size=Pt(8))
        set_cell_left(t.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in t.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(5.9)
    doc.add_paragraph()

    # Factors table
    add_para('因素导向表', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    ft_data = [['因素', '有利方', '理由']]
    ft_data.extend(m['factors'])
    ft = doc.add_table(rows=len(ft_data), cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(ft_data[0]):
        set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(8.5), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(ft_data[1:], 1):
        set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8))
        clr = None
        if '[火][火][火]' in row[1] or '[火][火]' in row[1]:
            clr = ACCENT
        set_cell(ft.rows[i].cells[1], row[1], size=Pt(8), color=clr)
        set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in ft.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(5.0)
    doc.add_paragraph()

    # Strength + Resilience
    add_para(f'强队分类: {m["strength"]}', size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))

    add_para(m['resilience_label'], bold=True, size=Pt(9), color=DARK, space_after=Pt(2))
    res_data = [['维度', '评分', '说明']]
    res_data.extend(m['resilience'])
    rt = doc.add_table(rows=len(res_data), cols=3)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(res_data[0]):
        set_cell(rt.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(res_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(rt.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell(rt.rows[i].cells[1], row[1], size=Pt(8), bg=bg)
        set_cell_left(rt.rows[i].cells[2], row[2], size=Pt(7.5))
    for row in rt.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(7.5)
    doc.add_paragraph()

    # Upset risk
    risk_color = ACCENT if m['upset_risk'] in ('中', '中高', '高') else DARK
    add_para(f'冷门风险: {m["upset_risk"]}', bold=True, size=Pt(10), color=risk_color, space_after=Pt(2))
    add_para(m['upset_detail'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # Predictions table
    add_para('比分预测', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    pred_header = ['比分', '概率', '说明']
    pred_rows = []
    for p in m['predictions']:
        pred_rows.append(list(p))
    pt = doc.add_table(rows=1 + len(pred_rows), cols=3)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(pred_header):
        set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(pred_rows, 1):
        bg = GREEN_BG if i == 1 else None
        set_cell(pt.rows[i].cells[0], row[0], bold=True, size=Pt(10), bg=bg)
        set_cell(pt.rows[i].cells[1], row[1], bold=(i==1), size=Pt(9), bg=bg)
        set_cell_left(pt.rows[i].cells[2], row[2], size=Pt(8))
    for row in pt.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(8.4)
    doc.add_paragraph()

    # Injuries
    add_para('伤病/停赛', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    inj_data = m['injuries']
    it = doc.add_table(rows=len(inj_data), cols=4)
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(inj_data[0]):
        set_cell(it.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(inj_data[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(it.rows[i].cells[0], row[0], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[1], row[1], size=Pt(7.5))
        set_cell(it.rows[i].cells[2], row[2], size=Pt(8), bg=bg)
        set_cell_left(it.rows[i].cells[3], row[3], size=Pt(7.5))
    for row in it.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(5.2)
    doc.add_paragraph()

    # Coach tactics
    add_para('教练博弈', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    for c in m['coach']:
        add_para(c, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(1))
    doc.add_paragraph()

    # Set piece
    add_para('定位球攻防', bold=True, size=Pt(10), color=DARK, space_after=Pt(2))
    add_para(m['set_piece'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(4))
    doc.add_paragraph()

    # Upset path
    add_para('冷门路径', bold=True, size=Pt(10), color=ACCENT, space_after=Pt(2))
    add_para(m['upset_path'], size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

    # separator between matches — v6 style blue rule
    sep_para = doc.add_paragraph()
    sep_para.paragraph_format.space_before = Pt(12)
    sep_para.paragraph_format.space_after = Pt(12)
    add_para_border(sep_para, sz='6')

# RISK SUMMARY
add_heading_styled('风险提示', level=2)

risk_notes = [
    ('澳大利亚 vs 埃及 [注意] 萨拉赫的腿筋=唯一决定性变量', '如果萨拉赫100%首发->埃及2-1或3-1。如果萨拉赫替补(最可能)->埃及2-1(首选)或1-1(加时)。如果萨拉赫缺阵->澳大利亚可能2-0。预测基于"萨拉赫替补60分钟后上场"的假设。埃及防线3-4人伤缺+拉欣停赛->澳大利亚定位球(苏塔尔1.98m)能进1球。'),
    ('阿根廷 vs 佛得角 [注意] 已验证大巴(0-0西班牙) (有) 已验证平局', '佛得角0-0西班牙=大巴验证。但阿根廷有梅西(6球)->梅西的破大巴能力比西班牙的佩德里+奥尔莫更强。半场0-0可能但90分钟0-0概率~25%。2-0首选意味着梅西在第70-80分钟破局。"漏洞需终结能力"——佛得角进攻=2球->无法惩罚阿根廷。'),
    ('哥伦比亚 vs 加纳 [注意] 本日最不确定——没有之一', '加纳对英格兰0-0=已验证大巴。哥伦比亚对葡萄牙0-0=哥伦比亚自己面对大巴也可被限制。这场是两个防守强队的对决——0-0或1-0的比赛。哥伦比亚1-0首选但加纳1-0备选与首选概率最接近(40%-30%差距)。如果这场比赛进超过2球->将是非常意外的结果。'),
]

for title, detail in risk_notes:
    add_para(title, bold=True, size=Pt(9), color=ACCENT, space_after=Pt(1))
    add_para(detail, size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

# page bottom separator — v6 blue rule
sep_para = doc.add_paragraph()
sep_para.paragraph_format.space_before = Pt(16)
sep_para.paragraph_format.space_after = Pt(4)
add_para_border(sep_para, sz='6')

add_para('数据源: ESPN API + Sports Mole + BBC Sport + ABC News + RotoWire + FIFA  |  框架: CLAUDE.md v16 + 7月3日三场复盘教训(漏洞价值公式+已验证vs被迫大巴)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('预测时间: 2026年7月3日 14:00 北京时间  |  教训加权: 漏洞价值=漏洞x终结者 / 教练知识优势(有)比赛优势 / 大巴分"已验证"和"被迫"',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('预测仅供分析参考  |  生成工具: python-docx v6美化版',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

# Fix: add w:percent to zoom (validation requirement)
settings = doc.settings._element
zoom = settings.find(qn('w:zoom'))
if zoom is not None:
    zoom.set(qn('w:percent'), '100')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

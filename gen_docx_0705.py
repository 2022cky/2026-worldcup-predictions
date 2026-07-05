# -*- coding: utf-8 -*-
"""Generate 2026-07-05 predictions DOCX — Chinese names, no emoji markers, corrected widths."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月5日_2场预测.docx")

doc = Document()

# ── Colours ──
NAVY   = RGBColor(0x1A, 0x2E, 0x3D)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW = RGBColor(0xF5, 0xF7, 0xFA)
PREF_ROW= RGBColor(0xE8, 0xF5, 0xE9)
BORDER  = RGBColor(0xD0, 0xD5, 0xDD)
RED_T   = RGBColor(0xC0, 0x39, 0x2B)
DARK    = RGBColor(0x1A, 0x1A, 0x2E)
META    = RGBColor(0x7F, 0x8C, 0x8D)
META_L  = RGBColor(0x95, 0xA5, 0xA6)
BLUE_T  = RGBColor(0x2E, 0x75, 0xB6)
ORANGE  = RGBColor(0xE6, 0x7E, 0x22)
GREEN_T = RGBColor(0x27, 0xAE, 0x60)

FONT = '微软雅黑'

# Page setup — landscape
for sec in doc.sections:
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
    sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)

CONTENT_W = 27.3  # cm (29.7 - 2*1.2)

# ─── helpers ───
def sf(run, size=None, color=None, bold=False, font=FONT):
    run.font.name = font; run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.bold = bold

def arun(p, text, **kw):
    r = p.add_run(text); sf(r, **kw)
    return r

def mp(before=0, after=0, align=None, border=False, bc='2E75B6'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    if border:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        for a, v in [('w:val','single'),('w:sz','4'),('w:space','4'),('w:color',bc)]:
            bot.set(qn(a), v)
        pBdr.append(bot)
        for tag in ['w:spacing','w:ind','w:jc','w:rPr']:
            el = pPr.find(qn(tag))
            if el is not None: pPr.insert(list(pPr).index(el), pBdr); break
        else: pPr.append(pBdr)
    return p

def h1(t): return arun(mp(before=16,after=4,border=True,bc='C0392B'), t, size=16, bold=True, color=RED_T)
def h2(t): return arun(mp(before=12,after=3), t, size=12, bold=True, color=DARK)
def body(t, size=8.5, color=None, indent=False):
    p = mp(before=1,after=2)
    if indent: p.paragraph_format.left_indent = Cm(0.5)
    arun(p, t, size=size, color=color or DARK)
    return p
def meta(t): return arun(mp(before=2,after=2), t, size=7, color=META)

def _shade(cell, clr):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), str(clr)); shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def _cb(cell, clr=BORDER, sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        b = OxmlElement(f'w:{edge}')
        for a, v in [('w:val','single'),('w:sz',sz),('w:color',str(clr)),('w:space','0')]:
            b.set(qn(a), v)
        borders.append(b)
    tcPr.append(borders)

def cpara(cell, text, size=8, bold=False, color=None, align=None):
    for pp in cell.paragraphs: pp.clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    if align is not None: p.alignment = align
    r = p.add_run(str(text)); sf(r, size=size, bold=bold, color=color or DARK)
    return p

def make_table(headers, rows, widths, pref=None):
    """widths list in cm, must sum to CONTENT_W"""
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    # Header
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; _shade(c, NAVY)
        cpara(c, h, size=7.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER); _cb(c, clr=RGBColor(0x15,0x29,0x38))
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            if pref is not None and ri == pref: _shade(c, PREF_ROW)
            elif ri % 2 == 0: _shade(c, ALT_ROW)
            al = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else None
            cpara(c, val, size=7.5, align=al); _cb(c)
    # Widths
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    return tbl

# ── header / footer ──
for sec in doc.sections:
    h = sec.header; h.is_linked_to_previous = False
    hp0 = h.paragraphs[0]; hp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp0.paragraph_format.space_after = Pt(2)
    r = hp0.add_run('2026 FIFA 世界杯 · 7月5日 两场预测报告')
    sf(r, size=7, color=META_L, font='Arial'); r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pPr = hp0._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    for a, v in [('w:val','single'),('w:sz','4'),('w:space','4'),('w:color','95A5A6')]:
        bot.set(qn(a), v)
    pBdr.append(bot); pPr.insert(0, pBdr)

    f = sec.footer; f.is_linked_to_previous = False
    fp = f.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    rp = fp.add_run(); sf(rp, size=7, color=META_L, font='Arial'); rp._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); rp._r.append(fc1)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; rp._r.append(it)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); rp._r.append(fc2)

# ═══════════ COVER ═══════════
p = mp(before=60, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '2026 FIFA 世界杯', size=26, bold=True, color=RED_T)
p = mp(before=2, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '淘汰赛十六强 · 7月5日 两场预测报告', size=15, bold=True, color=DARK)
p = mp(before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '加拿大 vs 摩洛哥  |  巴拉圭 vs 法国', size=12, bold=True, color=BLUE_T)

p = mp(before=16, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '预测时间: 2026年7月4日 14:00 北京时间', size=8.5, color=META)
p = mp(before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '数据来源: FIFA API + ESPN + RotoWire + SI + Sporting News + The Standard + Al Jazeera', size=7, color=META_L)

p = mp(before=8, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
arun(p, '[注意] 7/4教训: 大巴进球看对手不看小组赛总量 | 了解对手≠优势 | 强队破大巴类型不同 | 中卫进攻是破大巴变量 | 教练保守声明=烟雾弹', size=7, color=RED_T)

# ═══════════ SUMMARY ═══════════
h1('一、预测汇总')
meta('数据截至: 2026-07-04 14:00 CST | 淘汰赛十六强 | 所有时间均为北京时间')
make_table(
    ['#', '时间', '阶段', '比赛', '身价比', '首选', '概率', '备选', '冷门'],
    [
        ['1', '01:00', '十六强', '加拿大 vs 摩洛哥', '1:2.3', '摩洛哥 2-0', '~42%', '摩洛哥 1-0 / 2-1 / 1-1(加时)', '低-中'],
        ['2', '05:00', '十六强', '巴拉圭 vs 法国', '1:10', '法国 3-0', '~48%', '法国 2-0 / 4-0 / 法国 2-1', '极低'],
    ],
    widths=[0.5, 1.3, 1.0, 3.0, 1.1, 2.0, 1.0, 4.5, 1.2],
)

# ═══════════ BRACKET ═══════════
h1('二、淘汰赛路径')
make_table(
    ['日期(北京时间)', '阶段', '对阵', '场地'],
    [
        ['7/5 01:00', '十六强', '加拿大 vs 摩洛哥', 'NRG体育场, 休斯顿'],
        ['7/5 05:00', '十六强', '巴拉圭 vs 法国', '林肯金融球场, 费城'],
        ['7/7 03:00', '十六强', '西班牙 vs 葡萄牙', 'AT&T球场, 阿灵顿'],
        ['7/7 08:00', '十六强', '比利时 vs 美国', '流明球场, 西雅图'],
        ['7/8 00:00', '十六强', '阿根廷 vs 埃及', '梅赛德斯-奔驰球场, 亚特兰大'],
        ['7/8 04:00', '十六强', '哥伦比亚 vs 瑞士', 'BC Place, 温哥华'],
    ],
    widths=[3.0, 1.2, 4.5, 4.5],
)
p = mp(before=4, after=0); arun(p, '四分之一决赛路径: 法国若晋级 → 对加拿大或摩洛哥 → 无论对手是谁法国均占优(身价比>3:1)', size=8, bold=True, color=BLUE_T)

# ═══════════ MATCH 1 ═══════════
h1('三、比赛1: 加拿大 vs 摩洛哥')
meta('7月5日 01:00 北京时间 | NRG体育场, 休斯顿 | FIFA #30 vs #6 | 历史: 摩洛哥3胜1平')

h2('3.1 基本信息')
make_table(
    ['项目', '内容'],
    [['时间', '7月5日 01:00 北京时间 (7/4 13:00 EDT)'],
     ['场地', 'NRG体育场, 休斯顿, 美国'],
     ['FIFA排名', '加拿大 #30 vs 摩洛哥 #6'],
     ['阵容身价', '加拿大 1.99亿欧元 vs 摩洛哥 4.48亿欧元 (约1:2.3)'],
     ['历史交锋', '摩洛哥3胜1平 — 最近一次: 2022世界杯小组赛 摩洛哥2-1加拿大'],
     ['晋级奖励', '四分之一决赛 vs 巴拉圭/法国胜者 (7/9 波士顿)']],
    widths=[2.5, 10.8],
)

h2('3.2 小组赛及淘汰赛回顾')
make_table(
    ['球队', '对手', '比分', '关键'],
    [['加拿大', '卡塔尔 (小组赛)', '6-0', '对手9人 — 加拿大世界杯首胜'],
     ['加拿大', '波黑 (小组赛)', '1-1', '逼平'],
     ['加拿大', '瑞士 (小组赛)', '0-2', '不敌瑞士 — 但仍以4分晋级'],
     ['加拿大', '南非 (三十二强)', '1-0', '欧斯塔基奥 92分钟绝杀 — 首次淘汰赛胜利'],
     ['摩洛哥', '巴西 (小组赛)', '1-1', '逼平身价10倍对手 — 小组最强表现'],
     ['摩洛哥', '苏格兰 (小组赛)', '2-0', '稳扎稳打'],
     ['摩洛哥', '海地 (小组赛)', '3-0', '7分小组第二'],
     ['摩洛哥', '荷兰 (三十二强)', '1-1(点球3-2)', '120分钟 — 点球淘汰荷兰 — 非洲冠军韧性']],
    widths=[1.2, 2.2, 2.5, 5.2],
)

h2('3.3 核心球员评分')
make_table(
    ['球队', '球员', '评分', '位置/俱乐部', '关键信息'],
    [['加拿大', '阿方索·戴维斯 [C]', '8.5 [关键]', '左后卫 / 拜仁', '腿筋缺席全部小组赛, 三十二强替补20分钟, 状态生锈 — 可能首次首发'],
     ['加拿大', '斯蒂芬·欧斯塔基奥', '8.0 [最佳]', '中场 / 波尔图', '三十二强92分钟制胜球 — 加拿大英雄 — 定位球+远射+后排插上'],
     ['加拿大', '乔纳森·戴维', '8.0', '前锋 / 里尔', '最可靠得分手 — 本届已进球'],
     ['加拿大', '塔容·布坎南', '7.3', '边锋 / 国际米兰', '速度+突破 — 右路威胁'],
     ['加拿大', '马克西姆·克雷波', '7.5', '门将 / 波特兰', '三十二强多次关键扑救'],
     ['摩洛哥', '伊斯梅尔·赛巴里', '9.0 [最佳]', '前锋 / 拜仁新援', '本届3球 — 十六强最危险的前锋 — 对荷兰罚进制胜点球'],
     ['摩洛哥', '布拉欣·迪亚兹', '8.5', '攻击中场 / 皇家马德里', '2助攻 — 创造力+传球 — 任何防守中都能找到空隙'],
     ['摩洛哥', '阿什拉夫·哈基米 [C]', '8.5', '右后卫 / 巴黎圣日耳曼', '世界前三右后卫 — 攻防一体 — 120分钟消耗仅休3天'],
     ['摩洛哥', '努赛尔·马兹拉维', '7.8', '左后卫 / 曼联', '与哈基米组成世界级边卫组合'],
     ['摩洛哥', '博诺', '7.8', '门将', '2022世界杯英雄 — 点球大战经验丰富']],
    widths=[1.0, 3.2, 1.5, 2.5, 5.1],
)

h2('3.4 因素导向')
make_table(
    ['因素', '有利', '理由'],
    [['赛巴里(3球)+迪亚兹(2助)+哈基米: 三维度攻击', '摩洛哥 ★★★', '边路推进+创造力+终结 — 加拿大防线从未面对此级别组合'],
     ['戴维斯(可能首发)状态生锈', '摩洛哥 ★★★', '教练马什原话"他看起来很生疏" — 若仅六成状态则左路攻防不如预期'],
     ['摩洛哥2022四强淘汰赛经验', '摩洛哥 ★★★', '淘汰赛心理已验证 — 加拿大首次十六强属未知领域'],
     ['戴维斯若百分百首发', '加拿大 ★★★', '本场最大变量 — 健康戴维斯=世界最佳左后卫'],
     ['摩洛哥120分钟点球消耗', '加拿大 ★★', '哈基米+埃尔艾纳维120分钟 — 休仅3天 — 加拿大90分钟+绝杀士气'],
     ['欧斯塔基奥绝杀能力', '加拿大 ★★', '三十二强92分钟绝杀 — 加拿大最后时刻有得分能力'],
     ['加拿大主场(联合主办国)', '加拿大 ★★', '休斯顿大批加拿大球迷南下 — 场面优势'],
     ['身价比 1:2.3', '摩洛哥 ★★', '明显差距但不碾压 — 摩洛哥占优非稳赢']],
    widths=[4.5, 2.0, 6.8],
)

h2('3.5 韧性评估')
make_table(
    ['球队', '低位防守', '速度反击', '定位球', '前30分钟', '被压不崩', '综合', '说明'],
    [['摩洛哥', '世界级', '世界级', '良好', '良好', '世界级', '5/5', '2022四强验证 — 进攻韧性远超佛得角'],
     ['加拿大', '一般', '良好', '良好', '良好', '良好', '3/5', '首次十六强 — 防线面对世界级攻击未验证']],
    widths=[1.2, 1.1, 1.1, 0.9, 1.1, 1.1, 0.8, 4.2],
)

h2('3.6 比分预测')
make_table(
    ['首选', '概率', '半场', '逻辑'],
    [['摩洛哥 2-0', '~42%', '0-0', '半场0-0(摩洛哥体能影响开局) — 55-70分钟赛巴里破门 — 加拿大压出 — 哈基米反击锁定'],
     ['摩洛哥 1-0', '~20%', '0-0', '摩洛哥体能影响严重 — 一球小胜 — 加拿大无法破门'],
     ['摩洛哥 2-1', '~15%', '0-0', '摩洛哥2-0 — 戴维斯助攻扳回 — 摩洛哥守住'],
     ['1-1 (加时)', '~13%', '—', '加拿大定位球先破 — 摩洛哥追平 — 加时摩洛哥板凳晋级'],
     ['加拿大 1-0', '~10%', '—', '戴维斯百分百+欧斯塔基奥定位球 — 历史性四分之一决赛']],
    widths=[2.2, 1.0, 0.8, 8.1],
    pref=0,
)
p = mp(before=4, after=0); arun(p, '冷门风险: 低-中 | 平局概率约20% | 加拿大赢约22% | 条件: 戴维斯百分百+欧斯塔基奥定位球+摩洛哥体能崩+克雷波英雄', size=8, bold=True, color=BLUE_T)

h2('3.7 伤病与停赛')
make_table(
    ['球队', '球员', '状态', '影响'],
    [['加拿大', '伊斯梅尔·科内 (中场)', '[缺] 断腿已离队', '严重 — 中场核心缺失 — 萨利巴替代'],
     ['加拿大', '阿方索·戴维斯 (左后卫)', '[疑] 腿筋 — 三十二强替补20分钟', '严重 — 5月6日以来首次可能首发 — 状态生疏'],
     ['摩洛哥', '全队健康', '[可]', '注意 — 但多名核心120分钟消耗(休整仅3天)']],
    widths=[1.5, 3.5, 3.0, 5.3],
)

h2('3.8 冷门路径')
body('加拿大赢需同时满足: (1)戴维斯百分百状态首发,左路压制哈基米 (2)欧斯塔基奥定位球或角球,乔纳森·戴维头球破门 1-0 (3)摩洛哥120分钟消耗致进球效率下降 (4)克雷波多次世界级扑救 (5)加拿大主场心理优势 — 概率约一成', size=8)

# ═══════════ MATCH 2 ═══════════
h1('四、比赛2: 巴拉圭 vs 法国')
meta('7月5日 05:00 北京时间 | 林肯金融球场, 费城 | FIFA #35 vs #2 | 历史: 法国2胜 (1958/1998)')

h2('4.1 基本信息')
make_table(
    ['项目', '内容'],
    [['时间', '7月5日 05:00 北京时间 (7/4 17:00 EDT)'],
     ['场地', '林肯金融球场 (Lincoln Financial Field), 费城, 美国'],
     ['FIFA排名', '巴拉圭 #35 vs 法国 #2'],
     ['阵容身价', '巴拉圭 1.54亿欧元 vs 法国 15.2亿欧元 (约1:10)'],
     ['历史交锋', '2次 — 1958年法国7-3 / 1998年十六强法国1-0(金球制)'],
     ['晋级奖励', '四分之一决赛 vs 加拿大/摩洛哥胜者 (7/9 波士顿)']],
    widths=[2.5, 10.8],
)

h2('4.2 小组赛及淘汰赛回顾')
make_table(
    ['球队', '对手', '比分', '关键'],
    [['巴拉圭', '美国 (小组赛)', '1-4', '被东道主碾压 — 面对顶级攻击力大巴崩溃'],
     ['巴拉圭', '土耳其 (小组赛)', '1-0', '稳扎稳打'],
     ['巴拉圭', '澳大利亚 (小组赛)', '0-0', '大巴拿1分 — 后来证明是防守质量'],
     ['巴拉圭', '德国 (三十二强)', '1-1(点球4-3)', '淘汰四届冠军 — 恩西索先破 — 吉尔点球英雄'],
     ['法国', '塞内加尔 (小组赛)', '3-1', '开门红'],
     ['法国', '伊拉克 (小组赛)', '3-0', '碾压式'],
     ['法国', '挪威 (小组赛)', '4-1', '哈兰德在仍轻松取胜'],
     ['法国', '瑞典 (三十二强)', '3-0', '姆巴佩+登贝莱+奥利塞不可阻挡']],
    widths=[1.2, 2.2, 2.5, 5.2],
)
p = mp(before=4, after=0); arun(p, '[关键] 法国: 12进球2失球=世界最佳攻防平衡。姆巴佩6球(并列金靴), 登贝莱4球, 奥利塞5助(距贝利纪录差1)。4战4胜场均3球从未落后 — 本届最强球队。', size=8, bold=True, color=RED_T)

h2('4.3 核心球员评分')
make_table(
    ['球队', '球员', '评分', '位置/俱乐部', '关键信息'],
    [['巴拉圭', '胡利奥·恩西索', '8.0 [最佳]', '前锋 / 斯特拉斯堡', '19岁 — 对德国首开纪录 — 巴拉圭唯一能威胁法国的攻击手'],
     ['巴拉圭', '迭戈·戈麦斯', '7.5', '中场 / 布莱顿', '停赛复出 — 对德国缺阵 — 本场回归=中场控制力升级'],
     ['巴拉圭', '奥兰多·吉尔', '8.5', '门将', '对德国点球英雄 — 扑出2个 — 但面对姆巴佩(6球)=完全不同级别'],
     ['巴拉圭', '安德烈斯·库巴斯 [C]', '7.5', '后腰 / 温哥华', '对德国120分钟防守核心 — 中场屏障'],
     ['巴拉圭', '古斯塔沃·戈麦斯', '7.5', '中卫 / 帕尔梅拉斯', '防空+身体对抗 — 面对姆巴佩最可靠屏障'],
     ['法国', '基利安·姆巴佩 [C]', '9.8 [最佳]', '前锋 / 皇家马德里', '6球=与梅西并列金靴 — 巴拉圭无人可单防 — 需2-3人包夹'],
     ['法国', '迈克尔·奥利塞', '9.5', '攻击中场 / 拜仁', '5助攻=十六强最高 — 距贝利单届纪录仅差1 — 创造力+传球+任意球'],
     ['法国', '奥斯曼·登贝莱', '9.0', '右边锋 / 巴黎圣日耳曼', '4球 — 速度+盘带 — 右路=第三维度攻击'],
     ['法国', '布拉德利·巴科拉', '8.5', '左边锋 / 巴黎圣日耳曼', '2球 — 左路=第四维度攻击'],
     ['法国', '威廉·萨利巴', '8.0', '中卫 / 阿森纳', '世界前三中卫 — 零封瑞典']],
    widths=[1.0, 3.0, 1.5, 2.5, 5.3],
)

h2('4.4 因素导向')
make_table(
    ['因素', '有利', '理由'],
    [['身价比 1:10 (1.54亿 vs 15.2亿欧元)', '法国 ★★★', '>10:1=巨大差距。法国正常发挥即胜 — 不做平局预测'],
     ['姆巴佩(6球)+登贝莱(4)+奥利塞(5助)+巴科拉(2)', '法国 ★★★', '四个攻击维度 — 包夹姆巴佩则奥利塞+登贝莱空间更大 — 大巴无法全封'],
     ['巴拉圭对美国1-4', '法国 ★★★', '面对顶级速度+创造力=大巴崩溃 — 法国攻击力≥美国数倍'],
     ['法国4战4胜场均3球从未落后', '法国 ★★★', '巴拉圭需领先才能激活大巴 — 但可能永远无法领先'],
     ['恩西索(对德国进球)+阿尔米隆速度', '巴拉圭 ★★', '法国防线可能的弱点 — 孔德进攻倾向强 — 恩西索可寻此侧空间'],
     ['戈麦斯(布莱顿)停赛复出', '巴拉圭 ★★', '对德国中场控制不足 — 戈麦斯回归=中场升级'],
     ['1998十六强: 法国1-0巴拉圭(金球)', '巴拉圭 ★', '历史暗示可拖入加时 — 但那是28年前齐达内时代']],
    widths=[4.5, 2.0, 6.8],
)

h2('4.5 韧性评估')
make_table(
    ['球队', '低位防守', '速度反击', '定位球', '前30分钟', '被压不崩', '综合', '说明'],
    [['巴拉圭', '良好', '良好', '一般', '良好', '良好', '3.5/5', '对德国验证了面对体系型能力 — 但法国是超巨型×4'],
     ['法国', '世界级', '世界级', '世界级', '世界级', '世界级', '5/5', '本届最强 — 全部维度世界级 — 唯一风险: 从未落后过']],
    widths=[1.2, 1.1, 1.1, 0.9, 1.1, 1.1, 0.8, 4.2],
)

h2('4.6 比分预测')
make_table(
    ['首选', '概率', '半场', '逻辑'],
    [['法国 3-0', '~48%', '1-0', '半场1-0(姆巴佩20-30分钟) — 下半场登贝莱反击 — 巴科拉/奥利塞锁定3-0'],
     ['法国 2-0', '~25%', '0-0', '巴拉圭大巴比预期好 — 55-65分钟姆巴佩破局 — 登贝莱锁定'],
     ['法国 4-0', '~15%', '2-0', '法国早早破门(10-15分钟) — 巴拉圭被迫放弃大巴 — 与美国1-4节奏一致'],
     ['法国 2-1', '~7%', '—', '法国2-0 — 恩西索定位球追回1球 — 法国不慌'],
     ['巴拉圭1-0/点球', '~5%', '—', '需全部条件同时发生 — 概率仅佛得角对阿根廷的一半']],
    widths=[2.5, 1.0, 0.8, 7.8],
    pref=0,
)
p = mp(before=4, after=0)
arun(p, '冷门风险: 极低 [低于阿根廷对佛得角] | 平局概率约5% | 巴拉圭赢约5% | 身价比1:10 — 不做平局预测', size=8, bold=True, color=RED_T)

h2('4.7 伤病与停赛')
make_table(
    ['球队', '球员', '状态', '影响'],
    [['巴拉圭', '迭戈·戈麦斯 (中场)', '[可] 停赛复出→首发', '利好 — 布莱顿中场回归 — 控制力升级'],
     ['法国', '马库斯·图拉姆 (前锋)', '[疑] 小腿轻伤', '轻微 — 国米轮换球员 — 不影响首发'],
     ['法国', '全队健康无停赛', '[可]', '利好 — 德尚完整武器库 — 板凳: 科曼/杜埃/福法纳均为世界级']],
    widths=[1.5, 3.5, 3.0, 5.3],
)

h2('4.8 冷门路径')
body('巴拉圭赢需同时满足: (1)恩西索15-25分钟利用孔德前压空间远射破门 1-0 (2)吉尔踢出比德国战更出色的120分钟 (3)姆巴佩+登贝莱+奥利塞三人同时低效(极低概率) (4)戈麦斯回归中场=拦截升级 (5)法国首次面对淘汰赛逆境', size=8)
body('概率约5% — 佛得角有对乌拉圭和阿根廷的攻击韧性 — 巴拉圭只有打德国的一球。法国有4个超巨而阿根廷只有1个。巴拉圭冷门路径比佛得角窄50%。', size=8, color=RED_T)

# ═══════════ RISK ═══════════
h1('五、风险提示')
make_table(
    ['#', '场景', '风险', '应用'],
    [['1', '加拿大对摩洛哥: 戴维斯是否首发', '若戴维斯百分百→加拿大胜率升至25%。若状态生疏→摩洛哥2-0大概率', '赛前1小时确认首发后更新'],
     ['2', '巴拉圭对法国: 并非可能翻车的比赛', '1:10身价比。法国4战4胜场均3球从未落后。巴拉圭对德1-1≠对法能守住', '仅当法国多名核心伤缺则冷门概率>10%'],
     ['3', '大巴进球看对手不看总量', '巴拉圭对美仅1球→对强队终结力不如佛得角。恩西索被锁=零威胁', '淘汰赛进球预期=小组赛对强队进球数'],
     ['4', '"了解对手=优势"零数据', '本日无此前主帅/前球员场景→记录已存', '未来禁用此因子']],
    widths=[0.8, 3.5, 4.5, 4.5],
)

# ═══════════ FOOTER ═══════════
p = mp(before=18, after=0); arun(p, '数据来源: FIFA API + ESPN API + RotoWire + SI + Sporting News + The Standard + Al Jazeera + WhoScored + Goal.com + FotMob', size=7, color=META_L)
p = mp(before=2, after=0); arun(p, '分析框架: CLAUDE.md v16 + 7月4日三场复盘教训(大巴进球看对手+了解≠优势+强队破大巴类型不同+中卫进攻+教练烟雾弹)', size=7, color=META_L)
p = mp(before=2, after=0); arun(p, '生成时间: 2026年7月4日 14:00 CST | 引擎: python-docx v6 (横向A4) | 首发确认: 两场均未公布 — 赛前1-2小时需更新', size=7, color=META_L)

doc.save(OUT)
print(f'OK: {OUT} ({os.path.getsize(OUT):,} bytes)')

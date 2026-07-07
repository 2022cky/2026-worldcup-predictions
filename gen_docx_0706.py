# -*- coding: utf-8 -*-
"""Generate 2026-07-06 World Cup predictions as a formatted .docx file — OFFICIAL LINEUP UPDATE."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月6日_两场预测.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# colours
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GOLD   = RGBColor(0xD4, 0xA0, 0x17)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0x27, 0xAE, 0x60)
RED_BG     = RGBColor(0xE7, 0x4C, 0x3C)
ORANGE_BG  = RGBColor(0xF3, 0x9C, 0x12)
LIGHT_GREEN_BG = RGBColor(0xE8, 0xF8, 0xF5)

def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('7月6日 16强淘汰赛 预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('巴西 vs 挪威  ·  墨西哥 vs 英格兰', bold=False, size=Pt(11), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(10))
add_para('Round of 16  ·  MetLife Stadium / Estadio Azteca', bold=False, size=Pt(10), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(18))

add_para('生成时间: 2026年7月6日 北京时间 02:50  |  首发: 官方确认 (搜狐体育/AS USA/Sporting News)',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(4))
add_para('数据源: FIFA API + ESPN API + Sporting News + AS USA + ETtoday  |  框架: CLAUDE.md v17',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(6))
add_para('[注意] 高温41C体感  |  [注意] 拉菲尼亚回归替补!  |  [注意] 吕尔松伤愈回归挪威首发  |  裁判: 埃尔法特(美国)  |  预测仅供分析参考',
         size=Pt(8), color=ACCENT, align='center', space_after=Pt(20))

# ═══════════════════════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════════════════════
add_heading_styled('预测汇总', level=2)

summary_data = [
    ['#', '比赛 (BJT)', '身价比', '强队分类', '首选比分', '半场', '备选比分', '冷门风险'],
    ['91', 'BRA vs NOR (04:00)', '1.51:1', '双方均\n超巨型', '巴西 2-1\n挪威', '1-0', '2-2(加时)\n1-0 / 2-0', '中'],
    ['92', 'MEX vs ENG (08:00)', '1:6.8', 'ENG:\n超巨星型', '英格兰 2-0\n墨西哥', '1-0', '1-0 / 2-1\n1-1(加时)', '中'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 1.85, 0.55, 0.72, 1.05, 0.50, 1.30, 0.50]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        align = 'left' if j == 1 else 'center'
        bold = (j == 4)
        if j == 1:
            set_cell_left(table.rows[i].cells[j], text, bold=False, size=Pt(9))
        else:
            set_cell(table.rows[i].cells[j], text, bold=bold, size=Pt(9), bg=bg)

# Cold risk highlight
for ri in [1, 2]:
    set_cell(table.rows[ri].cells[7], '中', bold=True, size=Pt(9), color=WHITE, bg=ORANGE_BG)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. MATCH 91: BRAZIL vs NORWAY
# ═══════════════════════════════════════════════════
add_heading_styled('比赛91: BRA 巴西 vs NOR 挪威  (04:00 BJT  |  MetLife Stadium, 新泽西)', level=2)

add_para('16强淘汰赛  |  八强对手: MEX vs ENG 胜者  |  裁判: 伊斯梅尔·埃尔法特(美国)  |  [注意] 高温体感41C  |  身价比: 巴西€909M vs 挪威€601M = 1.51:1',
         size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(6))

# ── Confirmed starting XI ──
add_para('官方确认首发阵容', bold=True, size=Pt(10), color=DARK)
add_para('BRA (4-2-3-1): 阿利松 | 达尼洛[黄] 马尔基尼奥斯(C) 加布里埃尔 道格拉斯·桑托斯 | 卡塞米罗[黄] 吉马良斯 | 拉扬 马丁内利 维尼修斯 | 库尼亚',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('NOR (4-3-3): 尼兰 | 吕尔松 阿耶尔 赫格姆 默勒·沃尔夫 | 桑德·贝格 帕特里克·贝格 厄德高(C) | 索尔洛特 哈兰德 努萨[黄]',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('[注意] 拉菲尼亚在巴西替补席(刚伤愈)! 吕尔松伤愈回归挪威首发!  卡塞米罗确认通过体检.',
         size=Pt(8), color=ACCENT, space_after=Pt(8))

# ── Player Ratings ──
add_para('球员评分', bold=True, size=Pt(10), color=DARK)
pr_data = [
    ['#', '球员', '评分', '位置', '关键标注'],
    ['7', '维尼修斯 (BRA)', '9.0', 'LW', '[MOTM候选] 4球+1助. 身价€180M. 对吕尔松仍有优势'],
    ['9', '哈兰德 (NOR)', '9.5', 'ST', '[MOTM候选] 5球/3场. 世界最佳中锋. 仅需一次机会'],
    ['10', '厄德高 (NOR)', '8.5', 'AM', '创造力核心. 阿森纳队长'],
    ['1', '阿利松 (BRA)', '8.0', 'GK', '世界前三门将. 本届稳定'],
    ['8', '吉马良斯 (BRA)', '7.8', 'DM', '中场发动机. 赛前表示必须盯住哈兰德'],
    ['4', '马尔基尼奥斯 (BRA, C)', '7.8', 'CB', '防线领袖+队长. 对日本多次关键解围'],
    ['26', '吕尔松 (NOR)', '7.5', 'RB', '[注意] 多特蒙德RB伤愈回归! 原预测佩德森6.5'],
    ['22', '马丁内利 (BRA)', '7.5', '#10', '[注意] 非自然位置! 阿森纳边锋顶替伤缺帕奎塔打#10'],
    ['7', '索尔洛特 (NOR)', '7.5', 'RW', '194cm第二高点. 定位球核心威胁'],
    ['5', '卡塞米罗 (BRA)', '7.3', 'DM', '[黄] 34岁+刚伤愈+背牌+高温=四重风险'],
    ['26', '拉扬 (BRA)', '7.0', 'RW', '19岁伯恩茅斯新人. 世界杯淘汰赛首秀'],
    ['13', '达尼洛 (BRA)', '7.0', 'RB', '[黄] 背牌收敛面对努萨速度'],
    ['—', '拉菲尼亚 (BRA)', '—', '替补', '[注意][注意] 伤愈在巴西替补席! 安切洛蒂70分钟王牌!'],
    ['—', '内马尔 (BRA)', '—', '替补', '仅14分钟出场. 状态未知'],
]
prt = doc.add_table(rows=len(pr_data), cols=5)
prt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(pr_data[0]):
    set_cell(prt.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(pr_data[1:], 1):
    for j, text in enumerate(row):
        c = 'left' if j in [1, 4] else 'center'
        bg = LIGHT_GREEN_BG if 'MOTM' in row[4] else (LIGHT_GRAY if i % 2 == 0 else None)
        clr = ACCENT if '四重风险' in row[4] or '非自然位置' in row[4] else None
        if c == 'left':
            set_cell_left(prt.rows[i].cells[j], text, size=Pt(8), color=clr)
        else:
            set_cell(prt.rows[i].cells[j], text, size=Pt(8), color=clr, bg=bg)
for row in prt.rows:
    row.cells[0].width = Inches(0.35)
    row.cells[1].width = Inches(2.2)
    row.cells[2].width = Inches(0.45)
    row.cells[3].width = Inches(0.55)
    row.cells[4].width = Inches(5.0)

doc.add_paragraph()

# ── Factors table ──
add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
ft_data = [
    ['因素', '有利方', '理由'],
    ['哈兰德状态: 5球/3场', 'NOR [火][火][火]', '仅需一次机会. 加布里埃尔+马尔基尼奥斯有英超经验但未曾面对此状态的哈兰德'],
    ['巴西伤病: 帕奎塔缺阵', 'NOR [火][火][火]', '马丁内利打#10是非自然位置 — 巴西攻击维度从4降为2'],
    ['维尼修斯 vs 吕尔松', 'BRA [火][火]', '[降级:原[火][火][火]] 吕尔松(多特蒙德)远强于佩德森. 优势仍在但不再降维打击'],
    ['卡塞米罗确认首发', 'BRA [火][火]', '[升级:原[疑]] 赛前通过体检. 中场控制力增强. 双后腰可限制厄德高'],
    ['拉菲尼亚回归替补!', 'BRA [火][火]', '[新增] 安切洛蒂70分钟王牌! 巴塞罗那边锋伤愈. 原预测标注缺席'],
    ['高温41C体感', 'BRA [火][火]', '[注][修正!原均偏NOR->BRA] 科研:高温削弱防守>进攻(Kyklos 2024+EJSS 2025+Schwarz 2025). 防守压迫/抢断全降 进球+射门转化升. 巴西控球方可休息. CLAUD.md:高温不帮大巴'],
    ['挪威防线: 7失球/4场', 'BRA [火][火]', '对法国1-4暴露防线. 但吕尔松回归改善了右侧防守'],
    ['背牌风险: 3人背黄', '均偏NOR(-)', '卡塞米罗背牌铲球收敛(StatsBomb验证-30%). 达尼洛+努萨同理'],
    ['裁判: 埃尔法特(美国)', 'NOR [火]', '[新增/确认] 美国本土裁判. 欧式尺度允许身体对抗'],
    ['挪威定位球: 三高塔', 'NOR [火]', '[调整] 哈兰德195+索尔洛特194+阿耶尔. 挪威定位球优势被低估! 修正:挪威略优'],
]
ft = doc.add_table(rows=len(ft_data), cols=3)
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(ft_data[0]):
    set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(ft_data[1:], 1):
    set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8.5))
    clr = ACCENT if '[火][火][火]' in row[1] else None
    set_cell(ft.rows[i].cells[1], row[1], size=Pt(8.5), color=clr)
    set_cell_left(ft.rows[i].cells[2], row[2], size=Pt(8.5))
for row in ft.rows:
    row.cells[0].width = Inches(3.0)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(6.0)

doc.add_paragraph()

# ── Strength + Upset ──
add_para('强队分类: 巴西(超级巨星型, €909M) vs 挪威(超级巨星型, €601M). 双方均有>=2名超巨 — 罕见的淘汰赛配置.',
         size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(2))
add_para('冷门风险: 中 (原中高->下调: 卡塞米罗确认+拉菲尼亚回归=巴西利好; 吕尔松回归+高温=挪威利好; 双向抵消)',
         bold=True, size=Pt(9), color=ORANGE_BG, space_after=Pt(6))

# ── Predictions ──
add_para('比分预测', bold=True, size=Pt(10), color=DARK)
pred_data = [
    ['类型', '比分', '半场', '进球者与逻辑'],
    ['[首选]', '巴西 2-1 挪威', '1-0', '维尼修斯34\'内切远射 / 哈兰德58\'厄德高直塞(卡塞米罗60\'被换下窗口) / 库尼亚73\''],
    ['备选1', '2-2 (加时/点球)', '1-1', '厄德高28\' / 维尼修斯41\' / 哈兰德63\'(点) / 拉菲尼亚82\'替补建功'],
    ['备选2', '巴西 1-0 挪威', '0-0', '拉菲尼亚78\'替补即破门. 高温使比赛节奏降->双方保守->替补灵光'],
    ['冷门', '挪威 2-1 巴西', '0-1', '哈兰德22\'定位球+67\' / 维尼修斯55\'. 挪威三高塔定位球+高温拖垮巴西'],
]
pt = doc.add_table(rows=len(pred_data), cols=4)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(pred_data[0]):
    set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(pred_data[1:], 1):
    bg = LIGHT_GREEN_BG if i == 1 else None
    fg = GREEN_BG if i == 1 else (ACCENT if i == 4 else None)
    set_cell(pt.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg, color=fg)
    set_cell(pt.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
    set_cell(pt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
    set_cell_left(pt.rows[i].cells[3], row[3], size=Pt(8))
for row in pt.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(1.4)
    row.cells[2].width = Inches(0.7)
    row.cells[3].width = Inches(7.7)

doc.add_paragraph()

# ── Key battles ──
add_para('关键对位与战术要点', bold=True, size=Pt(10), color=DARK)
battles = [
    ['对位', '详情', '预判'],
    ['维尼修斯 vs 吕尔松', '€180M超巨 vs 多特蒙德RB(刚伤愈). 吕尔松德甲级别远强佩德森. 高温下刚伤愈体能可能仅撑60分钟', '维尼修斯占优但非碾压. 下半场末段空间增大'],
    ['加布里埃尔 vs 哈兰德', '阿森纳队友对决. 英超最熟悉的对手. 安帅称不需要反哈兰德计划 — 真信还是烟雾?', '哈兰德5球/3场状态. 加布里埃尔需全场专注'],
    ['拉扬(19岁) vs 默勒·沃尔夫', '新人右翼vs狼队左后卫. 两人均无淘汰赛经验', '巴西右路最大变数. 若被锁死->巴西变单翼'],
    ['卡塞米罗管理', '34岁+刚伤愈+背牌+高温=四重风险. 安帅大概率60\'用法比尼奥换下', '挪威最佳攻击窗口: 60-75分钟'],
    ['挪威定位球三高塔', '哈兰德195+索尔洛特194+阿耶尔 vs 巴西仅双中卫高点', '挪威最可靠得分路径. 角球/任意球=最大威胁'],
]
bt = doc.add_table(rows=len(battles), cols=3)
bt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(battles[0]):
    set_cell(bt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(battles[1:], 1):
    set_cell(bt.rows[i].cells[0], row[0], bold=True, size=Pt(8.5))
    set_cell_left(bt.rows[i].cells[1], row[1], size=Pt(8.5))
    clr = ACCENT if '最大威胁' in row[2] or '四重风险' in row[1] else None
    set_cell_left(bt.rows[i].cells[2], row[2], size=Pt(8.5), color=clr)
for row in bt.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(5.0)
    row.cells[2].width = Inches(3.7)

doc.add_paragraph()

# ── Injuries ──
add_para('伤病/停赛/背牌 — 官方确认', bold=True, size=Pt(10), color=DARK)
inj_data = [
    ['球队', '球员', '状态'],
    ['BRA', '帕奎塔 (Paqueta)', '[缺席] 大腿拉伤'],
    ['BRA', '拉菲尼亚 (Raphinha)', '[注意] [替补可用] 伤愈在巴西替补席! 安切洛蒂确认可出战不首发'],
    ['BRA', '卡塞米罗 (Casemiro)', '[确认出战] 通过体检. [黄] 背牌'],
    ['BRA', '达尼洛 (Danilo)', '[黄] 背牌'],
    ['NOR', '吕尔松 (Ryerson)', '[确认出战] [注意] 伤愈回归首发! 多特蒙德RB'],
    ['NOR', '佩德森 (Pedersen)', '[注意] 退回替补+呼吸道不适'],
    ['NOR', '努萨 (Nusa)', '[黄] 背牌'],
]
it = doc.add_table(rows=len(inj_data), cols=3)
it.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(inj_data[0]):
    set_cell(it.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(inj_data[1:], 1):
    clr = ACCENT if '[注意]' in row[2] or '[黄]' in row[2] else None
    set_cell(it.rows[i].cells[0], row[0], size=Pt(8.5))
    set_cell_left(it.rows[i].cells[1], row[1], size=Pt(8.5))
    set_cell_left(it.rows[i].cells[2], row[2], size=Pt(8.5), color=clr)
for row in it.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(2.3)
    row.cells[2].width = Inches(7.7)

doc.add_paragraph()

# ── Coach ──
add_para('教练博弈 + 高温管理', bold=True, size=Pt(10), color=DARK)
coach_data = [
    ['场景', '安切洛蒂 (BRA)', '索尔巴肯 (NOR)'],
    ['落后时', '60-70\'上内马尔+拉菲尼亚变4-2-4. 法比尼奥换卡塞米罗(体能+背牌双重管理)', '鲍勃换帕特里克·贝格变4-2-4. 拉森换索尔洛特增加禁区双高'],
    ['领先时', '法比尼奥换卡塞米罗控制节奏. 恩德里克换库尼亚保持反击', '佩德森换索尔洛特变5-4-1([注意]佩德森身体存疑). 奥尔斯内斯换努萨'],
    ['替补改变比赛', '[注意][高] 拉菲尼亚(刚伤愈/10-15分钟足够). 内马尔(技术/14分钟体能未知)', '[中] 鲍勃(创造力)+拉森(194cm支点). 佩德森([注意]身体存疑)'],
    ['高温策略(科研修正)', '[注] 高温削弱防守>进攻(多研究验证). 巴西控球方可在球上休息—体能优势. 防守方(挪威)体力消耗更大—防线60分钟后更松散', '定位球是低能耗得分路径. 挪威必须在上半场取得领先—否则高温累积效应越来越有利于巴西攻击手'],
]
ct = doc.add_table(rows=len(coach_data), cols=3)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(coach_data[0]):
    set_cell(ct.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(coach_data[1:], 1):
    set_cell(ct.rows[i].cells[0], row[0], bold=True, size=Pt(8.5))
    set_cell_left(ct.rows[i].cells[1], row[1], size=Pt(8))
    set_cell_left(ct.rows[i].cells[2], row[2], size=Pt(8))
for row in ct.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(4.7)
    row.cells[2].width = Inches(4.6)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. MATCH 92: MEXICO vs ENGLAND
# ═══════════════════════════════════════════════════
add_heading_styled('比赛92: MEX 墨西哥 vs ENG 英格兰  (08:00 BJT  |  Estadio Azteca, 墨西哥城  海拔2,240m)', level=2)

add_para('16强淘汰赛  |  八强对手: BRA vs NOR 胜者  |  身价比: 墨西哥~€192M vs 英格兰€1.310B = 1:6.8  |  [注意] 高原+墨西哥4场零封',
         size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(6))
add_para('[待更新] 首发尚未官方公布. 以下基于媒体预测(SI/RotoWire/ESPN/Goal.com).',
         size=Pt(8), color=ACCENT, space_after=Pt(8))

# ── Factors ──
add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
ft2_data = [
    ['因素', '有利方', '理由'],
    ['高原主场(2240m): 球飞行+6%, 60分钟后体能断崖', 'MEX [火][火][火]', '英格兰仅4天适应. 墨西哥球员在此出生/训练. 高原是额外球员'],
    ['墨西哥4场零封: 本届仅两支零失球球队', 'MEX [火][火][火]', '兰赫尔+巴斯克斯+蒙特斯防守三角. 但4场对手均非顶级攻击力'],
    ['英格兰伤病: 詹姆斯+匡萨缺席右后卫', 'MEX [火][火]', '斯彭斯淘汰赛首秀. 奎尼奥内斯(3球)对残破右路'],
    ['凯恩状态: 4球, 对刚果双响', 'ENG [火][火][火]', '无论高原如何 — 凯恩总能进球'],
    ['身价比: 1:6.8 (明显差距)', 'ENG [火][火][火]', '3:1-10:1范围 — 正常发挥英格兰显著占优'],
    ['贝林厄姆: 身价>墨西哥全队', 'ENG [火][火]', '单兵能力可在任何防守中找到空间'],
    ['赖斯若缺阵: 腿筋紧', 'MEX [火][火]', '英格兰最关键防守中场 — 若缺阵安德森无法填坑'],
    ['英格兰高原历史: 2-2-2', 'MEX [火]', '在墨西哥城/拉巴斯从未赢过'],
]
ft2 = doc.add_table(rows=len(ft2_data), cols=3)
ft2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(ft2_data[0]):
    set_cell(ft2.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(ft2_data[1:], 1):
    set_cell_left(ft2.rows[i].cells[0], row[0], size=Pt(8.5))
    clr = ACCENT if '[火][火][火]' in row[1] else None
    set_cell(ft2.rows[i].cells[1], row[1], size=Pt(8.5), color=clr)
    set_cell_left(ft2.rows[i].cells[2], row[2], size=Pt(8.5))
for row in ft2.rows:
    row.cells[0].width = Inches(3.2)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(5.8)

doc.add_paragraph()

add_para('强队分类: 英格兰(超级巨星型, €1.310B) vs 墨西哥(非强队方, ~€192M). 身价比1:6.8(明显差距). 高原因素是罕见的场地等效因子.',
         size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(2))
add_para('冷门风险: 中 — 高原+墨西哥4场零封=真防守体系. 但墨西哥从未面对贝林厄姆+凯恩级别的攻击组合.',
         bold=True, size=Pt(9), color=ORANGE_BG, space_after=Pt(6))

# ── Predictions ──
add_para('比分预测', bold=True, size=Pt(10), color=DARK)
pred2_data = [
    ['类型', '比分', '半场', '进球者与逻辑'],
    ['[首选]', '英格兰 2-0 墨西哥', '1-0', '凯恩38\'(点) / 贝林厄姆67\'. 高原上半场体能尚存. 墨西哥进攻能力不足以威胁'],
    ['备选1', '英格兰 1-0 墨西哥', '0-0', '凯恩72\'. 墨西哥大巴死守上半场->英格兰需在55-75分钟体能窗口破局'],
    ['备选2', '英格兰 2-1 墨西哥', '0-1', '奎尼奥内斯22\'反击 / 凯恩58\' 81\'. 墨西哥先破门->英格兰逆转'],
    ['冷门', '1-1 (墨西哥加时/点球)', '0-0', '阿尔瓦拉多77\' / 萨卡90+3\'. 高原拖入加时->指数级放大->墨西哥点球'],
]
pt2 = doc.add_table(rows=len(pred2_data), cols=4)
pt2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(pred2_data[0]):
    set_cell(pt2.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(pred2_data[1:], 1):
    bg = LIGHT_GREEN_BG if i == 1 else None
    fg = GREEN_BG if i == 1 else (ACCENT if i == 4 else None)
    set_cell(pt2.rows[i].cells[0], row[0], bold=(i==1), size=Pt(9), bg=bg, color=fg)
    set_cell(pt2.rows[i].cells[1], row[1], bold=True, size=Pt(10), bg=bg)
    set_cell(pt2.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
    set_cell_left(pt2.rows[i].cells[3], row[3], size=Pt(8))
for row in pt2.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(0.7)
    row.cells[3].width = Inches(7.6)

doc.add_paragraph()

# ── Injuries ──
add_para('伤病/停赛 (墨西哥vs英格兰)', bold=True, size=Pt(10), color=DARK)
inj2_data = [
    ['球队', '球员', '状态'],
    ['ENG', '里斯·詹姆斯 (Reece James)', '[缺席] RB — 小组赛以来缺阵'],
    ['ENG', '贾雷尔·匡萨 (Jarell Quansah)', '[缺席] RB/CB'],
    ['ENG', '德克兰·赖斯 (Declan Rice)', '[疑] 腿筋紧 — 预计出战'],
    ['ENG', '布卡约·萨卡 (Bukayo Saka)', '[疑] 恢复中 — 可能替补'],
    ['MEX', '全队健康', '[可]'],
]
it2 = doc.add_table(rows=len(inj2_data), cols=3)
it2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(inj2_data[0]):
    set_cell(it2.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(inj2_data[1:], 1):
    set_cell(it2.rows[i].cells[0], row[0], size=Pt(8.5))
    set_cell_left(it2.rows[i].cells[1], row[1], size=Pt(8.5))
    set_cell_left(it2.rows[i].cells[2], row[2], size=Pt(8.5))
for row in it2.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(7.2)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 4. CHANGES FROM ORIGINAL PREDICTION
# ═══════════════════════════════════════════════════
add_heading_styled('巴西vs挪威 — 官方首发确认后 vs 原预测变更', level=2)

changes_data = [
    ['#', '变更项目', '原预测', '官方确认', '影响'],
    ['1', '挪威RB', '佩德森 (6.5/弱点)', '吕尔松 (7.5/多特蒙德)', '[注意] 维尼修斯优势[火][火][火]->[火][火]'],
    ['2', '卡塞米罗', '[疑] 肌肉问题', '确认首发! 通过体检', '[升级] 巴西中场控制力增强'],
    ['3', '拉菲尼亚', '[缺席] 腿筋', '[注意] 替补可用! 伤愈', '[新增] 安切洛蒂70分钟王牌!'],
    ['4', '巴西阵型', '4-3-3', '4-2-3-1 (马丁内利#10)', '马丁内利打非自然位置'],
    ['5', '裁判', '未确认', '埃尔法特(美国)', '[新增] 美国本土裁判'],
    ['6', '高温41C', '未提及', '体感41C极端高温', '[注][修正] 原分析低估—科研:高温削弱防守. 巴西控球方受益'],
    ['7', '高温因子方向', '均偏挪威(-)', '巴西[火][火]', '[注] 科研修正. 高温削弱防守组织>进攻. 控球方体能优势(在球上休息)'],
    ['7', '定位球评估', '巴西略优', '挪威略优 (三高塔)', '[调整] 挪威定位球优势被低估'],
    ['8', '佩德森', '首发', '替补+呼吸道不适', '挪威替补RB状态存疑'],
    ['9', '帕奎塔伤情', '大腿二级拉伤', '大腿拉伤(较轻)', '可能提前复出'],
    ['10', '冷门风险', '中高', '中 (下调)', '卡塞米罗+拉菲尼亚=利好. 高温=巴西利好(科研修正). 吕尔松=挪威利好. 整体下调'],
]
ctbl = doc.add_table(rows=len(changes_data), cols=5)
ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(changes_data[0]):
    set_cell(ctbl.rows[0].cells[j], text, bold=True, size=Pt(8), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(changes_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row):
        c = 'left' if j in [1,2,3,4] else 'center'
        clr = ACCENT if '[注意]' in text else None
        if c == 'left':
            set_cell_left(ctbl.rows[i].cells[j], text, size=Pt(7.5), color=clr)
        else:
            set_cell(ctbl.rows[i].cells[j], text, size=Pt(7.5), color=clr, bg=bg)
for row in ctbl.rows:
    row.cells[0].width = Inches(0.25)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(2.0)
    row.cells[3].width = Inches(2.4)
    row.cells[4].width = Inches(4.85)

doc.add_paragraph()

# ── Upset paths ──
add_para('巴西vs挪威 — 冷门路径', bold=True, size=Pt(10), color=ACCENT)
upset_items = [
    '1. 上半场挪威定位球破门 — 厄德高角球/任意球->哈兰德/索尔洛特头球(三高塔优势)',
    '2. 高温41C削弱防守组织(Kyklos 2024+EJSS 2025验证) — 但若挪威上半场定位球领先,高温效应反转(领先方可收缩省体能,落后方需更多跑动攻坚)',
    '3. 卡塞米罗背牌收敛+高温体能早衰 — 60\'被换下 — 挪威在60-75分钟窗口扩大领先',
    '4. 吕尔松限制维尼修斯 — 巴西被逼依赖拉扬(19岁新人)+马丁内利(非自然#10)',
    '5. 巴西依赖内马尔(仅14分钟)+拉菲尼亚(刚伤愈) — 两大替补状态未验证 — 无法扭转',
    '6. 延续巴西对欧洲淘汰赛22年不胜的历史惯性',
]
for item in upset_items:
    add_para(item, size=Pt(8.5), color=RGBColor(0x66,0x66,0x66), space_after=Pt(1))

add_para('最大不确定性: 吕尔松刚伤愈—90分钟高温比赛体能? 拉菲尼亚替补状态? 高温实际影响(赛前训练有空调—无法模拟)',
         size=Pt(8), color=ACCENT, space_after=Pt(8))

add_para('墨西哥vs英格兰 — 冷门路径', bold=True, size=Pt(10), color=ACCENT)
upset2_items = [
    '1. 高原让英格兰体能在60分钟后崩溃 — 传控精度断崖下降',
    '2. 墨西哥大巴守住上半场(4场零封已验证) — 60-75分钟利用奎尼奥内斯反击破门',
    '3. 斯彭斯(淘汰赛首秀)被奎尼奥内斯打爆 — 墨西哥边路是唯一但可靠的攻击通道',
    '4. 凯恩高原表现低于预期(32岁/体能消耗更大)',
    '5. 比赛拖入加时 — 高原影响指数级放大 — 墨西哥点球胜',
]
for item in upset2_items:
    add_para(item, size=Pt(8.5), color=RGBColor(0x66,0x66,0x66), space_after=Pt(1))

add_para('最大不确定性: 墨西哥4场零封的对手(南非/韩国/捷克/厄瓜多尔)无一是顶级攻击力 — 面对贝林厄姆+凯恩是全新测试',
         size=Pt(8), color=ACCENT, space_after=Pt(8))

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: 搜狐体育(官方首发) + AS USA + Sporting News + The Standard + ETtoday(高温) + FIFA官网(裁判)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('框架: CLAUDE.md v17 (身价量化+强队三类+10项必填清单+背牌研究+教训验证)  |  生成工具: python-docx',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('版权声明: 预测仅供分析参考，不构成任何投资或博彩建议。  |  FIFA World Cup 2026 is a trademark of FIFA.',
         size=Pt(6), color=RGBColor(0xCC,0xCC,0xCC), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

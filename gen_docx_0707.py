# -*- coding: utf-8 -*-
"""Generate 2026-07-07 predictions DOCX — Portugal vs Spain + USA vs Belgium.
All lineups verified against player_database_0707.md. C罗 rating corrected per database (€8M, 0 non-penalty goals vs strong teams)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年7月7日_两场预测_v3.docx")

doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0x27, 0xAE, 0x60)
ORANGE_BG  = RGBColor(0xF3, 0x9C, 0x12)
LIGHT_GREEN_BG = RGBColor(0xE8, 0xF8, 0xF5)

def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None):
    set_cell(cell, text, bold=bold, size=size, color=color, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══ TITLE ═══
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('7月7日 16强淘汰赛 预测报告', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('葡萄牙 vs 西班牙  |  美国 vs 比利时', bold=False, size=Pt(11), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(10))
add_para('Round of 16  |  AT&T Stadium (达拉斯) / Lumen Field (西雅图)', bold=False, size=Pt(10), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(18))

add_para('生成: 2026年7月7日 BJT  |  阵容: player_database_0707.md + 媒体交叉验证',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(4))
add_para('数据源: player_database_0707.md + Sporting News + Fox Sports + RotoWire  |  CLAUDE.md v18',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(6))
add_para('[修正] C罗评分8.0->7.0(€8M/对强队0运动战进球) | 西班牙首发修正(波罗+奥尔莫) | 比利时首发修正(梅赫勒+瓦纳肯+德克特拉雷)',
         size=Pt(8), color=ACCENT, align='center', space_after=Pt(20))

# ═══ SUMMARY ═══
add_heading_styled('预测汇总', level=2)
sd = [['#','比赛 (BJT)','身价比','强队分类','首选比分','半场','备选比分','冷门风险'],
      ['93','POR vs ESP (03:00)','1:1.32','双方均\n超巨型','西班牙 2-1\n葡萄牙','1-0','1-1(加时/点球)\n2-0 / 葡萄牙2-1','中高'],
      ['94','USA vs BEL (08:00)','1:1.54','BEL:\n体系型','美国 2-1\n比利时','1-0','1-1(加时)\n2-0 / 比利时2-1','中']]
table = doc.add_table(rows=3, cols=8)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cw = [0.28,1.85,0.55,0.72,1.05,0.50,1.30,0.50]
for i,w in enumerate(cw):
    for row in table.rows: row.cells[i].width = Inches(w)
for j,t in enumerate(sd[0]): set_cell(table.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,rd in enumerate(sd[1:],1):
    bg = LIGHT_GRAY if i%2==0 else None
    for j,t in enumerate(rd):
        a='left' if j==1 else 'center'
        if j==1: set_cell_left(table.rows[i].cells[j],t,size=Pt(9))
        else: set_cell(table.rows[i].cells[j],t,bold=(j==4),size=Pt(9),bg=bg)
for ri in [1,2]: set_cell(table.rows[ri].cells[7],sd[ri][7],bold=True,size=Pt(9),color=WHITE,bg=ORANGE_BG)
doc.add_paragraph()

# ═══ POR vs ESP ═══
add_heading_styled('比赛93: POR 葡萄牙 vs ESP 西班牙  (03:00 BJT | AT&T Stadium, 达拉斯 | 室内)', level=2)
add_para('16强 | QF: USA/BEL胜者 | FIFA #7 vs #3 | €957M vs €1.267B (1:1.32) | ESP 4场零封(12-0) | POR 防守存疑(4场失4球)',
         size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D))

add_para('预测首发 (player_database_0707.md + 媒体交叉)', bold=True, size=Pt(10), color=DARK)
add_para('POR (4-2-3-1): 迪奥戈·科斯塔 | 坎塞洛 鲁本·迪亚斯[黄] 雷纳托·韦加 努诺·门德斯 | 维蒂尼亚 若昂·内维斯 | 佩德罗·内托 B费 拉斐尔·莱昂 | C罗',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('ESP (4-3-3): 乌奈·西蒙 | 佩德罗·波罗 库巴西 拉波尔特 库库雷利亚 | 罗德里 佩德里 奥尔莫 | 亚马尔 奥亚萨瓦尔 巴埃纳',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('[修正] ESP首发修正: RB波罗(热刺€50M,对奥地利进球!)替代略伦特 | CM奥尔莫(巴萨€50M)替代梅里诺. POR LW莱昂(€90M,对克罗地亚助攻)修正为费利克斯.',
         size=Pt(7), color=ACCENT, space_after=Pt(6))

# Player Ratings
add_para('核心球员评分 (来源: player_database_0707.md)', bold=True, size=Pt(10), color=DARK)
pr = [['#','球员','评分','位置','关键标注'],
      ['16','罗德里 (ESP)','9.0','DM','[MOTM候选] 金球奖得主. 曼城€130M'],
      ['19','拉明·亚马尔 (ESP)','9.0','RW','[MOTM候选] 18岁€200M. 本届最佳边锋'],
      ['21','奥亚萨瓦尔 (ESP)','8.5','ST','4球. 近16场首发17球. 皇社€45M'],
      ['20','佩德里 (ESP)','8.5','CM','巴萨核心€80M'],
      ['8','B费 (POR)','8.5','AM','[关键] 曼联€85M. 创造力中枢'],
      ['4','鲁本·迪亚斯 (POR)','8.5','CB','[注意][黄] 曼城€80M. 背牌收敛30%(StatsBomb)'],
      ['2','佩德罗·波罗 (ESP)','8.0','RB','[修正] 热刺€50M. 对奥地利头球破门!'],
      ['23','乌奈·西蒙 (ESP)','8.0','GK','519分钟不失球纪录. 毕尔巴鄂€25M'],
      ['17','拉斐尔·莱昂 (POR)','8.0','LW','[修正] AC米兰€90M. 对克罗地亚助攻. 防守存疑但攻击顶级'],
      ['10','达尼·奥尔莫 (ESP)','8.0','AM','[修正] 巴萨€50M. 7/4训练减量但可出战'],
      ['7','C罗 (POR)[C]','7.0','ST','[修正] 41岁€8M(沙特). 对弱队2球(运动战)+克罗地亚点球. 对强队0运动战进球'],
      ['9','贡萨洛·拉莫斯 (POR)','7.5','ST','[超级替补] PSG€50M. 对克罗地亚90+4绝杀'],
      ['—','尼科·威廉斯 (ESP)','—','LW','[缺席] 内收肌—西班牙左路爆破减半']]
prt = doc.add_table(rows=len(pr), cols=5); prt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(pr[0]): set_cell(prt.rows[0].cells[j],t,bold=True,size=Pt(8),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(pr[1:],1):
    for j,t in enumerate(row):
        c='left' if j in[1,4] else 'center'
        bg=LIGHT_GREEN_BG if 'MOTM候选' in row[4] else (LIGHT_GRAY if i%2==0 else None)
        clr=ACCENT if '[注意]' in row[4] or '[修正]' in row[4] else None
        if c=='left': set_cell_left(prt.rows[i].cells[j],t,size=Pt(8),color=clr)
        else: set_cell(prt.rows[i].cells[j],t,size=Pt(8),color=clr,bg=bg)
for row in prt.rows:
    row.cells[0].width=Inches(0.35); row.cells[1].width=Inches(2.2)
    row.cells[2].width=Inches(0.45); row.cells[3].width=Inches(0.7); row.cells[4].width=Inches(4.8)
doc.add_paragraph()

# Factors
add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
ft = [['因素','有利方','理由'],
      ['西班牙体系完整: 4场12-0(零封)','ESP ★★','铁律12降级: 零封对手均非超巨型. 从未面对C罗+B费+莱昂'],
      ['尼科·威廉斯缺阵: 左路爆破减半','POR ★★','巴埃纳(€50M)远不如威廉斯. 坎塞洛防守压力骤减'],
      ['C罗+B费+莱昂: 3核心攻击线','POR ★★★','莱昂€90M(对克罗地亚助攻)+B费€85M. C罗€8M对强队0运动战进球—非超巨'],
      ['C罗对强队0运动战进球: 41岁€8M','ESP ★★','[修正] 仅对乌兹别克斯坦运动战破门. 面对4场零封防线=运动战破门概率极低. 威胁限于点球+定位球'],
      ['葡萄牙防守不稳: 4场失4球','ESP ★★','对刚果金+哥伦比亚均失球. 面对西班牙12-0控制力=最大考验'],
      ['亚马尔 vs 努诺·门德斯: €200M vs PSG','ESP ★★','西班牙最强对位. 亚马尔可切可传'],
      ['罗德里 vs 维蒂尼亚: 金球奖vs欧冠冠军','ESP ★★','罗德里世界第一. 维蒂尼亚+内维斯=PSG欧冠级不弱'],
      ['室内球场: 气候受控','均势','纯实力对决. 西班牙体系优势被放大'],
      ['2025欧国联决赛: 葡萄牙点球胜','POR ★','最近交手心理优势']]
ftt = doc.add_table(rows=len(ft), cols=3); ftt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(ft[0]): set_cell(ftt.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(ft[1:],1):
    set_cell_left(ftt.rows[i].cells[0],row[0],size=Pt(8.5))
    clr=ACCENT if '★★★' in row[1] else None
    set_cell(ftt.rows[i].cells[1],row[1],size=Pt(8.5),color=clr)
    set_cell_left(ftt.rows[i].cells[2],row[2],size=Pt(8.5))
for row in ftt.rows:
    row.cells[0].width=Inches(3.2); row.cells[1].width=Inches(1.5); row.cells[2].width=Inches(5.8)
doc.add_paragraph()

add_para('强队分类: 葡萄牙(超级巨星型, €957M: B费€85M+莱昂€90M+迪亚斯€80M) vs 西班牙(超级巨星型, €1.267B: 亚马尔€200M+罗德里€130M+佩德里€80M)',
         size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(2))
add_para('冷门风险: 中高 | ESP 4场零封含金量待检验(铁律12) | POR超巨破局能力 | 威廉斯缺阵=左路爆破减半',
         bold=True, size=Pt(9), color=ORANGE_BG, space_after=Pt(6))

add_para('比分预测', bold=True, size=Pt(10), color=DARK)
pd = [['类型','比分','半场','进球者与逻辑'],
      ['[首选]','西班牙 2-1 葡萄牙','1-0','奥亚萨瓦尔31\' / C罗58\'(点) / 佩德里73\'. ESP体系控场. POR唯一进球路径=点球+定位球'],
      ['备选1','1-1 (POR加时/点球)','0-0','C罗67\'(点) / 亚马尔82\'. 双方谨慎→POR加时深度(莱昂+拉莫斯+B席)'],
      ['备选2','西班牙 2-0 葡萄牙','1-0','亚马尔24\' / 奥亚萨瓦尔55\'. 零封延续至6场'],
      ['冷门','葡萄牙 2-1 西班牙','0-1','亚马尔18\' / C罗44\'(点) / 莱昂79\'(替补绝杀)']]
ptt = doc.add_table(rows=5, cols=4); ptt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(pd[0]): set_cell(ptt.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(pd[1:],1):
    bg=LIGHT_GREEN_BG if i==1 else None; fg=GREEN_BG if i==1 else (ACCENT if i==4 else None)
    set_cell(ptt.rows[i].cells[0],row[0],bold=(i==1),size=Pt(9),bg=bg,color=fg)
    set_cell(ptt.rows[i].cells[1],row[1],bold=True,size=Pt(10),bg=bg)
    set_cell(ptt.rows[i].cells[2],row[2],size=Pt(9),bg=bg)
    set_cell_left(ptt.rows[i].cells[3],row[3],size=Pt(8))
for row in ptt.rows:
    row.cells[0].width=Inches(0.7); row.cells[1].width=Inches(1.4)
    row.cells[2].width=Inches(0.7); row.cells[3].width=Inches(7.7)
doc.add_paragraph()

# Key battles
add_para('关键对位', bold=True, size=Pt(10), color=DARK)
bt_data = [['对位','详情','预判'],
           ['库巴西(18) vs C罗(41)','18岁巴萨天才 vs 41岁6届传奇. 经验差23年','C罗运动战难破门. 禁区经验可制造点球'],
           ['亚马尔 vs 努诺·门德斯','€200M超巨 vs PSG左后卫€65M','ESP最大优势. 亚马尔传中=奥亚萨瓦尔终结'],
           ['罗德里 vs B费','金球奖后腰 vs 创造力核心. 若B费被锁=POR攻击瘫痪','罗德里大概率限制B费'],
           ['波罗 vs 莱昂','热刺RB€50M(对奥地利进球) vs AC米兰€90M','莱昂速度优势. 波罗进攻强于防守']]
bt = doc.add_table(rows=5, cols=3); bt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(bt_data[0]): set_cell(bt.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(bt_data[1:],1):
    set_cell(bt.rows[i].cells[0],row[0],bold=True,size=Pt(8.5))
    set_cell_left(bt.rows[i].cells[1],row[1],size=Pt(8.5))
    set_cell_left(bt.rows[i].cells[2],row[2],size=Pt(8.5))
for row in bt.rows:
    row.cells[0].width=Inches(2.0); row.cells[1].width=Inches(4.5); row.cells[2].width=Inches(4.0)
doc.add_paragraph()

# ═══ USA vs BEL ═══
add_heading_styled('比赛94: USA 美国 vs BEL 比利时  (08:00 BJT | Lumen Field, 西雅图 | 68,740人)', level=2)
add_para('16强 | QF: POR/ESP胜者 | FIFA #11 vs #5 | ~€345M vs €530M (1:1.54) | [注意] 巴洛贡红牌撤销复出 | 多库vs里姆(37)致命对位',
         size=Pt(9), color=RGBColor(0x7F,0x8C,0x8D))

add_para('预测首发 (player_database_0707.md + 媒体交叉)', bold=True, size=Pt(10), color=DARK)
add_para('USA (4-2-3-1): 弗里兹 | 弗里曼 理查兹 里姆(37) 罗宾逊 | 亚当斯 蒂尔曼 | 德斯特 麦肯尼 普利西奇(C) | 巴洛贡',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('BEL (4-2-3-1): 库尔图瓦 | 卡斯塔涅 梅赫勒 泰特 德克伊珀 | 蒂勒曼斯(C) 瓦纳肯 | 多库 德布劳内(C) 特罗萨尔 | 德克特拉雷',
         size=Pt(9), color=DARK, space_after=Pt(2))
add_para('[修正] BEL首发修正: CB梅赫勒(€5M/布鲁日)替代法斯 | CM瓦纳肯(€6M/工兵)替代奥纳纳 | CF德克特拉雷(€40M)替代卢卡库. 卢卡库=超级替补(86\'救命球). 德布劳内=那不勒斯(非曼城)/0助攻/未踢满90\'. 多库本届0球0助.',
         size=Pt(7), color=ACCENT, space_after=Pt(6))

# Player Ratings
add_para('核心球员评分 (来源: player_database_0707.md)', bold=True, size=Pt(10), color=DARK)
pr2 = [['#','球员','评分','位置','关键标注'],
       ['1','库尔图瓦 (BEL)','9.0','GK','[MOTM候选] 皇马€25M. 2014在此对美国15扑'],
       ['10','普利西奇 (USA)[C]','8.5','LW','[关键] AC米兰€50M. 美国队长'],
       ['9','巴洛贡 (USA)','8.5','ST','[注意] 3球/4场. 红牌撤销复出! 摩纳哥€50M. 状态待验证'],
       ['10','多库 (BEL)','8.5','RW','[MOTM候选] 曼城€60M. [注意]本届0球0助/R32被56\'换下. 但对里姆速度碾压'],
       ['7','德布劳内 (BEL)[C]','7.5','AM','[修正] 那不勒斯€55M. 35岁/0助攻/未踢满90\'/R32被56\'换下后球队逆转'],
       ['5','罗宾逊 (USA)','8.0','LB','富勒姆€30M. 本届最佳左后卫之一'],
       ['6','麦肯尼 (USA)','8.0','AM','尤文€25M. 对抗+跑动覆盖巨大'],
       ['8','蒂勒曼斯 (BEL)[C]','8.5','CM','[MOTM候选] 维拉€35M. R32英雄:89\'扳平+120+5\'点球'],
       ['11','特罗萨尔 (BEL)','8.0','LW','[关键] 阿森纳€45M. 2球1助—常规攻击最强点'],
       ['4','梅赫勒 (BEL)','6.5','CB','[弱点][修正] 布鲁日€5M. 对塞内加尔被打穿—速度不足'],
       ['13','里姆 (USA)','7.0','CB','[注意] 37岁€1M. 速度vs多库=致命对位. 亚当斯须协防'],
       ['9','卢卡库 (BEL)','7.5','ST','[超级替补] 罗马€20M. 86\'救命球. 仅适合30分钟冲刺'],
       ['17','德克特拉雷 (BEL)','7.0','ST','[修正] 亚特兰大€40M. 首发CF—R32无进球被换下']]
prt2 = doc.add_table(rows=len(pr2), cols=5); prt2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(pr2[0]): set_cell(prt2.rows[0].cells[j],t,bold=True,size=Pt(8),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(pr2[1:],1):
    for j,t in enumerate(row):
        c='left' if j in[1,4] else 'center'
        bg=LIGHT_GREEN_BG if 'MOTM候选' in row[4] else (LIGHT_GRAY if i%2==0 else None)
        clr=ACCENT if '[注意]' in row[4] or '[弱点]' in row[4] or '[修正]' in row[4] else None
        if c=='left': set_cell_left(prt2.rows[i].cells[j],t,size=Pt(8),color=clr)
        else: set_cell(prt2.rows[i].cells[j],t,size=Pt(8),color=clr,bg=bg)
for row in prt2.rows:
    row.cells[0].width=Inches(0.35); row.cells[1].width=Inches(2.2)
    row.cells[2].width=Inches(0.45); row.cells[3].width=Inches(0.7); row.cells[4].width=Inches(4.8)
doc.add_paragraph()

# Factors
add_para('因素导向表', bold=True, size=Pt(10), color=DARK)
ft2d = [['因素','有利方','理由'],
        ['美国主场: 68,740人满座','USA ★★★','本届主场全胜. 波切蒂诺:"美国足球史上最大比赛"'],
        ['巴洛贡复出: 红牌撤销, 3球射手回归','USA ★★★','[注意] 攻击力恢复双核. 但停赛期训练状态待验证(铁律11)'],
        ['多库 vs 里姆(37): 速度碾压','BEL ★★★','最致命对位. 多库(€60M)本届0球0助但速度优势真实'],
        ['库尔图瓦: 世界前三门将','BEL ★★★','淘汰赛门将模式. 但对塞内加尔失2球'],
        ['德布劳内状态: 35岁/0助攻/未踢满90\'','USA ★★','R32被56\'换下后球队逆转. 亚当斯可针对性锁死'],
        ['比利时防线速度慢: 梅赫勒被打穿','USA ★★','梅赫勒(€5M)对塞内加尔被爆. 巴洛贡+普利西奇速度可碾压'],
        ['美国Pochettino高位逼抢','USA ★★','比老龄后场出球=噩梦. 逼抢→失误→反击'],
        ['比利时攻击便秘: R32前86分钟0射正','USA ★★','对伊朗10人0-0/对埃及1-1. 有组织防守时进攻乏力']]
ftt2 = doc.add_table(rows=len(ft2d), cols=3); ftt2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(ft2d[0]): set_cell(ftt2.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(ft2d[1:],1):
    set_cell_left(ftt2.rows[i].cells[0],row[0],size=Pt(8.5))
    clr=ACCENT if '★★★' in row[1] else None
    set_cell(ftt2.rows[i].cells[1],row[1],size=Pt(8.5),color=clr)
    set_cell_left(ftt2.rows[i].cells[2],row[2],size=Pt(8.5))
for row in ftt2.rows:
    row.cells[0].width=Inches(3.2); row.cells[1].width=Inches(1.5); row.cells[2].width=Inches(5.8)
doc.add_paragraph()

add_para('强队分类: 比利时=体系型(德布劳内€55M未达超巨阈值+0助攻. 多库€60M核心级0球0助. 对塞内加尔86分钟0射正). 美国=非强队方(主场等效~€345M)',
         size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(2))
add_para('冷门风险: 中 | 身价比1:1.54=五五开 | 比利时老龄防线vs美国年轻锋线 | 库尔图瓦+德布劳内可单场决定 | 多库vs里姆致命',
         bold=True, size=Pt(9), color=ORANGE_BG, space_after=Pt(6))

add_para('比分预测', bold=True, size=Pt(10), color=DARK)
pd2 = [['类型','比分','半场','进球者与逻辑'],
       ['[首选]','美国 2-1 比利时','1-0','巴洛贡27\'(普利西奇助攻) / 多库52\' / 麦肯尼68\'. 主场+高位逼抢vs老龄防线'],
       ['备选1','1-1 (加时/点球)','0-0','卢卡库72\'(替补) / 巴洛贡85\'. 双方谨慎. 加时美国深度稍优'],
       ['备选2','美国 2-0 比利时','1-0','普利西奇34\' / 巴洛贡61\'. 比利时防线被逼抢撕碎'],
       ['冷门','比利时 2-1 美国','0-1','德布劳内38\' / 卢卡库55\' / 雷纳90+2\'. 库尔图瓦5+神扑+多库爆里姆']]
ptt2 = doc.add_table(rows=5, cols=4); ptt2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(pd2[0]): set_cell(ptt2.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(pd2[1:],1):
    bg=LIGHT_GREEN_BG if i==1 else None; fg=GREEN_BG if i==1 else (ACCENT if i==4 else None)
    set_cell(ptt2.rows[i].cells[0],row[0],bold=(i==1),size=Pt(9),bg=bg,color=fg)
    set_cell(ptt2.rows[i].cells[1],row[1],bold=True,size=Pt(10),bg=bg)
    set_cell(ptt2.rows[i].cells[2],row[2],size=Pt(9),bg=bg)
    set_cell_left(ptt2.rows[i].cells[3],row[3],size=Pt(8))
for row in ptt2.rows:
    row.cells[0].width=Inches(0.7); row.cells[1].width=Inches(1.4)
    row.cells[2].width=Inches(0.7); row.cells[3].width=Inches(7.7)
doc.add_paragraph()

# Key battles
add_para('关键对位', bold=True, size=Pt(10), color=DARK)
bt2d = [['对位','详情','预判'],
        ['多库 vs 里姆(37)','BEL最强对位. 多库速度+盘带 vs 37岁老将. 亚当斯须向右协防','多库大概率爆点. 美国需双人包夹'],
        ['巴洛贡 vs 梅赫勒','3球射手(刚复出) vs €5M被打穿过中卫. 巴洛贡逼抢+速度=噩梦','美国高位逼抢下梅赫勒出球失误率高'],
        ['亚当斯 vs 德布劳内','伯恩茅斯DM vs 35岁那不勒斯0助攻. 若亚当斯协防多库=德布劳内获空间','德布劳内本届低迷但一脚传球仍致命'],
        ['库尔图瓦 vs 美国攻击群','世界级门将 vs 巴洛贡+普利西奇+麦肯尼. 2014对美国15扑','库尔图瓦会神扑但无法阻止所有']]
bt2 = doc.add_table(rows=5, cols=3); bt2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j,t in enumerate(bt2d[0]): set_cell(bt2.rows[0].cells[j],t,bold=True,size=Pt(9),color=WHITE,bg=HEADER_BG)
for i,row in enumerate(bt2d[1:],1):
    set_cell(bt2.rows[i].cells[0],row[0],bold=True,size=Pt(8.5))
    set_cell_left(bt2.rows[i].cells[1],row[1],size=Pt(8.5))
    set_cell_left(bt2.rows[i].cells[2],row[2],size=Pt(8.5))
for row in bt2.rows:
    row.cells[0].width=Inches(2.0); row.cells[1].width=Inches(4.5); row.cells[2].width=Inches(4.0)
doc.add_paragraph()

# Risk
add_heading_styled('修正说明与风险提示', level=2)
add_para('C罗评分修正 (v18): C罗41岁, €8M身价(沙特联赛). 本届: 5-0乌兹别克斯坦2球(运动战/弱队) + 对克罗地亚点球. 对强队0运动战进球. 评分从8.0/MOTM候选→7.0. 主要威胁仅限点球+定位球头球, 非运动战破局者.',
         size=Pt(8.5), color=ACCENT, space_after=Pt(4))
add_para('西班牙首发修正: RB波罗(€50M热刺, 对奥地利进球)替代略伦特. AM奥尔莫(€50M巴萨)替代梅里诺. 来源: player_database_0707.md — R32实际首发.',
         size=Pt(8.5), color=ACCENT, space_after=Pt(4))
add_para('比利时首发修正: CB梅赫勒(€5M)替代法斯. CM瓦纳肯(€6M工兵)替代奥纳纳. CF德克特拉雷(€40M)替代卢卡库. 卢卡库=超级替补(86\'救命球). 德布劳内=那不勒斯/0助攻. 多库=本届0球0助. 来源: player_database_0707.md.',
         size=Pt(8.5), color=ACCENT, space_after=Pt(4))
add_para('系统综述交叉验证 (Illmer & Daumann 2022, 21篇): 降雨无效应[已确认] | 海拔降身体不降技术[已确认] | 温度需条件分支[已确认] | 风影响进攻[待纳入]',
         size=Pt(8.5), color=RGBColor(0x55,0x55,0x55), space_after=Pt(6))

# Footer
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: player_database_0707.md + Sporting News + Fox Sports + RotoWire + Illmer & Daumann (2022)',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('框架: CLAUDE.md v18 (铁律10-14) | 生成: python-docx | 分析仅供参考',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('版权声明: 预测仅供分析参考，不构成任何投资或博彩建议。 | FIFA World Cup 2026 is a trademark of FIFA.',
         size=Pt(6), color=RGBColor(0xCC,0xCC,0xCC), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')

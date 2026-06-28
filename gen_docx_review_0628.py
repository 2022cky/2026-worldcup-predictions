# -*- coding: utf-8 -*-
"""Generate 6月28日 六场复盘 as .docx — 沿用 gen_docx.py 预测v7同款样式."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "2026年6月28日_6场复盘.docx")

doc = Document()

# ── Page setup (同预测v7) ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

# ── Styles (同预测v7) ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

DARK  = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT= RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_BG   = RGBColor(0x27, 0xAE, 0x60)

# ─── helpers (同预测v7) ───
def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{str(bg)[2:]}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def set_cell_left(cell, text, bold=False, size=Pt(9), color=None, bg=None):
    set_cell(cell, text, bold=bold, size=size, color=color, bg=bg, align='left')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=Pt(10), color=None, align='left', space_after=Pt(4)):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_para('2026 FIFA 世界杯', bold=True, size=Pt(28), color=ACCENT, align='center', space_after=Pt(2))
add_para('6月28日 六场全复盘', bold=True, size=Pt(18), color=DARK, align='center', space_after=Pt(6))
add_para('小组赛收官日  ·  L / K / J 组  ·  第3轮', bold=False, size=Pt(11), color=RGBColor(0x7F,0x8C,0x8D), align='center', space_after=Pt(18))
add_para('生成时间: 2026年6月28日 北京时间 14:00  |  数据源: FIFA API + ESPN API  |  框架: CLAUDE.md v17',
         size=Pt(8), color=RGBColor(0x95,0xA5,0xA6), align='center', space_after=Pt(20))

# ═══════════════════════════════════════════════════
# 1. SUMMARY TABLE
# ═══════════════════════════════════════════════════
add_heading_styled('复盘成绩总览', level=2)

summary_data = [
    ['#', '时间', '比赛', '预测首选', '实际比分', '方向', '精确比分', '复盘'],
    ['1', '05:00', '克罗地亚 vs 加纳', '1-1', '2-1', '[错]', '[错]', '首选平局→克罗地亚赢'],
    ['2', '05:00', '巴拿马 vs 英格兰', '0-3', '0-2', '[对]', '[对]备选', '备选2-0命中!'],
    ['3', '07:30', '哥伦比亚 vs 葡萄牙', '1-1', '0-0', '[对]', '[错]', '方向对(平局)、比分0-0非首选'],
    ['4', '07:30', '刚果金 vs 乌兹别克斯坦', '刚果金 2-0', '3-1', '[对]', '[错]', '方向对、过程全反'],
    ['5', '10:00', '阿尔及利亚 vs 奥地利', '0-0', '3-3', '[错]', '[错]', '[火]默契球预判全错!'],
    ['6', '10:00', '约旦 vs 阿根廷', '0-3', '1-3', '[对]', '[错]', '方向对、约旦进球意外'],
]

table = doc.add_table(rows=len(summary_data), cols=len(summary_data[0]))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [0.28, 0.45, 1.80, 0.85, 0.72, 0.42, 0.52, 1.80]
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = Inches(w)

for j, text in enumerate(summary_data[0]):
    set_cell(table.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)

for i, row_data in enumerate(summary_data[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    for j, text in enumerate(row_data):
        if j == 2:
            set_cell_left(table.rows[i].cells[j], text, size=Pt(8))
        elif j == 7:
            set_cell_left(table.rows[i].cells[j], text, size=Pt(8))
        elif j in [5, 6]:
            clr = GREEN_BG if '[对]' in text else ACCENT
            set_cell(table.rows[i].cells[j], text, bold=True, size=Pt(8), color=clr, bg=bg)
        elif j == 4:
            set_cell(table.rows[i].cells[j], text, bold=True, size=Pt(9), bg=bg)
        elif j == 3:
            set_cell(table.rows[i].cells[j], text, bold=True, size=Pt(8), bg=bg)
        else:
            set_cell(table.rows[i].cells[j], text, size=Pt(8), bg=bg)

doc.add_paragraph()
add_para('汇总: 方向 4/6 (67%)  |  精确比分 1/6 (17%, 含备选)  |  半场正确 2/6 (33%)',
         size=Pt(10), bold=True, color=DARK, align='center', space_after=Pt(12))

# ═══════════════════════════════════════════════════
# 2. MATCH-BY-MATCH REVIEW
# ═══════════════════════════════════════════════════

matches = [
    {
        'num': 1, 'title': '克罗地亚 2-1 加纳 (L组)',
        'goals': [('31\'', '苏契奇 (Petar SUCIC)', '克罗地亚'), ('73\'', '卢卡森 (Derrick LUCKASSEN)', '加纳'), ('83\'', '弗拉希奇 (Nikola VLASIC)', '克罗地亚')],
        'events': '半场 克罗地亚 1-0 加纳 | 黄牌: 佩里希奇 68\'(克), 奥庞 90+4\'(加) | 换人: 克罗地亚4换(格瓦迪奥尔88\'上), 加纳5换',
        'pred_vs_actual': [
            ('首选比分', '1-1', '2-1', '[错]'),
            ('方向', '平局', '克罗地亚赢', '[错]'),
            ('半场', '0-0', '1-0', '[错]'),
        ],
        'factors': [
            ['格瓦迪奥尔替补→防守削弱', '仅丢1球,三后卫未崩溃', '[半对]'],
            ['加纳威廉姆斯+耶伦基替补→反击减半', '加纳仅1球(后卫定位球),无反击进球', '[对]'],
            ['莫德里奇定位球是唯一破局手段', '两球均来自中前场运动战配合', '[半对]'],
            ['1-1首选逻辑', '结果偏克罗地亚——替补深度差异更大', '[错]'],
        ],
        'lessons': [
            '核心缺席不对称: 克罗地亚替代者(苏契奇/弗拉希奇)是首发级别,加纳替代者不是。替补深度差 > 缺席人数对称。',
            '弗拉希奇隐藏胜负手: 赛前仅标注"中场"的#13球员,83\'绝杀。首发中靠前的中场值得专门评估。',
            '大巴持久力看替补防守深度: 首发大巴撑70分钟,替补防守下降导致最后阶段丢球。',
        ],
    },
    {
        'num': 2, 'title': '巴拿马 0-2 英格兰 (L组)',
        'goals': [('62\'', '贝林厄姆 (Jude BELLINGHAM)', '英格兰'), ('67\'', '凯恩 (Harry KANE)', '英格兰')],
        'events': '半场 巴拿马 0-0 英格兰 | 黄牌: 法哈多 53\'(巴), 昆沙 60\'(英), 安德拉德 83\'(巴) | 换人: 英格兰5换(凯恩84\'下), 巴拿马5换',
        'pred_vs_actual': [
            ('首选比分', '0-3', '0-2', '[错]'),
            ('备选比分', '0-2', '0-2', '[对]命中!'),
            ('方向', '英格兰赢', '英格兰赢', '[对]'),
            ('半场', '0-1', '0-0', '[错]'),
        ],
        'factors': [
            ['凯恩+萨卡+拉什福德→要进球', '贝林厄姆先破僵,凯恩跟进', '[对]'],
            ['轮换配合生疏→半场1-0', '上半场0-0!轮换影响大于预期', '[错]'],
            ['身价比碾压→英格兰必赢', '方向正确,赢球幅度低于预期', '[对]'],
        ],
        'lessons': [
            '轮换中场控制力下降真实: 安德森+罗杰斯上半场无法穿透巴拿马大巴。换上斯彭斯/马杜埃克后才打开。',
            '贝林厄姆中场定海神针: 轮换阵容中不可替代,62\'打破僵局。',
            '大幅轮换→上半场0-0: 不管对手多弱,全新中场+防线必然生涩。以后默认上半场进球困难。',
        ],
    },
    {
        'num': 3, 'title': '哥伦比亚 0-0 葡萄牙 (K组)',
        'goals': [],
        'events': '比分 哥伦比亚 0-0 葡萄牙 | 半场 0-0 | 黄牌: 普埃尔塔 86\'(哥) | 换人: 哥伦比亚5换(哈梅斯76\'下),葡萄牙5换(莱昂70\'上)',
        'pred_vs_actual': [
            ('首选比分', '1-1', '0-0', '[错]'),
            ('方向', '平局', '平局', '[对]'),
            ('半场', '0-0', '0-0', '[对]'),
        ],
        'factors': [
            ['哥伦比亚打平即头名→不冒险', '全场0-0,射门均不多', '[对]'],
            ['双方均已出线→平局双赢', '完全验证——典型默契平局', '[对]'],
            ['C罗可能被保护性使用', '打满全场但未进球,强度不足', '[半对]'],
        ],
        'lessons': [
            '"双方均可接受平局"判断准确: 识别出双方已出线不想冒险。但高估了进攻意愿——1-1非0-0。',
            '第3轮"已出线"比赛0-0概率翻倍: 历史基线9.4%在此类比赛可能×2。双方无必须进球理由。',
            '超巨星在"无所谓"的比赛中不会强行破局: 葡萄牙超巨型但未进球——比赛强度不够。',
        ],
    },
    {
        'num': 4, 'title': '刚果金 3-1 乌兹别克斯坦 (K组)',
        'goals': [('10\'', '绍穆罗多夫 (SHOMURODOV)', '乌兹别克'), ('68\'', '维萨 (Yoane WISSA)', '刚果金'), ('78\'', '马耶莱 (Fiston MAYELE)', '刚果金'), ('90+1\'', '维萨 (Yoane WISSA)', '刚果金')],
        'events': '半场 刚果金 0-1 乌兹别克 | 黄牌: 萨迪基21\'\姆布库45+4\'\穆图萨米62\'(刚) | 换人: 马耶莱51\'换入→78\'反超!',
        'pred_vs_actual': [
            ('首选比分', '刚果金 2-0', '3-1', '[错]'),
            ('方向', '刚果金赢', '刚果金赢', '[对]'),
            ('零封', '预测零封', '10\'即丢球!', '[错]'),
        ],
        'factors': [
            ['刚果金防守已验证', '10\'先丢球!开场松懈', '[错]'],
            ['乌兹别克已淘汰无威胁', '10\'先拔头筹!无压力反而放开', '[错]'],
            ['刚果金赢球动力明确', '下半场连进3球逆转——动力兑现', '[对]'],
        ],
        'lessons': [
            '"已淘汰球队"≠"无威胁": 乌兹别克10\'进球是响亮警告。无压力下可能超常发挥。',
            '无淘汰球队下半场崩盘真实: 60\'后体能/专注度下降——刚果金逆转窗口(68\'/78\'/90+1\')。',
            '上半场全错,全场正确: 足球逻辑有时延迟兑现。马耶莱替补奇兵(51\'→78\')。',
        ],
    },
    {
        'num': 5, 'title': '阿尔及利亚 3-3 奥地利 (J组) [火] 本日最具戏剧性',
        'goals': [('28\'', '阿瑙托维奇 (ARNAUTOVIC)', '奥地利'), ('45\'', '贝尔加利 (BELGHALI)', '阿尔及利亚'), ('55\'', '萨比策 (SABITZER)', '奥地利'), ('60\'', '马赫雷斯 (MAHREZ)', '阿尔及利亚'), ('90+3\'', '马赫雷斯 (MAHREZ)', '阿尔及利亚'), ('90+6\'', '卡拉季奇 (KALAJDZIC)', '奥地利')],
        'events': '半场 1-1 | 黄牌: 阿瑙托维奇 11\'(奥) | 换人: 卡拉季奇90+5\'入→90+6\'头球绝平!',
        'pred_vs_actual': [
            ('首选比分', '0-0', '3-3', '[错][错][错]'),
            ('方向', '平局(默契球)', '平局(生死战)', '巧合!逻辑全错'),
            ('核心逻辑', 'J2=西班牙→都不想赢', '3分不够→必须争胜', '[错][错][错]'),
        ],
        'factors': [
            ['J2=西班牙→双方都不想第2', '双方全力争胜!奥地利赢仍对西班牙', '[错]'],
            ['0-0默契球预判', '28\'即破僵,全场6球', '[错]'],
            ['阿瑙托维奇首发', '28\'进球!关键球员', '[对]'],
            ['马赫雷斯能力', '梅开二度!含90+3\'看似绝杀', '[对]'],
        ],
        'lessons': [
            '[致命]出线数学必须精确: 赛前"都不想赢"前提=输球J3=4分——但输球=3分! 3分大概率淘汰。',
            '3分≠安全、4分≠保险: 48队取8个第3名,第3轮所有球队底线是4分。',
            '预测0-0→实际3-3=本日最大失准: 生存本能压倒选对手算计。',
        ],
    },
    {
        'num': 6, 'title': '约旦 1-3 阿根廷 (J组)',
        'goals': [('19\'', '洛塞尔索 (LO CELSO)', '阿根廷'), ('31\'', '劳塔罗·马丁内斯 (Lautaro MARTINEZ)', '阿根廷'), ('55\'', '阿尔塔马里 (AL-TAMARI)', '约旦'), ('80\'', '梅西 (Lionel MESSI)', '阿根廷')],
        'events': '半场 约旦 0-2 阿根廷 | 黄牌: 约旦17\'/64\'/90+4\' | 换人: 梅西60\'入→80\'任意球破门=第6球!',
        'pred_vs_actual': [
            ('首选比分', '0-3', '1-3', '[错]差1球'),
            ('方向', '阿根廷赢', '阿根廷赢', '[对]'),
            ('半场', '0-2', '0-2', '[对]'),
        ],
        'factors': [
            ['身价比41:1→碾压', '1-3全程控制', '[对]'],
            ['阿根廷可能大幅轮换', '9人轮换!梅西替补', '[对]'],
            ['梅西替补出场刷1球', '60\'上场→80\'任意球=第6球!', '[对][对]'],
            ['约旦已淘汰无抵抗', '55\'扳回1球!3场均进球', '[半错]'],
        ],
        'lessons': [
            '约旦韧性被低估: 0-2落后仍扳回1球。小组赛3场全部进球。',
            '梅西替补=阿根廷奢侈: 60\'出场80\'破门。第6球领跑金靴,连续7场世界杯进球创历史。',
            '阿根廷9人轮换仍19\'破僵: 阵容深度世界顶级,淘汰赛提供极大弹性。',
        ],
    },
]

# ── Render ──
for m in matches:
    add_heading_styled(f'比赛{m["num"]}: {m["title"]}', level=2)

    # ─ Goals table ─
    if m['goals']:
        add_para('进球记录', bold=True, size=Pt(10), color=DARK)
        gdata = [['时间', '球员', '进球方']]
        gdata.extend(m['goals'])
        gt = doc.add_table(rows=len(gdata), cols=3)
        gt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, text in enumerate(gdata[0]):
            set_cell(gt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
        for i, row in enumerate(gdata[1:], 1):
            bg = LIGHT_GRAY if i % 2 == 0 else None
            set_cell(gt.rows[i].cells[0], row[0], size=Pt(9), bg=bg)
            set_cell(gt.rows[i].cells[1], row[1], size=Pt(9), bg=bg)
            set_cell(gt.rows[i].cells[2], row[2], size=Pt(9), bg=bg)
        for row in gt.rows:
            row.cells[0].width = Inches(0.6)
            row.cells[1].width = Inches(2.5)
            row.cells[2].width = Inches(1.0)
        doc.add_paragraph()

    # ─ Events ─
    add_para(m['events'], size=Pt(8.5), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(6))

    # ─ Pred vs Actual ─
    add_para('原始预测 vs 实际', bold=True, size=Pt(10), color=DARK)
    pdata = [['维度', '预测', '实际', '评估']]
    pdata.extend(m['pred_vs_actual'])
    pt = doc.add_table(rows=len(pdata), cols=4)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(pdata[0]):
        set_cell(pt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(pdata[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell(pt.rows[i].cells[0], row[0], bold=True, size=Pt(8.5), bg=bg)
        set_cell(pt.rows[i].cells[1], row[1], size=Pt(8.5), bg=bg)
        set_cell(pt.rows[i].cells[2], row[2], bold=True, size=Pt(9), bg=bg)
        ev = row[3]
        clr = GREEN_BG if '[对]' in ev and '[错]' not in ev else (ACCENT if '[错]' in ev else None)
        set_cell(pt.rows[i].cells[3], ev, bold=True, size=Pt(8.5), color=clr, bg=bg)
    for row in pt.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(1.5)
        row.cells[3].width = Inches(2.0)
    doc.add_paragraph()

    # ─ Factor checks ─
    add_para('赛前因素验证', bold=True, size=Pt(10), color=DARK)
    fdata = [['预测因素', '实际结果', '验证']]
    fdata.extend(m['factors'])
    ft = doc.add_table(rows=len(fdata), cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(fdata[0]):
        set_cell(ft.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
    for i, row in enumerate(fdata[1:], 1):
        bg = LIGHT_GRAY if i % 2 == 0 else None
        set_cell_left(ft.rows[i].cells[0], row[0], size=Pt(8.5))
        set_cell_left(ft.rows[i].cells[1], row[1], size=Pt(8.5))
        v = row[2]
        clr = GREEN_BG if v.startswith('[对]') else (ACCENT if v.startswith('[错]') else None)
        set_cell(ft.rows[i].cells[2], v, bold=True, size=Pt(8.5), color=clr, bg=bg)
    for row in ft.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(0.9)
    doc.add_paragraph()

    # ─ Lessons ─
    add_para('[关键教训]', bold=True, size=Pt(10), color=ACCENT)
    for idx, lesson in enumerate(m['lessons'], 1):
        add_para(f'{idx}. {lesson}', size=Pt(9), color=RGBColor(0x44,0x44,0x44), space_after=Pt(2))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. SUMMARY ANALYSIS
# ═══════════════════════════════════════════════════
add_heading_styled('六场汇总洞察', level=2)

# accuracy
add_para('方向准确率', bold=True, size=Pt(10), color=DARK)
ddata = [['指标', '数值'],
         ['方向正确', '4/6 (67%)'],
         ['方向错误', '2/6 (33%) — 比赛1(1-1→2-1)、比赛5(0-0→3-3逻辑全错)'],
         ['精确比分', '1/6 (17%, 含备选)'],
         ['半场正确', '2/6 (33%) — 比赛3(0-0)、比赛6(0-2)']]
dt = doc.add_table(rows=len(ddata), cols=2)
dt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(ddata[0]):
    set_cell(dt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(ddata[1:], 1):
    set_cell(dt.rows[i].cells[0], row[0], bold=True, size=Pt(9))
    set_cell_left(dt.rows[i].cells[1], row[1], size=Pt(9))
for row in dt.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.0)
doc.add_paragraph()

# dimensions
add_para('被高估/低估的维度', bold=True, size=Pt(10), color=DARK)
mdata = [['维度', '趋势', '典型案例'],
         ['第3轮出线数学', '[致命]错误', '比赛5: 3分≠安全,"双方都不想赢"逻辑完全错误'],
         ['已淘汰球队威胁', '被低估', '比赛4/6: 乌兹别克10\'进球 + 约旦回敬1球'],
         ['轮换影响', '被低估', '比赛2: 英格兰上半场0-0'],
         ['"双方都可接受平局"', '判断准确', '比赛3: 哥伦比亚0-0葡萄牙完全验证'],
         ['大巴持久力', '被高估', '比赛1: 加纳73\'扳平→83\'被绝杀'],
         ['核心缺席对称性', '被高估', '比赛1: 克罗地亚替代者 > 加纳替代者'],
         ['无淘汰球队下半场崩盘', '新发现', '比赛4: 乌兹别克上半场1-0→下半场连丢3球']]
mt = doc.add_table(rows=len(mdata), cols=3)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(mdata[0]):
    set_cell(mt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(mdata[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell_left(mt.rows[i].cells[0], row[0], bold=True, size=Pt(8.5), bg=bg)
    clr = ACCENT if '致命' in row[1] else None
    set_cell(mt.rows[i].cells[1], row[1], size=Pt(8.5), color=clr, bg=bg)
    set_cell_left(mt.rows[i].cells[2], row[2], size=Pt(8.5), bg=bg)
for row in mt.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(1.2)
    row.cells[2].width = Inches(4.0)
doc.add_paragraph()

# new rules
add_para('值得记录的新规律', bold=True, size=Pt(10), color=DARK)
rules = [
    '1. 第3轮"双方已出线+打平双赢"=0-0概率极高 — 哥伦比亚0-0葡萄牙 [对]',
    '2. 大幅轮换→上半场大概率0-0 — 巴拿马0-0英格兰(半场) [对]',
    '3. 已淘汰球队→上半场可能制造惊喜→下半场崩盘 [对]',
    '4. 大巴持久力≠全场→70\'后疲劳导致防守失误增加 [对]',
    '5. [新] 第3轮出线数学必须逐组精确计算,不能凭感觉 — 3分=大概率淘汰',
    '6. [新] "不想对强队" < "不能小组淘汰" — 生存本能压倒选对手算计',
]
for r in rules:
    clr = ACCENT if '[新]' in r else RGBColor(0x44,0x44,0x44)
    add_para(r, size=Pt(9), color=clr, space_after=Pt(2))
doc.add_paragraph()

# best/worst
add_heading_styled('本日最佳/最差', level=2)
bdata = [['奖项', '得主', '理由'],
         ['最佳预测', '巴拿马 0-2 英格兰', '备选比分完全命中,方向+走势正确'],
         ['最差预测', '阿尔及利亚 0-0 奥地利', '逻辑全错(0-0→3-3),出线数学算错'],
         ['最佳比赛', '阿尔及利亚 3-3 奥地利', '90+3\'→90+6\'绝平,本日最具戏剧性'],
         ['最佳球员', '马赫雷斯(阿尔及利亚)', '梅开二度,含90+3\'看似绝杀'],
         ['最佳换人', '卡拉季奇(奥地利)', '90+5\'换入→90+6\'头球绝平'],
         ['最大惊喜', '梅西替补任意球', '连续7场世界杯进球,第6球领跑金靴']]
bt = doc.add_table(rows=len(bdata), cols=3)
bt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, text in enumerate(bdata[0]):
    set_cell(bt.rows[0].cells[j], text, bold=True, size=Pt(9), color=WHITE, bg=HEADER_BG)
for i, row in enumerate(bdata[1:], 1):
    bg = LIGHT_GRAY if i % 2 == 0 else None
    set_cell_left(bt.rows[i].cells[0], row[0], bold=True, size=Pt(8.5), bg=bg)
    set_cell_left(bt.rows[i].cells[1], row[1], size=Pt(8.5), bg=bg)
    set_cell_left(bt.rows[i].cells[2], row[2], size=Pt(8.5), bg=bg)
for row in bt.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(3.8)

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('───', size=Pt(8), color=RGBColor(0xBD,0xBD,0xBD), align='center')
add_para('数据源: FIFA Official API (https://api.fifa.com/api/v3/live/football/) + ESPN API  |  框架: CLAUDE.md v17 复盘规范',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('原始预测: 2026年6月28日_6场预测.md (v7, FIFA首发确认版)  |  半场分析: 6月28日_10点两场半场与预测.md',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')
add_para('生成工具: python-docx  |  仅供分析参考  |  [对]=正确 [错]=错误 [半对]=部分正确',
         size=Pt(7), color=RGBColor(0xAA,0xAA,0xAA), align='center')

doc.save(OUT)
print(f'Saved to {OUT}')
print(f'   File size: {os.path.getsize(OUT):,} bytes')
